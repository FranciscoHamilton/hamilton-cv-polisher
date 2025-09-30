# app.py
import os, json, re, tempfile, traceback, zipfile, io
from pathlib import Path
from datetime import datetime, timedelta
from flask import Flask, request, send_file, render_template_string, abort, jsonify, make_response
from flask import session, redirect, url_for  # <-- ADDED earlier
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import HTTPException
def is_admin() -> bool:
    """Return True if the logged-in session user matches APP_ADMIN_USER."""
    try:
        return (session.get("user") or "").lower() == (os.getenv("APP_ADMIN_USER") or "").lower()
    except Exception:
        return False
# --- Database (Postgres via psycopg2) ---
import psycopg2
from psycopg2.pool import SimpleConnectionPool

DATABASE_URL = os.getenv("DATABASE_URL", "").strip()
DB_POOL = None
if DATABASE_URL:
    try:
        DB_POOL = SimpleConnectionPool(minconn=1, maxconn=5, dsn=DATABASE_URL)
        print("DB pool initialized")
    except Exception as e:
        print("DB pool init failed:", e)
        DB_POOL = None

INIT_SQL = """
CREATE TABLE IF NOT EXISTS users (
  id SERIAL PRIMARY KEY,
  username TEXT UNIQUE NOT NULL,
  password_hash TEXT NOT NULL,
  email TEXT,
  company TEXT,
  stripe_customer_id TEXT,
  active BOOLEAN DEFAULT TRUE,
  created_at TIMESTAMP DEFAULT NOW()
);

-- Orgs + org_id on users
CREATE TABLE IF NOT EXISTS orgs (
  id SERIAL PRIMARY KEY,
  name TEXT UNIQUE NOT NULL,
  active BOOLEAN DEFAULT TRUE,
  created_at TIMESTAMP DEFAULT NOW()
);
ALTER TABLE users ADD COLUMN IF NOT EXISTS org_id INTEGER;

-- Per-user usage events (+ optional org_id)
CREATE TABLE IF NOT EXISTS usage_events (
  id SERIAL PRIMARY KEY,
  user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
  filename TEXT,
  candidate TEXT,
  ts TIMESTAMP DEFAULT NOW()
);
ALTER TABLE usage_events ADD COLUMN IF NOT EXISTS org_id INTEGER;

-- Per-user credits ledger (kept for history; also stores org_id when known)
CREATE TABLE IF NOT EXISTS credits_ledger (
  id SERIAL PRIMARY KEY,
  user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
  delta INTEGER NOT NULL,
  reason TEXT,
  ext_ref TEXT,
  org_id INTEGER,
  created_at TIMESTAMP DEFAULT NOW()
);

-- NEW: Org-wide credits pool (one shared balance per org)
CREATE TABLE IF NOT EXISTS org_credits_ledger (
  id SERIAL PRIMARY KEY,
  org_id INTEGER NOT NULL,
  delta INTEGER NOT NULL,
  reason TEXT,
  created_by INTEGER,
  created_at TIMESTAMP DEFAULT NOW()
);

-- NEW: Optional per-user monthly caps (within an org)
CREATE TABLE IF NOT EXISTS org_user_limits (
  id SERIAL PRIMARY KEY,
  org_id INTEGER NOT NULL,
  user_id INTEGER NOT NULL,
  monthly_cap INTEGER,
  active BOOLEAN DEFAULT TRUE,
  created_at TIMESTAMP DEFAULT NOW()
);
-- Backfill 'active' columns for older databases (safe to run anytime)
ALTER TABLE users           ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE;
ALTER TABLE orgs            ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE;
ALTER TABLE org_user_limits ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE;

-- Leads (contact / trial requests)
CREATE TABLE IF NOT EXISTS leads (
  id          BIGSERIAL PRIMARY KEY,
  created_at  TIMESTAMP DEFAULT NOW(),
  company     TEXT,
  email       TEXT,
  name        TEXT,
  volume      TEXT,
  users       TEXT,
  templates   TEXT,
  need_sso    TEXT,
  message     TEXT,
  filename    TEXT,
  ip          TEXT,
  user_agent  TEXT
);

-- Helpful indexes (idempotent)
CREATE INDEX IF NOT EXISTS idx_users_org_id           ON users(org_id);
CREATE INDEX IF NOT EXISTS idx_usage_month_user       ON usage_events(user_id, ts);
CREATE INDEX IF NOT EXISTS idx_usage_org_id           ON usage_events(org_id);
CREATE INDEX IF NOT EXISTS idx_cred_user              ON credits_ledger(user_id);
CREATE INDEX IF NOT EXISTS idx_cred_org               ON credits_ledger(org_id);
CREATE INDEX IF NOT EXISTS idx_orgcred_org            ON org_credits_ledger(org_id);
CREATE INDEX IF NOT EXISTS idx_orglimits_org_user     ON org_user_limits(org_id, user_id);
CREATE INDEX IF NOT EXISTS idx_orglimits_active       ON org_user_limits(active);

-- Seed a default org (id=1) if you want Hamilton as org 1
INSERT INTO orgs (id, name, active)
VALUES (1, 'Hamilton', TRUE)
ON CONFLICT (id) DO NOTHING;
"""

def init_db():
    """Create tables if they don't exist. Safe to run on every boot."""
    if not DB_POOL:
        print("No DATABASE_URL set; skipping DB init.")
        return
    conn = DB_POOL.getconn()
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(INIT_SQL)
        print("DB init OK")
    except Exception as e:
        print("DB init failed:", e)
    finally:
        DB_POOL.putconn(conn)


# --- Small DB helpers ---
def db_conn():
    """Get a DB connection from the pool (or None if DB unused)."""
    if not DB_POOL:
        return None
    return DB_POOL.getconn()

def db_put(conn):
    """Return a DB connection to the pool safely."""
    if DB_POOL and conn:
        DB_POOL.putconn(conn)

def db_query_one(sql, params=()):
    """Run a SELECT that returns one row (as a tuple) or None."""
    conn = db_conn()
    if not conn:
        return None
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, params)
                row = cur.fetchone()
                return row
    except Exception as e:
        print("db_query_one error:", e)
        return None
    finally:
        db_put(conn)
def db_query_all(sql, params=()):
    """Run a SELECT that returns many rows (list of tuples)."""
    conn = db_conn()
    if not conn:
        return []
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, params)
                rows = cur.fetchall()
                return rows
    except Exception as e:
        print("db_query_all error:", e)
        return []
    finally:
        db_put(conn)
def db_execute(sql, params=()):
    """Run an INSERT/UPDATE/DELETE. Returns True/False."""
    conn = db_conn()
    if not conn:
        return False
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, params)
        return True
    except Exception as e:
        print("db_execute error:", e)
        return False
    finally:
        db_put(conn)

def seed_admin_user():
    """
    Ensure the env admin exists in Postgres with a hashed password.
    Uses APP_ADMIN_USER / APP_ADMIN_PASS for initial seed.
    Safe to run on every boot.
    """
    usr = os.getenv("APP_ADMIN_USER", "admin").strip()
    pw  = os.getenv("APP_ADMIN_PASS", "hamilton")
    if not usr or not pw:
        return
    # already exists?
    row = db_query_one("SELECT id FROM users WHERE username=%s", (usr,))
    if row:
        return
    # create it
    pw_hash = generate_password_hash(pw)
    ok = db_execute(
        "INSERT INTO users (username, password_hash, email, company, active) VALUES (%s,%s,%s,%s,%s)",
        (usr, pw_hash, None, "Hamilton Recruitment", True)
    )
    if ok:
        print(f"Seeded admin user in DB: {usr}")
    else:
        print("Failed to seed admin user in DB")

def get_user_db(username: str):
    """Return a dict for the DB user or None."""
    row = db_query_one(
        "SELECT id, username, password_hash, active FROM users WHERE username=%s",
        (username.strip(),)
    )
    if not row:
        return None
    return {
        "id": row[0],
        "username": row[1],
        "password_hash": row[2],
        "active": bool(row[3]),
    }
def get_user_plan_credits_and_overage(user_id: int):
    """
    Return (plan_credits, overage_rate) for this user from Postgres.
    If no plan is active or DB is unavailable, returns (0, 0.0).
    """
    if not user_id:
        return (0, 0.0)
    try:
        row = db_query_one(
            """
            SELECT monthly_credits, overage_rate
              FROM plans
             WHERE user_id = %s
               AND active = TRUE
             ORDER BY created_at DESC
             LIMIT 1
            """,
            (user_id,),
        )
        if not row:
            return (0, 0.0)
        return (int(row[0] or 0), float(row[1] or 0.0))
    except Exception as e:
        print("get_user_plan_credits_and_overage error:", e)
        return (0, 0.0)

# --- Per-user usage helpers (Postgres) ---
def get_user_month_usage(user_id: int) -> int:
    """
    Count usage events for this user in the current calendar month.
    Falls back to 0 if query fails.
    """
    row = db_query_one(
        """
        SELECT COUNT(*) FROM usage_events
         WHERE user_id = %s
           AND date_trunc('month', ts) = date_trunc('month', now())
        """,
        (user_id,),
    )
    try:
        return int(row[0]) if row and row[0] is not None else 0
    except Exception:
        return 0


def count_usage_this_month(user_id):
    """
    How many events this calendar month for this user.
    """
    if not DB_POOL or not user_id:
        return 0
    row = db_query_one("""
        SELECT COUNT(*)
        FROM usage_events
        WHERE user_id=%s
          AND date_trunc('month', ts) = date_trunc('month', now())
    """, (user_id,))
    return int(row[0]) if row and row[0] is not None else 0

def last_event_for_user(user_id):
    """
    Return (candidate, ts_str) for the most recent event of this user, else (None, None).
    """
    if not DB_POOL or not user_id:
        return (None, None)
    row = db_query_one("""
        SELECT candidate, to_char(ts, 'YYYY-MM-DD HH24:MI:SS')
        FROM usage_events
        WHERE user_id=%s
        ORDER BY ts DESC
        LIMIT 1
    """, (user_id,))
    if not row:
        return (None, None)
    return (row[0] or None, row[1] or None)

def log_usage_event(user_id: int, filename: str, candidate: str) -> bool:
    """
    Insert a usage_events row for this user.
    If the user has an org_id, store it too.
    """
    try:
        uid = int(user_id or 0)
    except Exception:
        uid = 0
    if not (DB_POOL and uid):
        return False
    try:
        # sanitize a bit
        fn = (filename or "")[:200]
        cand = (candidate or "")[:200]

        # get org (if any)
        row = db_query_one("SELECT org_id FROM users WHERE id=%s", (uid,))
        oid = int(row[0]) if row and row[0] is not None else None

        if oid:
            return db_execute(
                "INSERT INTO usage_events (user_id, ts, candidate, filename, org_id) VALUES (%s, now(), %s, %s, %s)",
                (uid, cand, fn, oid),
            )
        else:
            return db_execute(
                "INSERT INTO usage_events (user_id, ts, candidate, filename) VALUES (%s, now(), %s, %s)",
                (uid, cand, fn),
            )
    except Exception as e:
        # don't break the app if DB insert fails
        print("log_usage_event failed:", e)
        return False

def credits_add(user_id: int, delta: int, reason: str = "polish", ext_ref: str = "") -> bool:
    """
    Append a row to credits_ledger for this user (positive = grant, negative = charge).
    If the user has an org_id, store it too.
    """
    try:
        uid = int(user_id or 0)
    except Exception:
        uid = 0
    if not (DB_POOL and uid):
        return False
    try:
        # look up org for this user (if any)
        row = db_query_one("SELECT org_id FROM users WHERE id=%s", (uid,))
        oid = int(row[0]) if row and row[0] is not None else None

        # sanitize
        d = int(delta)
        r = (reason or "")[:50]
        x = (ext_ref or "")[:200]

        if oid:
            return db_execute(
                "INSERT INTO credits_ledger (user_id, delta, reason, ext_ref, org_id) VALUES (%s,%s,%s,%s,%s)",
                (uid, d, r, x, oid),
            )
        else:
            return db_execute(
                "INSERT INTO credits_ledger (user_id, delta, reason, ext_ref) VALUES (%s,%s,%s,%s)",
                (uid, d, r, x),
            )
    except Exception as e:
        print("credits_add failed:", e)
        return False

def count_usage_month_db(user_id: int) -> int:
    """
    Count usage_events for this user in the current calendar month.
    Returns 0 on any error.
    """
    row = db_query_one(
        """
        SELECT COUNT(*) FROM usage_events
         WHERE user_id = %s
           AND date_trunc('month', ts) = date_trunc('month', now())
        """,
        (user_id,),
    )
    try:
        return int(row[0]) if row and row[0] is not None else 0
    except Exception:
        return 0
def list_users_usage_month():
    """
    Return a list of dicts:
      {'id', 'username', 'active', 'month_usage', 'total_usage'}
    Reads from Postgres. Returns [] if DB is missing or on error.
    """
    sql = """
      SELECT
        u.id,
        u.username,
        COALESCE(u.active, TRUE) AS active,
        COALESCE(SUM(CASE
          WHEN date_trunc('month', e.ts) = date_trunc('month', now()) THEN 1
          ELSE 0
        END), 0) AS month_usage,
        COALESCE(COUNT(e.id), 0) AS total_usage
      FROM users u
      LEFT JOIN usage_events e
        ON e.user_id = u.id
      GROUP BY u.id, u.username, u.active
      ORDER BY LOWER(u.username)
    """
    conn = db_conn()
    if not conn:
        return []
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql)
                rows = cur.fetchall()
        # rows: [(id, username, active, month_usage, total_usage), ...]
        out = []
        for r in rows:
            out.append({
                "id": r[0],
                "username": r[1],
                "active": bool(r[2]),
                "month_usage": int(r[3] or 0),
                "total_usage": int(r[4] or 0),
            })
        return out
    except Exception as e:
        print("list_users_usage_month error:", e)
        return []
    finally:
        db_put(conn)

def get_recent_usage_events(limit: int = 50):
    """
    Return the most recent usage events joined with usernames.
    Shape: [{"ts": "YYYY-MM-DD HH:MM:SS", "user_id": int, "username": str, "filename": str, "candidate": str}, ...]
    Safe no-op: returns [] if DB is missing or on error.
    """
    # Normalize and clamp the limit
    try:
        limit = int(limit)
    except Exception:
        limit = 50
    limit = max(1, min(limit, 500))

    conn = db_conn()
    if not conn:
        return []

    sql = """
        SELECT
            e.ts,
            e.user_id,
            COALESCE(u.username, '(unknown)') AS username,
            e.filename,
            e.candidate
        FROM usage_events e
        LEFT JOIN users u ON u.id = e.user_id
        ORDER BY e.ts DESC
        LIMIT %s
    """

    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, (limit,))
                rows = cur.fetchall()
        out = []
        for ts, user_id, username, filename, candidate in rows:
            # Make ts a consistent string
            ts_str = ts.isoformat(sep=" ", timespec="seconds") if hasattr(ts, "isoformat") else str(ts)
            out.append({
                "ts": ts_str,
                "user_id": int(user_id) if user_id is not None else None,
                "username": username,
                "filename": filename or "",
                "candidate": candidate or "",
            })
        return out
    except Exception as e:
        print("get_recent_usage_events error:", e)
        return []
    finally:
        try:
            db_put(conn)
        except Exception:
            pass
# Try fast PDF extraction first (PyMuPDF)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# Text extraction / DOCX tooling
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document as Docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = Path(__file__).parent.resolve()
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # used only if OPENAI_API_KEY is set

# --- simple persistent stats + history ---
STATS_FILE = PROJECT_DIR / "stats.json"
if STATS_FILE.exists():
    try:
        STATS = json.loads(STATS_FILE.read_text(encoding="utf-8"))
    except Exception:
        STATS = {"downloads": 0, "last_candidate": "", "last_time": "", "history": []}
else:
    STATS = {"downloads": 0, "last_candidate": "", "last_time": "", "history": []}
STATS.setdefault("history", [])
# NEW: credits bucket for director view (does not change polish behavior)
STATS.setdefault("credits", {"balance": 0, "purchased": 0})
STATS.setdefault("plan", {"name": "", "credits": 0})


def _save_stats():
    if len(STATS.get("history", [])) > 1000:
        STATS["history"] = STATS["history"][-1000:]
    STATS_FILE.write_text(json.dumps(STATS, indent=2), encoding="utf-8")

# NEW: simple users store (for recruiters you create in Director)
USERS_FILE = PROJECT_DIR / "users.json"
if USERS_FILE.exists():
    try:
        USERS_DB = json.loads(USERS_FILE.read_text(encoding="utf-8"))
    except Exception:
        USERS_DB = {"users": []}
else:
    USERS_DB = {"users": []}

def _save_users():
    USERS_FILE.write_text(json.dumps(USERS_DB, indent=2), encoding="utf-8")

def _get_user(u):
    for x in USERS_DB.get("users", []):
        if (x.get("username") or "").lower() == (u or "").lower():
            return x
    return None

# --- tiny trial request logger ---
TRIALS_FILE = PROJECT_DIR / "trials.json"

def _log_trial(data: dict):
    try:
        if TRIALS_FILE.exists():
            buf = json.loads(TRIALS_FILE.read_text(encoding="utf-8"))
        else:
            buf = []
        # keep it light; don’t store anything sensitive
        buf.append({
            "ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "company": (data.get("company") or "")[:200],
            "email": (data.get("email") or "")[:200],
            "name": (data.get("name") or "")[:200],
            "team_size": (data.get("team_size") or "")[:50],
            "notes": (data.get("notes") or "")[:1000],
        })
        # keep last 1000 only
        buf = buf[-1000:]
        TRIALS_FILE.write_text(json.dumps(buf, indent=2), encoding="utf-8")
    except Exception:
        # don’t break the flow if logging fails
        pass

# alias so either name works (covers older calls)
def _log_trial_request(data: dict):
    _log_trial(data)

# --- tiny trial sign-up logger (local JSON file) ---
TRIALS_FILE = PROJECT_DIR / "trials.json"
if TRIALS_FILE.exists():
    try:
        TRIALS = json.loads(TRIALS_FILE.read_text(encoding="utf-8"))
    except Exception:
        TRIALS = []
else:
    TRIALS = []

def _log_trial(entry: dict):
    TRIALS.append(entry)
    TRIALS_FILE.write_text(json.dumps(TRIALS, indent=2), encoding="utf-8")

# ------------------------ Public Home ------------------------
HOMEPAGE_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
  /* ===== Consistent site header (brand + right links) across all pages ===== */
:root{
  --site-max: 1200px;
  --site-pad-x: 24px;
  --site-pad-y: 16px;
}
.sitebar, .topbar, header .sitebar, header .topbar{
  max-width: var(--site-max);
  margin: 0 auto;
  padding: var(--site-pad-y) var(--site-pad-x);
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}
.sitebar .brand, .topbar .brand{
  display:flex; align-items:center; gap:10px;
  font-weight: 900; letter-spacing: .02em; font-size: 22px;
}
.sitebar nav, .topbar .nav{
  display:flex; align-items:center; gap:24px;
}
.sitebar a, .topbar a{
  text-decoration:none; font-weight:800;
}
    :root{
  --blue:#2563eb;      /* vivid indigo */
  --blue-2:#22d3ee;    /* bright cyan  */
  --ink:#0f172a; --muted:#5b677a; --line:#e5e7eb;
  --bg:#f5f8fd; --card:#ffffff; --shadow: 0 10px 28px rgba(13,59,102,.08);
}
    *{box-sizing:border-box}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}

    .wrap{max-width:1100px;margin:24px auto 64px;padding:0 20px}

    /* top nav */
    .nav{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}
    .brand{font-weight:900;color:var(--black);text-decoration:none;font-size:22px;letter-spacing:.2px}
    .nav a{color:var(--black);text-decoration:none;font-weight:800;margin-left:22px}

    /* hero */
    .hero{background:var(--card);border:1px solid var(--line);border-radius:22px;padding:28px;box-shadow:var(--shadow)}
    .kicker{font-size:12px;letter-spacing:.14em;font-weight:900;color:var(--blue);margin-bottom:10px}
    h1{font-size:48px;line-height:1.05;letter-spacing:-.01em;margin:0 0 10px}
    .lead{font-size:17px;color:var(--muted);max-width:720px;margin:6px 0 16px}
    .actions{display:flex;gap:12px;flex-wrap:wrap;margin-top:8px}
    .btn{display:inline-block;padding:12px 16px;border-radius:12px;font-weight:800;text-decoration:none}
    .btn.primary{background:linear-gradient(90deg,var(--blue),var(--blue-2));color:#fff}
    .btn.secondary{background:#fff;border:1px solid var(--line);color:var(--blue)}
    .meta{color:var(--muted);font-size:13px;margin-top:8px}

    /* features */
    .grid3{display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin-top:18px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:18px;padding:16px}
    .card h3{margin:0 0 6px;font-size:18px;color:var(--blue)}
    .card p{margin:0;color:var(--muted);font-size:14px}

    /* stepper */
    .stepper{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-top:18px}
    .step{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:12px;display:flex;align-items:center;gap:10px}
    .b{display:inline-flex;align-items:center;justify-content:center;width:28px;height:28px;border-radius:999px;border:1px solid var(--line);font-weight:900;color:var(--blue)}

    @media (max-width:900px){
      h1{font-size:36px}
      .lead{font-size:16px}
      .grid3,.stepper{grid-template-columns:1fr}
    }
  </style>
</head>
<body>
  <header>
  <div class="sitebar">
    <a class="brand" href="/">Lustra</a>
    <nav>
      <a href="/about">About</a>
      <a href="/pricing">Pricing</a>
      <a href="/login">Sign in</a>
    </nav>
  </div>
</header>

    <div class="hero">
      <div class="kicker">BUILT BY RECRUITERS, FOR RECRUITERS</div>
      <h1>Client-ready CVs.<br/>On your brand.<br/>In seconds.</h1>
      <p class="lead">Lustra turns that 10–20 minute task into seconds!</p>
      <div class="actions">
        <a class="btn primary" href="/start">Contact Us</a>
        <a class="btn secondary" href="/login">Sign in</a>
      </div>
      <div class="meta">Custom build per client · Keep your look · Fully tailored CVs</div>
    </div>

    <div class="grid3">
      <div class="card">
        <h3>On-brand output</h3>
        <p>Your logo, fonts and spacing. Clean and consistent across the team.</p>
      </div>
      <div class="card">
        <h3>No learning curve</h3>
        <p>Upload → Download. 15 minutes saved per CV, every time.</p>
      </div>
      <div class="card">
        <h3>Built for recruiters</h3>
        <p>No invented facts. We only structure what’s in the candidate’s CV.</p>
      </div>
    </div>

    <div class="stepper">
      <div class="step"><span class="b">1</span>Upload a CV</div>
      <div class="step"><span class="b">2</span>We extract &amp; structure</div>
      <div class="step"><span class="b">3</span>Download polished DOCX</div>
    </div>
  </div>
</body>
</html>
"""
DIRECTOR_HTML = r"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Director – Usage</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    body { font: 14px/1.4 -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif; margin: 20px; }
    h1 { margin: 0 0 12px; }
    .wrap { display: grid; grid-template-columns: 1fr; gap: 24px; }
    @media (min-width: 1000px) { .wrap { grid-template-columns: 1fr 1fr; } }
    .card { border: 1px solid #e5e7eb; border-radius: 10px; padding: 16px; background: #fff; box-shadow: 0 1px 2px rgba(0,0,0,0.03); }
    table { width: 100%; border-collapse: collapse; }
    th, td { text-align: left; padding: 8px; border-bottom: 1px solid #f1f5f9; }
    th { background: #f8fafc; position: sticky; top: 0; }
    .muted { color: #64748b; }
    .pill { display: inline-block; padding: 2px 8px; border-radius: 999px; font-size: 12px; background: #f1f5f9; }
    .topbar { display: flex; justify-content: space-between; align-items: center; margin-bottom: 16px; }
    a.btn { text-decoration: none; border: 1px solid #e5e7eb; padding: 8px 10px; border-radius: 8px; color: #0f172a; }
  </style>
</head>
<body>
  <div class="topbar">
    <h1>Director – Usage</h1>
    <div>
      <a class="btn" href="/app">← Back to App</a>
    </div>
  </div>

  <div class="wrap">
    <div class="card">
      <h2>Per-user totals <span class="pill">{{ users|length }} users</span></h2>
      {% if users and users|length > 0 %}
      <table>
        <thead>
          <tr>
            <th>User</th>
            <th class="muted">Active</th>
            <th>This month</th>
            <th>Total</th>
          </tr>
        </thead>
        <tbody>
          {% for u in users %}
          <tr>
            <td>{{ u.username }}</td>
            <td class="muted">{{ 'Yes' if u.active else 'No' }}</td>
            <td>{{ u.month_usage }}</td>
            <td>{{ u.total_usage }}</td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
      {% else %}
        <div class="muted">No users found (DB offline or empty).</div>
      {% endif %}
    </div>

    <div class="card">
      <h2>Recent usage events <span class="pill">{{ events|length }}</span></h2>
      {% if events and events|length > 0 %}
      <table>
        <thead>
          <tr>
            <th>When</th>
            <th>User</th>
            <th>Candidate</th>
            <th class="muted">File</th>
          </tr>
        </thead>
        <tbody>
          {% for e in events %}
          <tr>
            <td>{{ e.ts }}</td>
            <td>{{ e.username }}</td>
            <td>{{ e.candidate }}</td>
            <td class="muted">{{ e.filename }}</td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
      {% else %}
        <div class="muted">No events yet (or DB offline).</div>
      {% endif %}
    </div>
        <div class="card">
      <h2>Legacy JSON history <span class="pill">{{ legacy|length }}</span></h2>
      {% if legacy and legacy|length > 0 %}
      <table>
        <thead>
          <tr>
            <th>When</th>
            <th>Candidate</th>
            <th class="muted">File</th>
          </tr>
        </thead>
        <tbody>
          {% for h in legacy|reverse %}
          <tr>
            <td>{{ h.ts }}</td>
            <td>{{ h.candidate }}</td>
            <td class="muted">{{ h.filename }}</td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
      {% else %}
        <div class="muted">No legacy entries.</div>
      {% endif %}
    </div>
  </div>
</body>
</html>
"""

# ------------------------ About (already added earlier) ------------------------
ABOUT_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>About — Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
  /* ===== Consistent site header (brand + right links) across all pages ===== */
:root{
  --site-max: 1200px;
  --site-pad-x: 24px;
  --site-pad-y: 16px;
}
.sitebar, .topbar, header .sitebar, header .topbar{
  max-width: var(--site-max);
  margin: 0 auto;
  padding: var(--site-pad-y) var(--site-pad-x);
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}
.sitebar .brand, .topbar .brand{
  display:flex; align-items:center; gap:10px;
  font-weight: 900; letter-spacing: .02em; font-size: 22px;
}
.sitebar nav, .topbar .nav{ display:flex; align-items:center; gap:24px; }
.sitebar a, .topbar a{ text-decoration:none; font-weight:800; }
    :root{
  --blue:#2563eb;      /* vivid indigo */
  --blue-2:#22d3ee;    /* bright cyan  */
  --ink:#0f172a; --muted:#5b677a; --line:#e5e7eb;
  --bg:#f5f8fd; --card:#ffffff; --shadow: 0 10px 28px rgba(13,59,102,.08);
}
    *{box-sizing:border-box}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}
    .wrap{max-width:1100px;margin:24px auto 64px;padding:0 20px}

    /* nav like homepage */
    .nav{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}
    .brand{font-weight:900;color:var(--black);text-decoration:none;font-size:22px;letter-spacing:.2px}
    .nav a{color:var(--black);text-decoration:none;font-weight:800;margin-left:22px}

    .card{background:var(--card);border:1px solid var(--line);border-radius:22px;padding:0;box-shadow:var(--shadow)}
    .inner{max-width:780px;padding:24px}
    h1{margin:6px 0 12px;font-size:28px;color:var(--blue)}
h2{margin:18px 0 10px;font-size:18px;color:var(--blue)}
p{margin:8px 0;color:var(--ink);font-size:13.5px;line-height:1.65}
ul{margin:8px 0 16px 20px;color:var(--ink);font-size:13.5px;line-height:1.65}
    .btn{display:inline-block;margin-top:14px;padding:12px 16px;border-radius:12px;background:linear-gradient(90deg,var(--blue),var(--blue-2));border:none;text-decoration:none;font-weight:800;color:#fff}
  </style>
</head>
<body>
  <header>
  <div class="sitebar">
    <a class="brand" href="/">Lustra</a>
    <nav>
      <a href="/about">About</a>
      <a href="/pricing">Pricing</a>
      <a href="/login">Sign in</a>
    </nav>
  </div>
</header>
<div class="wrap">

    <div class="card">
      <div class="inner">
        <h1>Built by recruiters, for recruiters</h1>
        <p>Formatting CVs is necessary—but it’s not why you got into recruitment. After 10+ years running desks and a recruitment business, I’ve felt the pain first-hand: breaking flow to rework a CV, juggling fonts and spacing, fixing headers, and trying to keep branding consistent across the team.</p>
        <p><strong>This tool turns that 10–20 minute task into seconds.</strong> Upload a raw CV (PDF, DOCX, or TXT). We extract the content, structure it, and lay it out in your company’s template. You download a polished, on-brand DOCX—ready to send.</p>

        <h2>Why it matters</h2>
        <ul>
          <li><strong>Time back on the desk:</strong> 15 minutes per CV ≈ 0.25 hours. At £20–£40/hour, that’s £5–£10 per CV. 50 CVs/month ≈ 12.5 hours saved → £250–£500/month in recruiter time.</li>
          <li><strong>Consistency at scale:</strong> every consultant outputs the same, branded format.</li>
          <li><strong>Better candidate & client experience:</strong> clean, readable CVs that reflect your brand.</li>
        </ul>

        <h2>How it works</h2>
        <ul>
          <li><strong>Upload</strong> a raw CV (PDF / DOCX / TXT).</li>
          <li><strong>Extract & structure:</strong> we pull out the real content (experience, education, skills) without inventing facts.</li>
          <li><strong>Lay out in your template:</strong> headers/footers, fonts, sizes and spacing are applied automatically.</li>
          <li><strong>Download</strong> a polished DOCX.</li>
        </ul>

        <h2>Privacy & control</h2>
        <ul>
          <li>No CV content is stored by default—only basic usage metrics (filename + timestamp) for tracking volume and billing.</li>
          <li>Your company template and logo are stored securely to ensure output is always on-brand.</li>
        </ul>

        <h2>What’s on the site</h2>
        <ul>
          <li>Multi-company login (soon): per-company routes and templates.</li>
          <li>Director dashboards (soon): usage counts, CSV export, trends.</li>
          <li>Credit plans: pay-as-you-go or monthly bundles.</li>
          <li>Self-serve template builder (soon): upload a DOCX to switch branding instantly.</li>
        </ul>

        <a class="btn" href="/start">Start free trial</a>
      </div>
    </div>
  </div>
</body>
</html>
"""
  # (unchanged content from your script above)
# For brevity here, keep the exact ABOUT_HTML, PRICING_HTML, HTML, LOGIN_HTML strings from your script.

# ------------------------ Pricing (NEW) ------------------------
PRICING_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Pricing — Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
  /* ===== Consistent site header (brand + right links) across all pages ===== */
:root{
  --site-max: 1200px;
  --site-pad-x: 24px;
  --site-pad-y: 16px;
}
.sitebar, .topbar, header .sitebar, header .topbar{
  max-width: var(--site-max);
  margin: 0 auto;
  padding: var(--site-pad-y) var(--site-pad-x);
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}
.sitebar .brand, .topbar .brand{
  display:flex; align-items:center; gap:10px;
  font-weight: 900; letter-spacing: .02em; font-size: 22px;
}
.sitebar nav, .topbar .nav{ display:flex; align-items:center; gap:24px; }
.sitebar a, .topbar a{ text-decoration:none; font-weight:800; }
    :root{
      --brand:#2563eb; --brand-2:#22d3ee;
      --ink:#0f172a; --muted:#64748b; --line:#e5e7eb;
      --bg:#f6f9ff; --card:#fff; --shadow:0 10px 24px rgba(13,59,102,.08);
      --ok:#16a34a;
      --acc-start:#06b6d4; --acc-growth:#2563eb; --acc-scale:#8b5cf6; --acc-packs:#f59e0b;
    }
    *{box-sizing:border-box}
    html,body{margin:0;padding:0}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink)}

    /* NAV (match Home/About) */
    .wrap{max-width:1120px;margin:0 auto;padding:0 24px}
    .nav{display:flex;align-items:center;justify-content:space-between;padding:18px 0}
    .brand{font-weight:900;color:var(--ink);text-decoration:none;font-size:26px;letter-spacing:.2px}
    .nav-links a{color:var(--black);text-decoration:none;font-weight:800;margin-left:22px}

    /* PAGE BOX (like the white box on other pages) */
    .pagebox{
      background:var(--card);
      border:1px solid var(--line);
      border-radius:26px;
      box-shadow:var(--shadow);
      padding:24px 28px 26px;
      margin:14px 0 12px;   /* adds space under the nav and above the note */
    }

    /* Hero inside box */
    h1{margin:2px 0 8px;font-size:40px;letter-spacing:-.02em;color:#122033}
    .lead{margin:0 0 14px;color:var(--muted);font-size:15px;max-width:880px}

    /* Plans grid INSIDE the box */
    .grid4{display:grid;gap:16px;grid-template-columns:repeat(4,1fr);align-items:stretch}
    @media(max-width:1100px){ .grid4{grid-template-columns:repeat(2,1fr)} }
    @media(max-width:680px){ .grid4{grid-template-columns:1fr} }

    /* Card */
    .card{
      background:var(--card);
      border:1px solid var(--line);
      border-radius:20px;
      box-shadow:var(--shadow);
      overflow:hidden;
      display:flex;flex-direction:column;
      min-height:520px; /* smaller than before */
    }
    .inner{padding:14px 14px 16px;display:flex;flex-direction:column;height:100%}
    .name{font-weight:900;color:#0b1220;font-size:12px;letter-spacing:.08em;margin:0 0 8px}
    .qty{font-size:26px;font-weight:900;letter-spacing:-.01em}
    .per{font-size:13px;color:var(--muted);font-weight:700;margin-left:6px}

    .card.start  .name{color:var(--acc-start)}
    .card.growth .name{color:var(--acc-growth)}
    .card.scale  .name{color:var(--acc-scale)}
    .card.packs  .name{color:var(--acc-packs)}

    /* Price chip */
    .chip{
      display:inline-flex;align-items:baseline;gap:6px;align-self:flex-start;margin-top:8px;
      padding:8px 12px;border-radius:999px;background:#eef4ff;border:1px solid #dbeafe;color:#132a63;
      font-weight:700;font-size:12px;
    }
    .chip .price-month{font-size:1.05em;color:#0b1220;font-weight:800}
    .chip .dot{color:#8aa0c4;font-weight:700;line-height:1}
    .chip .price-cv{font-size:.9em;color:#667792;font-weight:600}

    /* Features */
    .feat{margin:12px 0 0 0;padding:0;list-style:none;color:#475569;font-size:12.5px}
    .feat li{display:flex;align-items:center;gap:8px;margin-top:8px}
    .tick{display:inline-flex;align-items:center;justify-content:center;width:18px;height:18px;border-radius:50%;
      background:rgba(34,211,238,.18);color:#0891b2;font-weight:900;font-size:12px}

    /* Packs block */
    .packs-info{margin:10px 0 0 0;padding:10px;border:1px dashed #dbeafe;border-radius:14px;background:#f8fbff}
    .packs-info .title{font-weight:900;color:#0b1220;font-size:12px;margin-bottom:6px;letter-spacing:.04em}
    .packs-info ul{margin:0;padding-left:18px;font-size:12.5px;color:#334155}
    .packs-info li{margin:4px 0}

    .select{margin-top:10px}
    .select select{width:100%;padding:10px 12px;border:1px solid var(--line);border-radius:12px;background:#fff;font-weight:700}

    /* Buttons closer to content */
    .btn{
  margin-top:auto;
  align-self:center; width:220px;
  padding:10px 14px; border-radius:18px;
  text-align:center; font-weight:900; font-size:14px;
  text-decoration:none; border:1px solid var(--line);
  color:#0b1220; background:#fff;
}
    .btn.primary{background:linear-gradient(90deg,var(--brand),var(--brand-2));color:#fff;border:none}
    .btn:hover{transform:translateY(-1px)}

    /* Note + Calculator spacing (calculator pulled up) */
    .note{color:var(--muted);font-size:12px;margin:6px 0 4px}

    /* Calculator (unchanged UI) */
    .card.calc{min-height:unset}
    .card.calc .inner{align-items:stretch;text-align:left}
    .card.calc .name{font-size:18px;color:var(--brand)}
    .calc-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}
    @media(max-width:980px){ .calc-grid{grid-template-columns:1fr} }
    .calc label{display:block;font-weight:900;margin-bottom:6px}
    .calc input[type=number]{width:100%;padding:12px;border:1px solid var(--line);border-radius:12px;background:#fff;box-shadow:inset 0 1px 2px rgba(2,6,23,.03)}
    .calc-out{display:flex;flex-wrap:wrap;gap:24px;align-items:center;margin-top:12px}
    .calc-out .n{font-weight:900;color:var(--brand);font-size:22px}
  </style>
</head>

<body>
  <header>
  <div class="sitebar">
    <a class="brand" href="/">Lustra</a>
    <nav>
      <a href="/about">About</a>
      <a href="/pricing">Pricing</a>
      <a href="/login">Sign in</a>
    </nav>
  </div>
</header>

  <!-- PAGE BOX with hero + plans (aligned with Lustra) -->
  <div class="wrap">
    <div class="pagebox">
      <h1>Plans</h1>
      <p class="lead">Pricing that fits your desk. Pick a predictable monthly plan, add non-expiring packs when you spike, and keep shipping on-brand CVs without surprises.</p>

      <div class="grid4">
        <!-- STARTER -->
        <div class="card start">
  <div class="inner">
    <div class="name">STARTER</div>
    <div class="qty">100 polished CVs <span class="per">/ mo</span></div>
    <span class="chip"><span class="price-month">£150/mo</span><span class="dot">·</span><span class="price-cv">£1.50 per CV</span></span>

    <ul class="feat" style="margin-top:12px">
      <li><span class="tick">✓</span><span>1 brand template included</span></li>
      <li><span class="tick">✓</span><span>Up to 10 users</span></li>
      <li><span class="tick">✓</span><span>Director dashboard</span></li>
      <li><span class="tick">✓</span><span>CSV export (usage, credits, history)</span></li>
      <li><span class="tick">✓</span><span>Supported files: PDF / DOCX / TXT</span></li>
      <li><span class="tick">✓</span><span>Email support</span></li>
      <li><span class="tick">✓</span><span>Overage: <strong>£1.60/CV</strong></span></li>
    </ul>

    <a class="btn primary" href="/start">Choose Starter</a>
  </div>
</div>

        <!-- GROWTH -->
        <div class="card growth">
  <div class="inner">
    <div class="name">GROWTH</div>
    <div class="qty">250 polished CVs <span class="per">/ mo</span></div>
    <span class="chip"><span class="price-month">£350/mo</span><span class="dot">·</span><span class="price-cv">£1.40 per CV</span></span>

    <ul class="feat" style="margin-top:12px">
      <li><span class="tick">✓</span><span>2 brand template included</span></li>
      <li><span class="tick">✓</span><span>Up to 20 users</span></li>
      <li><span class="tick">✓</span><span>Director dashboard</span></li>
      <li><span class="tick">✓</span><span>CSV export (usage, credits, history)</span></li>
      <li><span class="tick">✓</span><span>Supported files: PDF / DOCX / TXT</span></li>
      <li><span class="tick">✓</span><span>Priority support</span></li>
      <li><span class="tick">✓</span><span>Overage: <strong>£1.50/CV</strong></span></li>
    </ul>

    <a class="btn primary" href="/start">Choose Growth</a>
  </div>
</div>

        <!-- SCALE -->
        <div class="card scale">
  <div class="inner">
    <div class="name">SCALE</div>
    <div class="qty">500 polished CVs <span class="per">/ mo</span></div>
    <span class="chip"><span class="price-month">£650/mo</span><span class="dot">·</span><span class="price-cv">£1.30 per CV</span></span>

    <ul class="feat" style="margin-top:12px">
      <li><span class="tick">✓</span><span>3 brand template included</span></li>
      <li><span class="tick">✓</span><span>Up to 30 users</span></li>
      <li><span class="tick">✓</span><span>Director dashboard</span></li>
      <li><span class="tick">✓</span><span>CSV export (usage, credits, history)</span></li>
      <li><span class="tick">✓</span><span>Supported files: PDF / DOCX / TXT</span></li>
      <li><span class="tick">✓</span><span>Priority support</span></li>
      <li><span class="tick">✓</span><span>Overage: <strong>£1.40/CV</strong></span></li>
    </ul>

    <a class="btn primary" href="/start">Choose Scale</a>
  </div>
</div>

        <!-- BUY PACKS -->
        <div class="card packs">
          <div class="inner">
            <div class="name">BUY PACKS</div>
            <div class="qty">Non-expiring</div>
            <span class="chip"><span class="price-month">Org-wide</span><span class="dot">·</span><span class="price-cv">Used after monthly</span></span>

            <div class="packs-info">
              <div class="title">Choose your pack</div>
              <ul>
                <li><strong>100 CVs</strong> — £160 ( £1.60/CV )</li>
                <li><strong>300 CVs</strong> — £450 ( £1.50/CV )</li>
                <li><strong>500 CVs</strong> — £700 ( £1.40/CV )</li>
              </ul>
            </div>

            <div class="select">
              <select id="packSelect">
                <option value="100">100 CVs — £160</option>
                <option value="300">300 CVs — £450</option>
                <option value="500">500 CVs — £700</option>
              </select>
            </div>

            <ul class="feat" style="margin-top:12px">
              <li><span class="tick">✓</span><span>Never expires; consumed after monthly pool</span></li>
              <li><span class="tick">✓</span><span>Org-wide across users</span></li>
              <li><span class="tick">✓</span><span><strong>Extra templates £50 each</strong> (1 included per org)</span></li>
            </ul>

            <a class="btn primary" href="/start">Buy pack</a>
          </div>
        </div>
      </div> <!-- /.grid4 -->
    </div> <!-- /.pagebox -->
  </div> <!-- /.wrap -->

  <div class="wrap">
    <p class="note">Prices exclude VAT where applicable. Monthly credits reset each month. Packs never expire and are org-wide. We always use Monthly first, then Packs if available, otherwise Overage.</p>
  </div>

  <!-- CALCULATOR (unchanged; only constants updated) -->
  <div class="wrap section">
    <div class="card calc">
      <div class="inner">
        <div class="name">Savings & best plan</div>
        <div class="sub">Estimate time/payroll savings and see which plan fits your volume (with overage).</div>
        <div class="calc-grid" style="margin-top:10px">
          <div><label>CVs per month</label><input id="cvs" type="number" min="0" value="100"></div>
          <div><label>Minutes per CV (manual polish)</label><input id="minManual" type="number" min="0" value="15"></div>
          <div><label>Recruiter hourly cost</label><input id="hourRate" type="number" min="0" value="30"></div>
        </div>
        <div class="calc-out">
          <div><span class="n" id="outHours">25.0</span> hours saved / month</div>
          <div><span class="n">£<span id="outMoney">750</span></span> payroll saved / month</div>
        </div>
        <div class="sub" id="planPick" style="margin-top:8px"></div>
      </div>
    </div>
  </div>

  <script>
    /* Calculator unchanged — only plan constants updated */
    function fmt(n){ return new Intl.NumberFormat('en-GB',{maximumFractionDigits:0}).format(n); }
    function fmtGBP(n){ return '£' + new Intl.NumberFormat('en-GB',{maximumFractionDigits:0}).format(Math.round(n)); }

    const PLANS = [
      { kind:'Monthly', key:'Starter', baseCredits:100, baseCost:150, baseRate:1.50, overRate:1.60 },
      { kind:'Monthly', key:'Growth',  baseCredits:300, baseCost:420, baseRate:1.40, overRate:1.50 },
      { kind:'Monthly', key:'Scale',   baseCredits:500, baseCost:650, baseRate:1.30, overRate:1.40 }
      // Buy Packs are on-demand; not part of the "best monthly plan" picker.
    ];

    function costFor(plan, volume){
      const included=Math.min(volume, plan.baseCredits);
      const over=Math.max(0, volume - plan.baseCredits);
      const cost=plan.baseCost + over*plan.overRate;
      const percv=volume ? (cost/volume) : 0;
      const detail=over>0?`includes ${plan.baseCredits} + ${over} over @ £${plan.overRate.toFixed(2)}`:`up to ${plan.baseCredits} included`;
      return { name:plan.key, cost, percv, detail };
    }

    function calc(){
      const cvs=parseFloat(document.getElementById('cvs').value)||0;
      const mManual=parseFloat(document.getElementById('minManual').value)||0;
      const rate=parseFloat(document.getElementById('hourRate').value)||0;

      const timeSavedHours=(Math.max(0,mManual)*cvs)/60;
      const moneySaved=timeSavedHours*rate;
      document.getElementById('outHours').textContent=(Math.round(timeSavedHours*10)/10).toFixed(1);
      document.getElementById('outMoney').textContent=fmt(Math.round(moneySaved));

      const options=PLANS.map(p=>({plan:p,quote:costFor(p,cvs)})).sort((a,b)=>a.quote.cost-b.quote.cost);
      const pickEl=document.getElementById('planPick');
      if(!cvs){ pickEl.textContent=''; return; }
      const best=options[0].quote;
      const percv=best.percv?` (~£${(Math.round(best.percv*100)/100).toFixed(2)}/CV)`:'';
      pickEl.innerHTML=`Best option: <strong>${best.name}</strong> — <strong>${fmtGBP(best.cost)}</strong>/mo${percv}<br><span class="sub">${best.detail}</span>`;
    }
    document.addEventListener('input', (e)=>{ if(['cvs','minManual','hourRate'].includes(e.target.id)) calc(); });
    document.addEventListener('DOMContentLoaded', calc);
  </script>
</body>
</html>
"""

CONTACT_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Talk to Sales — Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{
      --brand:#2563eb; --brand-2:#22d3ee;
      --ink:#0f172a; --muted:#64748b; --line:#e5e7eb;
      --bg:#f6f9ff; --card:#fff; --shadow:0 10px 24px rgba(2,6,23,.06);
      --ok:#16a34a; --red:#ef4444;
    }
    *{box-sizing:border-box}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;margin:0;background:var(--bg);color:var(--ink)}
    .wrap{max-width:1180px;margin:28px auto 64px;padding:0 28px}
    .nav{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}
    .brand{font-weight:900;color:#0f172a;text-decoration:none;font-size:22px;letter-spacing:.2px}
    .nav a{color:var(--ink);text-decoration:none;font-weight:800;margin-left:22px}
    .pagebox{background:var(--card);border:1px solid var(--line);border-radius:20px;box-shadow:var(--shadow);padding:28px}
    h1{margin:2px 0 6px;font-size:40px;letter-spacing:-.01em}
    .sub{margin:0 0 16px;color:var(--muted)}
    .grid{display:grid;grid-template-columns:2fr 1fr;gap:26px}
    @media(max-width:900px){ .grid{grid-template-columns:1fr} }
    label{font-weight:800;font-size:14px;margin:12px 0 6px;display:block}
    input[type="text"], input[type="email"], select, textarea{
      width:100%;padding:12px 14px;border:1px solid var(--line);border-radius:14px;
      font-size:15px;outline:none;background:#fff;
    }
    textarea{min-height:110px;resize:vertical}
    .row{display:grid;grid-template-columns:1fr 1fr;gap:14px}
    .hint{font-size:12.5px;color:var(--muted);margin-top:4px}
    .btn{
      display:inline-block;width:100%;padding:14px 18px;border-radius:999px;text-align:center;
      font-weight:900;text-decoration:none;border:none;background:linear-gradient(90deg,var(--brand),var(--brand-2));color:#fff
    }
    .side{background:#fbfdff;border:1px solid var(--line);border-radius:16px;padding:16px}
    .pill{display:inline-flex;align-items:center;padding:4px 8px;border:1px solid var(--line);border-radius:999px;margin:6px 6px 0 0;font-size:12px;color:var(--ink)}
    .tick{display:inline-flex;align-items:center;justify-content:center;width:16px;height:16px;border-radius:50%;background:rgba(34,211,238,.15);color:#0891b2;font-weight:900;font-size:11px;margin-right:6px}
    .err{color:var(--red);font-weight:700;margin:6px 0 0}
    .ok{color:var(--ok);font-weight:800}
    .fineprint{margin-top:12px;color:var(--muted);font-size:12.5px}
/* --- FORM FIELD SIZE TUNE-UP (reduce field text; keep label size unchanged) --- */
.pagebox input[type="text"],
.pagebox input[type="email"],
.pagebox input[type="number"],
.pagebox select{
  height:42px;          /* a touch shorter */
  padding:8px 12px;
  font-size:14px;       /* smaller than label; titles unchanged */
  line-height:1.25;
  border-radius:12px;
}

.pagebox textarea{
  font-size:14px;       /* smaller than label */
  line-height:1.35;
  padding:10px 12px;
  min-height:110px;
}

.pagebox ::placeholder{font-size:14px;opacity:.65}
/* optional: keep or remove this if you liked the smaller button */
/* .pagebox .btn{padding:12px 16px} */
  </style>
</head>
<body>
  <div class="wrap">
    <div class="nav">
      <a class="brand" href="/">Lustra</a>
      <div>
        <a href="/about">About</a>
        <a href="/pricing" style="margin-left:18px">Pricing</a>
        <a href="/login" style="margin-left:18px">Sign in</a>
      </div>
    </div>

    <div class="pagebox">
      <h1>Talk to Sales</h1>
      <p class="sub">Tell us about your desk; we’ll get you live in a day. Or start a 5-CV trial — no card required.</p>

      <form method="POST" enctype="multipart/form-data" class="grid">
        <div>
          <label>Company</label>
          <input name="company" type="text" required>

          <div class="row">
            <div>
              <label>Work email</label>
              <input name="email" type="email" required>
            </div>
            <div>
              <label>Your name</label>
              <input name="name" type="text">
            </div>
          </div>

          <div class="row">
            <div>
              <label>Monthly CVs</label>
              <select name="volume" required>
                <option value="">Choose…</option>
                <option>100</option>
                <option>250</option>
                <option>500</option>
                <option>700+</option>
                <option>Not sure</option>
              </select>
            </div>
            <div>
              <label>Users</label>
              <select name="users" required>
                <option value="">Choose…</option>
                <option>1–5</option>
                <option>6–10</option>
                <option>11–20</option>
                <option>21–30</option>
                <option>30+</option>
              </select>
            </div>
          </div>

          <div class="row">
            <div>
              <label>Templates needed</label>
              <select name="templates">
                <option>1 (included)</option>
                <option>2</option>
                <option>3+</option>
              </select>
              <div class="hint">Extra templates £50 each.</div>
            </div>
            <div>
              <label>SSO / SLA required?</label>
              <select name="need_sso">
                <option>No</option>
                <option>Yes</option>
                <option>Maybe</option>
              </select>
            </div>
          </div>

          <label>Message (optional)</label>
          <textarea name="message" placeholder="Anything we should know?"></textarea>

          <label>Upload a sample CV (optional)</label>
          <input name="file" type="file" accept=".pdf,.doc,.docx,.txt">

          <!-- anti-spam -->
          <input type="text" name="website" style="display:none" tabindex="-1" autocomplete="off">

          <!-- CSRF -->
          <input type="hidden" name="csrf" value="{{csrf}}">

          <div style="margin-top:16px">
            <button class="btn" type="submit">Send to Sales / Start trial</button>
          </div>

          <div class="fineprint">We’ll reply within 1 business day. We never use your documents to train models.</div>
        </div>

        <aside class="side">
          <div style="font-weight:900;margin-bottom:10px">What you get</div>
          <div><span class="tick">✓</span>On-brand DOCX output</div>
          <div><span class="tick">✓</span>PDF / DOCX / TXT supported</div>
          <div><span class="tick">✓</span>Org-wide credits & usage</div>
          <div><span class="tick">✓</span>CSV exports and Director dashboard</div>
          <hr style="border:none;border-top:1px solid var(--line);margin:14px 0">
          <div style="font-weight:900;margin-bottom:6px">Plans</div>
          <div class="pill">Starter — 100 CVs · £150/mo</div>
          <div class="pill">Growth — 250 CVs · £350/mo</div>
          <div class="pill">Scale — 500 CVs · £650/mo</div>
          <div class="hint" style="margin-top:8px">Buy Packs available · non-expiring.</div>
          <hr style="border:none;border-top:1px solid var(--line);margin:14px 0">
          <div style="font-weight:900;margin-bottom:6px">Prefer email?</div>
          <div><a href="mailto:hello@lustra.uk">hello@lustra.uk</a></div>
          <div class="hint">We reply within 1 business day.</div>
        </aside>
      </form>
    </div>
  </div>
</body>
</html>
"""

# ------------------------ Start Free Trial (new) ------------------------
START_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Start free trial — CV Polisher</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{--blue:#003366;--ink:#111827;--muted:#6b7280;--line:#e5e7eb;--bg:#f2f6fb;--card:#fff}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;margin:0;background:var(--bg);color:var(--ink)}
    .wrap{max-width:620px;margin:36px auto;padding:0 18px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:18px}
    h1{margin:0 0 12px;font-size:26px;color:var(--blue)}
    label{font-weight:600;font-size:13px}
    input,select,textarea{width:100%;padding:10px;border:1px solid var(--line);border-radius:10px;margin-top:6px}
    button{width:100%;margin-top:12px;background:linear-gradient(90deg,#003366,#0a4d8c);color:#fff;border:none;border-radius:10px;padding:10px 16px;font-weight:700;cursor:pointer}
    .muted{color:var(--muted);font-size:12px;margin-top:8px}
    .row{display:grid;grid-template-columns:1fr 1fr;gap:12px}
    @media(max-width:640px){ .row{grid-template-columns:1fr} }
  </style>
</head>
<body>
  <div class="wrap">
    <a href="/">← Home</a>
    <div class="card">
      <h1>Start your free trial</h1>
      <p class="muted">5 free CVs, no card required. On submit you’ll be taken to Sign in. Your banner in the app will show trial credits left.</p>
      <form method="post" action="/start" autocomplete="off">
        <label>Company</label>
        <input name="company" required />
        <div class="row">
          <div>
            <label>Work email</label>
            <input name="email" type="email" required />
          </div>
          <div>
            <label>Your name</label>
            <input name="name" required />
          </div>
        </div>
        <div class="row">
          <div>
            <label>Team size</label>
            <select name="team_size">
              <option value="">Choose…</option>
              <option>Just me</option>
              <option>2–5</option>
              <option>6–15</option>
              <option>16–50</option>
              <option>50+</option>
            </select>
          </div>
          <div style="display:none">
            <!-- Honeypot (bots will fill this) -->
            <label>Website</label>
            <input name="website" />
          </div>
        </div>
        <label>Notes (optional)</label>
        <textarea name="notes" rows="4" placeholder="Anything we should know?"></textarea>
        <label style="display:flex;gap:8px;align-items:center;margin-top:10px">
          <input type="checkbox" name="agree" required style="width:auto"/> I agree to fair use of the free trial.
        </label>
        <button type="submit">Create free trial</button>
        <p class="muted" style="text-align:center">Already have an account? <a href="/login">Sign in</a></p>
      </form>
    </div>
  </div>
</body>
</html>
"""

# ------------------------ Branded App UI (unchanged except banner hook + Director button) ------------------------
HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Hamilton Recruitment — Executive Search & Selection</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{
  --bg:#f5f8fd;
  --ink:#0f172a;
  --muted:#64748b;
  --line:#e5e7eb;
  --card:#ffffff;
  --blue:#2563eb;
  --blue-2:#22d3ee;;
  --ok:#16a34a;
  --shadow:0 10px 28px rgba(13,59,102,.08);
}
*{box-sizing:border-box}
body{
  margin:0;
  color:var(--ink);
  background:var(--bg);
  font-family:Inter, system-ui, -apple-system, "Segoe UI", Roboto, Arial, sans-serif;
  letter-spacing:.2px;
}

/* shell */
.wrap{max-width:1160px;margin:28px auto;padding:0 20px}
.nav{display:flex;align-items:center;gap:12px;margin-bottom:18px}
.brand-logo{width:46px;height:46px;border-radius:10px;background:linear-gradient(135deg,var(--blue),var(--blue-2));display:flex;align-items:center;justify-content:center;overflow:hidden}
.brand-logo img{width:100%;height:100%;object-fit:contain}
.brand-head{line-height:1.05}
.brand-title{margin:0;font-size:24px;font-weight:900;color:var(--blue);letter-spacing:-.01em}
.brand-sub{margin:0;color:var(--muted);font-size:12.5px}

.nav-right{margin-left:auto;display:flex;gap:10px}
.nav-right a{
  display:inline-block;background:#fff;color:var(--blue);border:1px solid var(--line);
  border-radius:12px;padding:9px 14px;font-weight:800;text-decoration:none;box-shadow:var(--shadow)
}
.nav-right a:hover{transform:translateY(-1px)}

/* layout */
.grid{display:grid;grid-template-columns:1.2fr .8fr;gap:16px}
@media(max-width:980px){ .grid{grid-template-columns:1fr;gap:12px} }

/* cards */
.card{
  background:var(--card);
  border:1px solid var(--line);
  border-radius:22px;
  padding:18px 20px;
  box-shadow:var(--shadow)
}
.card h3{margin:0 0 12px;color:var(--blue);font-size:18px;letter-spacing:.2px}

/* form + buttons */
label{font-weight:700;font-size:13.5px}
input[type=file]{width:100%;padding:10px;border:1px solid var(--line);border-radius:12px;margin-top:6px;background:#fff}
.ts{color:var(--muted);font-size:12.5px}
button{
  background:linear-gradient(90deg,var(--blue),var(--blue-2));color:#fff;border:none;border-radius:12px;
  padding:12px 18px;font-weight:900;cursor:pointer;box-shadow:var(--shadow)
}
button:hover{transform:translateY(-1px)}
button[disabled]{opacity:.6;cursor:not-allowed}

/* progress */
.progress{display:none;margin-top:12px;border:1px solid var(--line);border-radius:14px;padding:12px;background:var(--card)}
.stage{display:flex;align-items:center;gap:8px;margin-bottom:6px;font-size:12.5px;color:var(--muted)}
.stage .dot{width:9px;height:9px;border-radius:999px;background:#cbd5e1}
.stage.active{color:var(--ink)}
.stage.active .dot{background:var(--blue)}
.stage.done{color:var(--ok)}
.stage.done .dot{background:var(--ok)}
.bar{height:12px;border-radius:999px;background:var(--line);overflow:hidden;margin-top:6px;position:relative}
.bar>span{display:block;height:100%;width:0;background:linear-gradient(90deg,var(--blue),var(--blue-2));transition:width .35s ease}
.pct{position:absolute;right:8px;top:50%;transform:translateY(-50%);font-size:11px;color:#fff;font-weight:800}
.success{display:none;margin-top:10px;color:var(--ok);font-weight:800}

/* stats */
.statsgrid {
  display: grid;
  grid-template-columns: repeat(4, 1fr);
  gap: 8px;                  /* slightly smaller gap */
  margin-bottom: 10px;
}

.stat {
  border: 1px solid var(--line);
  border-radius: 8px;         /* less rounded */
  padding: 8px 10px;          /* slimmer boxes */
  background: var(--card);
}

.stat .k {
  font-size: 12px;
  color: var(--muted);
  font-weight: 700;
}

.stat .v {
  font-size: 12px;            /* reduce size of numbers */
  font-weight: 600;           /* softer weight */
  margin-top: 2px;            /* tighter spacing */
  color: #333;                /* softer than brand ink */
}
.kicker{color:var(--muted);font-size:12.5px;margin:8px 0 6px}
.history{border:1px solid var(--line);border-radius:14px;max-height:300px;overflow:auto;background:var(--card)}
.row{display:flex;justify-content:space-between;gap:10px;padding:8px 12px;border-bottom:1px solid var(--line)}
.row:last-child{border-bottom:none}

/* Full history type sizes (scoped) */
#history { font-size: 11.5px; }           /* base for the list */
#history .row strong { font-size: 12px; } /* candidate name */
#history .muted { font-size: 11px; }      /* timestamp + filename */

.candidate{font-weight:700;font-size:13.5px}
.tsm{color:var(--muted);font-size:12px}

/* credits chip in stats */
.chip{
  display:inline-block;border:1px dashed var(--line);border-radius:12px;padding:6px 10px;font-weight:800;
  color:var(--blue);background:#fff
}
/* Skills list overrides (only inside the Skills card) */
#skillsCard .chip{
  font-size: 11px;     /* make smaller */
  font-weight: 600;    /* less bold than default 800 */
  color: #000;         /* black text */
}
/* Skills pills – smaller + black (only the unified list) */
#skillsAll .pill{ font-size:10px; color:#000; }
#skillsAll .pill .x{ font-size:10px; }

/* if any pill label ends up inside an <a>, make it black too */
#skillsAll .pill a{ color:#000; }
.pill{
  display:inline-flex;align-items:center;gap:6px;
  padding:3px 8px;border:1px solid var(--line);border-radius:999px;
  margin:3px 6px 0 0;font-weight:600;font-size:10.5px;background:#fff;line-height:1.1;color:#0f172a
}
.pill .x{
  cursor:pointer;border:none;background:transparent;font-weight:900;
  font-size:12px;padding:0 2px;line-height:1;color:#0f172a
}
/* (Optional cleanup) You can delete .pill.base and .pill.off if present */
  </style>
  <script>
    let timer=null, pct=0;
    function setProgress(p){
      pct = Math.max(0, Math.min(100, p));
      const bar = document.getElementById('barfill');
      const pctEl = document.getElementById('pct');
      if(bar){ bar.style.width = pct + '%'; }
      if(pctEl){ pctEl.textContent = Math.round(pct) + '%'; }
    }
    function setStage(i, total){
      const items = document.querySelectorAll('.stage');
      items.forEach((el,idx)=>{
        el.classList.remove('active','done');
        if(idx < i) el.classList.add('done');
        if(idx === i) el.classList.add('active');
      });
      setProgress((i/(total-1))*95);
    }
    function startProgress(){
      const prog = document.getElementById('progress');
      const ok = document.getElementById('success');
      const btn = document.getElementById('btn');
      if(prog) prog.style.display='block';
      if(ok) ok.style.display='none';
      if(btn) btn.disabled=true;
      const total = 5; let i = 0;
      setStage(0,total);
      if(timer){ clearInterval(timer); }
      timer = setInterval(()=>{
        if(i < total-1){ i++; setStage(i,total); } else { clearInterval(timer); }
      }, 700);
    }
    function stopProgressSuccess(){
      const prog = document.getElementById('progress');
      const ok = document.getElementById('success');
      const btn = document.getElementById('btn');
      if(timer){ clearInterval(timer); timer=null; }
      document.querySelectorAll('.stage').forEach(el=>el.classList.add('done'));
      setProgress(100);
      if(ok) ok.style.display='block';
      if(btn) btn.disabled=false;
      setTimeout(()=>{ if(prog) { prog.style.display='none'; setProgress(0);} }, 800);
      const form = document.getElementById('upload-form');
      if(form){ form.reset(); }
      const nameEl = document.getElementById('filenamePreview');
      if(nameEl){ nameEl.textContent = "—"; }
    }
    async function refreshStats(){
      // --- Fast path: single call to /me/dashboard; fallback to legacy below ---
  try {
    const d = await fetch('/me/dashboard').then(r => r.ok ? r.json() : Promise.reject());
    const setText = (sel, val) => { const el = document.querySelector(sel); if (el) el.textContent = (val ?? '').toString(); };

    setText('#downloadsMonth', d.downloadsMonth);
    setText('#lastCandidate', d.lastCandidate);

    if (d.lastTime) {
      const dt = new Date(d.lastTime);
      setText('#lastTime', isNaN(dt.getTime()) ? d.lastTime : dt.toLocaleString());
    } else {
      setText('#lastTime', '');
    }

    // Credits tile: show remaining balance (prefer this), fallback to used
if (typeof d.creditsBalance === 'number' && !Number.isNaN(d.creditsBalance)) {
  setText('#creditsUsed', d.creditsBalance);
} else if (typeof d.creditsUsed === 'number' && !Number.isNaN(d.creditsUsed)) {
  setText('#creditsUsed', d.creditsUsed);
} else {
  setText('#creditsUsed', '');
}

    return; // success -> stop here, skip legacy logic below
  } catch (e) {
    // Ignore and let the existing legacy fetches run
  }

  try{
    const r = await fetch('/stats', {cache:'no-store'});
    if(!r.ok) return;
    const s = await r.json();


    // Trial banner
    const tb = document.getElementById('trialBanner');
    if (tb) {
      const left = s.trial_credits_left || 0;
      if (left > 0) {
        tb.style.display = 'block';
        tb.querySelector('.left').textContent = left;
      } else {
        tb.style.display = 'none';
      }
    }

    // Top stats
    const dm = document.getElementById('downloadsMonth');
    if (dm) dm.textContent = (s.downloads_this_month ?? s.downloads);
    const lc = document.getElementById('lastCandidate');
    if (lc) lc.textContent = s.last_candidate || '—';
    const lt = document.getElementById('lastTime');
    if (lt) lt.textContent = s.last_time || '—';

    // Credits used (in-plan) — show "X / Y"
const cu = document.getElementById('creditsUsed');
if (cu) {
  const planCap = (s.plan && s.plan.credits) ? s.plan.credits : 0;
  const used = s.credits_used ?? 0;
  cu.textContent = `${used} / ${planCap}`;
}

    // History list
    const list = document.getElementById('history');
    if (list) {
      list.innerHTML = '';
      (s.history || []).slice().reverse().forEach(item => {
        const row = document.createElement('div'); row.className = 'row';
        const left = document.createElement('div');
        left.innerHTML = '<div class="candidate">' + (item.candidate || item.filename || '—') + '</div><div class="ts">' + (item.filename || '') + '</div>';
        const right = document.createElement('div'); right.className = 'ts';
        right.textContent = item.ts || '';
        row.appendChild(left); row.appendChild(right); list.appendChild(row);
      });
    }
    } catch(e) {}

  // === NEW: override with per-user DB usage ===
  try {
    const mu = await fetch('/me/usage', {cache:'no-store'});
    if (mu.ok) {
      const j = await mu.json();
          if (j && j.ok) {
      const dmEl = document.getElementById('downloadsMonth');
      if (dmEl) dmEl.textContent = j.month_usage ?? 0;
    }
    }
  } catch(e) {}

  // Optionally also refresh last candidate/time from DB
  try {
    const le = await fetch('/me/last-event', {cache:'no-store'});
    if (le.ok) {
      const j = await le.json();
      if (j && j.ok) {
        const lcEl = document.getElementById('lastCandidate');
const ltEl = document.getElementById('lastTime');
if (lcEl) lcEl.textContent = j.candidate || (s.last_candidate || '—');
if (ltEl) ltEl.textContent = j.ts || (s.last_time || '—');
      }
    }
  } catch(e) {}

// === Unified Skills rendering (single list) ===
let skillsState = null;

async function loadSkills(){
  const r = await fetch('/skills', {cache:'no-store'});
  if(!r.ok) return;
  skillsState = await r.json();
  renderSkillsUnified();

  // Hide old sections; show unified one
  const custom = document.getElementById('customSkills');
  const base = document.getElementById('baseSkills');
  const allH = document.getElementById('skillsAllHeader');
  const allC = document.getElementById('skillsAll');

  if (custom){
    if (custom.previousElementSibling) custom.previousElementSibling.style.display = 'none'; // "Custom skills (A–Z)"
    custom.style.display = 'none';
  }
  if (base){
    if (base.previousElementSibling) base.previousElementSibling.style.display = 'none';     // "Built-in skills (A–Z)"
    base.style.display = 'none';
  }
  if (allH) allH.style.display = 'block';
  if (allC) allC.style.display = 'block';
}

function makePill(label, actionLabel, onClick, extraClass){
  const span = document.createElement('span');
  span.className = 'pill' + (extraClass?(' '+extraClass):'');
  span.append(document.createTextNode(label+' '));
  const b = document.createElement('button');
  b.type='button'; b.className='x'; b.textContent = actionLabel;
  b.addEventListener('click', onClick);
  span.appendChild(b);
  return span;
}

function renderSkillsUnified(){
  const container = document.getElementById('skillsAll');
  if(!container || !skillsState) return;
  container.innerHTML = '';

  const list = (skillsState.effective || [])
    .slice()
    .sort((a,b)=>a.localeCompare(b, undefined, {sensitivity:'base'}));

  list.forEach(label=>{
    container.appendChild(
      makePill(label, '×', ()=> removeSkill(label), '')
    );
  });
}

// Add a skill -> always added as "custom"; unified list shows it with the rest
async function addSkill(label){
  if(!label) return;
  const fd = new FormData(); fd.append('skill', label);
  const r = await fetch('/skills/custom/add', {method:'POST', body: fd});
  if(r.ok){ await loadSkills(); }
}

// Remove from unified list:
// - if it’s custom: delete it
// - if it’s a built-in: disable it so it disappears from "effective"
async function removeSkill(label){
  if(!skillsState) return;
  const isCustom = new Set((skillsState.custom || []).map(s=>s.toLowerCase()));
  if(isCustom.has(label.toLowerCase())){
    const fd = new FormData(); fd.append('skill', label);
    const r = await fetch('/skills/custom/remove', {method:'POST', body: fd});
    if(r.ok){ await loadSkills(); }
  }else{
    await toggleBase(label, 'disable');
  }
}

/* Keep compatibility with existing calls (delegates) */
async function addCustom(skill){ return addSkill(skill); }
async function removeCustom(skill){ return removeSkill(skill); }
async function toggleBase(skill, action){
  const fd = new FormData(); fd.append('skill', skill); fd.append('action', action);
  const r = await fetch('/skills/base/toggle', {method:'POST', body: fd});
  if(r.ok){ skillsState = await r.json(); renderSkillsUnified(); }
}
    document.addEventListener('DOMContentLoaded',()=>{
      refreshStats();
      setInterval(refreshStats, 5000);

      const form = document.getElementById('upload-form');
      const fileInput = document.getElementById('cv');

      // fetch + Blob download
form.addEventListener('submit', async (e)=>{
  e.preventDefault();
  startProgress();              // show progress UI + begin staged animation
  try{
    const fd = new FormData(form);
    const r = await fetch('/polish', { method:'POST', body: fd, cache:'no-store' });
    if(!r.ok) throw new Error('Server error ('+r.status+')');

    // Download blob
    const blob = await r.blob();
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;
    a.download = 'polished_cv.docx';
    document.body.appendChild(a);
    a.click();
    a.remove();
    URL.revokeObjectURL(url);

    stopProgressSuccess();      // mark all stages done, briefly show “Done”
    refreshStats();             // refresh stats after a successful polish
  }catch(err){
    alert('Polishing failed: ' + (err?.message||'Unknown error'));
    // clean up UI on error
    const btn = document.getElementById('btn'); if(btn) btn.disabled=false;
    const prog = document.getElementById('progress'); if(prog) prog.style.display='none';
    setProgress(0);
  }
});

      fileInput.addEventListener('change',()=>{
        const v=fileInput.files?.[0]?.name||'';
        const name=v.replace(/[_-]/g,' ').replace(/\.(pdf|docx|txt)$/i,'');
        if(name){document.getElementById('filenamePreview').textContent=name;}
      });
     

// Skills manager toggle + events
const skillsToggle = document.getElementById('skillsToggle');
const skillsCard = document.getElementById('skillsCard');
if (skillsToggle && skillsCard){
  skillsToggle.addEventListener('click', ()=>{
    const show = skillsCard.style.display === 'none' || skillsCard.style.display === '';
    skillsCard.style.display = show ? 'block' : 'none';
    skillsToggle.textContent = show ? 'Hide' : 'Show';
    if (show) loadSkills();
  });
}
const skillForm = document.getElementById('skillAddForm');
if (skillForm){
  skillForm.addEventListener('submit', (e)=>{
    e.preventDefault();
    const inp = document.getElementById('skillInput');
    const val = (inp.value||'').trim();
    if (val){ addSkill(val); inp.value=''; }
  });
}

});
</script>
    
</head>
<body>
  <div class="wrap">
    <div class="nav">
      <div class="brand-logo"><img src="/logo" alt="Hamilton Logo" onerror="this.style.display='none'"/></div>
      <div class="brand-head">
        <p class="brand-title">Hamilton Recruitment</p>
        <p class="brand-sub">Executive Search &amp; Selection</p>
      </div>
      <div style="margin-left:auto; display:flex; gap:8px;">
        <!-- NEW: Director button -->
        <a href="/director" style="display:inline-block;background:#fff;color:var(--blue);border:1px solid var(--line);border-radius:10px;padding:8px 12px;font-weight:700;text-decoration:none">Director</a>
        <a href="/logout" style="display:inline-block;background:#fff;color:var(--blue);border:1px solid var(--line);border-radius:10px;padding:8px 12px;font-weight:700;text-decoration:none">Log out</a>
      </div>
    </div>

    <!-- NEW: free-trial banner -->
    <div id="trialBanner" class="card" style="display:none; margin-bottom:12px">
      <strong>Free trial:</strong> <span class="left">5</span> CVs left.
      <span class="ts">Need more? <a href="/pricing">See plans</a></span>
    </div>

    <div class="grid">
      <div class="card">
        <h3>Upload CV</h3>
        <form id="upload-form" method="post" action="/polish" enctype="multipart/form-data">
          <label for="cv">Raw CV (PDF / DOCX / TXT)</label><br/>
          <input id="cv" type="file" name="cv" accept=".pdf,.docx,.txt" required />
          <div class="ts" style="margin-top:4px">Header (logo &amp; bar) preserved from Hamilton template.</div>
          <div class="ts" style="margin-top:2px">Candidate: <span id="filenamePreview">—</span></div>

          <div id="progress" class="progress">
            <div class="stage"><div class="dot"></div><div>Uploading</div></div>
            <div class="stage"><div class="dot"></div><div>Extracting text</div></div>
            <div class="stage"><div class="dot"></div><div>Structuring</div></div>
            <div class="stage"><div class="dot"></div><div>Composing document</div></div>
            <div class="stage"><div class="dot"></div><div>Downloading</div></div>
            <div class="bar"><span id="barfill"></span><span id="pct" class="pct">0%</span></div>
          </div>

          <div id="success" class="success">Done. Your download should start automatically.</div>

          <div style="margin-top:12px"><button id="btn" type="submit">Polish & Download</button></div>
        </form>
      </div>

      <div class="card">
        <h3>Session Stats</h3>

<!-- 4 compact tiles -->
<div class="statsgrid">
  <div class="stat"><div class="k">Downloads this month</div><div class="v" id="downloadsMonth">0</div></div>
  <div class="stat"><div class="k">Last Candidate</div><div class="v" id="lastCandidate">—</div></div>
  <div class="stat"><div class="k">Last Polished</div><div class="v" id="lastTime">—</div></div>
  <div class="stat">
    <div class="k">Credits Left</div>
    <div class="v" id="creditsLeft">—</div>
  </div>
</div>

<!-- Full history: now collapsible (default hidden) -->
<div class="kicker" style="margin:10px 0 6px 2px; display:flex; align-items:center; justify-content:space-between">
  <span>Full history</span>
  <button id="historyToggle" type="button" class="chip">Show</button>
</div>
<div id="history" class="history" style="display:none"></div>

<!-- Skills manager: keep as collapsible (unchanged) -->
<div class="kicker" style="margin:12px 0 6px 2px; display:flex; align-items:center; justify-content:space-between">
  <span>Skills dictionary (matching keywords)</span>
  <button id="skillsToggle" type="button" class="chip">Show</button>
</div>
<div id="skillsCard" class="ts" style="display:none; border:1px dashed var(--line); border-radius:12px; padding:10px; background:#fff">
  <div class="ts" style="margin-bottom:8px">
    These keywords are used to surface <strong>Skills</strong> when polishing. Add your own, remove yours, or disable built-ins.
  </div>

  <form id="skillAddForm" style="display:flex; gap:8px; flex-wrap:wrap; margin-bottom:8px">
    <input id="skillInput" placeholder="Add a skill (e.g., ACCA)" style="flex:1; min-width:220px; padding:10px; border:1px solid var(--line); border-radius:10px"/>
    <button type="submit">Add</button>
  </form>

  <!-- Unified list will render here -->
  <div class="ts" id="skillsAllHeader" style="margin:6px 0 2px; display:none">Skills (A–Z)</div>
  <div id="skillsAll" style="display:none"></div>

  <!-- Hide legacy sections by default (JS also hides them when unified list shows) -->
  <div class="ts" style="margin:6px 0 2px; display:none">Custom skills (A–Z)</div>
  <div id="customSkills" style="display:none"></div>
  <div class="ts" style="margin:10px 0 2px; display:none">Built-in skills (A–Z)</div>
  <div id="baseSkills" style="display:none"></div>
</div>
      </div>
    </div>
  </div>
  <script>
(function(){
  const btn = document.getElementById('historyToggle');
  const panel = document.getElementById('history');
  if (!btn || !panel) return;

  let loaded = false, loading = false;

  async function loadOnce() {
    if (loaded || loading) return;
    loading = true;
    panel.innerHTML = '<div class="muted">Loading…</div>';
    try {
      const r = await fetch('/me/history', { cache: 'no-store' });
      const j = await r.json();
      const rows = Array.isArray(j?.history) ? j.history : Array.isArray(j) ? j : [];
      panel.innerHTML = rows.length
        ? rows.map(it => `
            <div class="row" style="padding:6px 0;border-bottom:1px solid var(--line)">
              <div>
                <div class="candidate">${it.candidate || it.filename || '—'}</div>
                ${it.filename ? `<div class="ts">${it.filename}</div>` : ''}
              </div>
              <div class="ts">${it.ts || ''}</div>
            </div>
          `).join('')
        : '<div class="muted">(no history yet)</div>';
      loaded = true;
    } catch(e) {
      panel.innerHTML = '<div class="muted">Could not load history.</div>';
    } finally {
      loading = false;
    }
  }

  btn.addEventListener('click', async () => {
    const opening = (panel.style.display === 'none' || panel.style.display === '');
    panel.style.display = opening ? 'block' : 'none';
    btn.textContent = opening ? 'Hide' : 'Show';
    if (opening) await loadOnce();
  });
})();
</script>
</body>
</html>
"""

# ------------------------ Login page HTML (unchanged, plus a small "Forgot password?" link) ------------------------
LOGIN_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Lustra — Sign in</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{
  --blue:#2563eb;      /* vivid indigo */
  --blue-2:#22d3ee;    /* bright cyan  */
  --ink:#0f172a; --muted:#5b677a; --line:#e5e7eb;
  --bg:#f5f8fd; --card:#ffffff; --shadow: 0 10px 28px rgba(13,59,102,.08);
}
    *{box-sizing:border-box}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}
    .wrap{max-width:1100px;margin:12px auto 56px;padding:0 20px}
    .nav{display:flex;align-items:center;justify-content:space-between;margin-bottom:8px}
    .brand{font-weight:900;color:var(--blue);text-decoration:none;font-size:22px;letter-spacing:.2px}
    .nav a{color:var(--ink);text-decoration:none;font-weight:800;margin-left:22px}

    .auth{max-width:460px;margin:28px auto 0;background:var(--card);border:1px solid var(--line);border-radius:18px;padding:18px;box-shadow:var(--shadow)}
    h1{margin:0 0 12px;font-size:28px;color:var(--blue)}
    label{font-weight:600;font-size:13px}
    input[type=text],input[type=password]{width:100%;padding:12px;border:1px solid var(--line);border-radius:12px;margin-top:6px}
    button{width:100%;margin-top:14px;background:linear-gradient(90deg,var(--blue),var(--blue-2));color:#fff;border:none;border-radius:12px;padding:12px 16px;font-weight:800;cursor:pointer;box-shadow:var(--shadow)}
    .muted{color:var(--muted);font-size:12px;text-align:center;margin-top:10px}
    .err{margin-top:8px;color:#b91c1c;font-weight:800;font-size:12px}
    a{color:var(--blue);text-decoration:none}
    /* === Compact Sign-in Card (only affects login page) === */
#signinCard{
  max-width: 680px;     /* <- make it narrower (try 640–720 to taste) */
  width: 100%;
  margin: 40px auto;    /* centers the card with comfortable vertical space */
  padding: 22px 26px;   /* <- reduce padding to make the card shorter */
  border-radius: 16px;  /* slightly tighter corners (optional) */
}

#signinCard h1{
  font-size: 36px;      /* was larger; this helps reduce height */
  margin: 6px 0 12px;
}

#signinCard .field,
#signinCard label{
  margin-bottom: 6px;
}

#signinCard input{
  padding: 12px 14px;   /* slightly smaller inputs = less height */
}

#signinCard .btn{
  padding: 14px 16px;   /* slightly smaller button = less height */
  border-radius: 12px;
}

@media (max-width: 640px){
  #signinCard{ 
    max-width: 94vw;    /* keep it tidy on mobile */
    margin: 24px auto;
    padding: 18px 18px;
  }
  #signinCard h1{ font-size: 30px; }
}
/* ===== Consistent site header (brand + right links) across all pages ===== */
:root{
  --site-max: 1200px;
  --site-pad-x: 24px;
  --site-pad-y: 16px;
}
.sitebar, .topbar, header .sitebar, header .topbar{
  max-width: var(--site-max);
  margin: 0 auto;
  padding: var(--site-pad-y) var(--site-pad-x);
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 16px;
}
.sitebar .brand, .topbar .brand{
  display:flex; align-items:center; gap:10px;
  font-weight: 900; letter-spacing: .02em; font-size: 22px;
}
.sitebar nav, .topbar .nav{
  display:flex; align-items:center; gap:24px;
}
.sitebar a, .topbar a{
  text-decoration:none; font-weight:800;
}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="nav">
      <a class="brand" href="/">Lustra</a>
      <div>
  <a href="/">Home</a>
  <a href="/about" style="margin-left:18px">About</a>
  <a href="/login" style="margin-left:18px">Sign in</a>
</div>
    </div>

    <div class="auth">
      <h1>Sign in</h1>
      <!--ERROR-->
      <form method="post" action="/login" autocomplete="off">
        <label for="username">Username</label>
        <input id="username" type="text" name="username" autofocus required />
        <div style="height:10px"></div>
        <label for="password">Password</label>
        <input id="password" type="password" name="password" required />
        <button type="submit">Continue</button>
      </form>
      <div class="muted">
        Default demo: admin / hamilton • <a href="/forgot">Forgot password?</a> • <a href="/">Home</a>
      </div>
    </div>
  </div>
</body>
</html>
"""

# ------------------------ Director: gate + pages ------------------------
DIRECTOR_PASS = os.getenv("DIRECTOR_PASS", "director")
RESET_CODE = os.getenv("RESET_CODE", "reset123")  # used for password resets

DIRECTOR_LOGIN_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Director access — Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{
      --brand:#2563eb; --brand-2:#22d3ee;
      --ink:#0f172a; --muted:#64748b; --line:#e5e7eb;
      --bg:#f6f9ff; --card:#ffffff; --shadow:0 10px 24px rgba(13,59,102,.08);
    }
    *{box-sizing:border-box}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}
    a{color:var(--brand);text-decoration:none}
    .wrap{max-width:560px;margin:56px auto;padding:0 18px}
    .logo{display:flex;align-items:center;gap:10px;margin-bottom:18px}
    .logo .dot{width:10px;height:10px;border-radius:999px;background:linear-gradient(90deg,var(--brand),var(--brand-2))}
    .logo .name{font-weight:900;letter-spacing:.06em;text-transform:uppercase;font-size:12px;color:#0b1220}
    .card{background:var(--card);border:1px solid var(--line);border-radius:20px;padding:22px;box-shadow:var(--shadow)}
    h1{margin:0 0 6px;font-size:22px;letter-spacing:-.01em}
    .sub{color:var(--muted);font-size:14px;margin:0 0 14px}
    label{font-weight:700;font-size:13px}
    input[type=password]{width:100%;padding:12px;border:1px solid var(--line);border-radius:12px;margin-top:6px}
    button{width:100%;margin-top:14px;background:linear-gradient(90deg,var(--brand),var(--brand-2));color:#fff;border:none;border-radius:12px;padding:12px 16px;font-weight:800;cursor:pointer}
    .muted{color:var(--muted);font-size:13px;margin-top:12px}
    .err{background:#fee2e2;border:1px solid #fecaca;color:#991b1b;padding:10px;border-radius:12px;font-weight:700;margin:10px 0}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="logo">
      <div class="dot"></div><div class="name">Lustra • Director</div>
    </div>
    <div class="card">
      <h1>Sign in</h1>
      <p class="sub">Enter the director password to access the console.</p>
      <!--DERR-->
      <form method="post" action="/director/login" autocomplete="off">
        <label for="password">Director password</label>
        <input id="password" type="password" name="password" autofocus required />
        <button type="submit">Continue</button>
      </form>
      <div class="muted">
        <a href="/director/forgot">Forgot director password?</a> · <a href="/app">Back to app</a>
      </div>
    </div>
  </div>
</body>
</html>
"""

DIRECTOR_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Director — Lustra Console</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{
      --brand:#2563eb; --brand-2:#22d3ee;
      --ink:#0f172a; --muted:#64748b; --line:#e5e7eb;
      --bg:#f6f9ff; --card:#ffffff; --shadow:0 10px 24px rgba(13,59,102,.08);
      --ok:#16a34a; --warn:#a16207; --bad:#b91c1c;
    }
    *{box-sizing:border-box}
    body{margin:0;background:var(--bg);color:var(--ink);font:14px/1.5 Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif}
    header .sitebar{
      max-width:1200px;margin:0 auto;padding:16px 24px;
      display:flex;align-items:center;justify-content:space-between;gap:16px;border-bottom:1px solid var(--line);background:#fff
    }
    .brand{display:flex;align-items:center;gap:10px;text-decoration:none;color:var(--ink);font-weight:900}
    .brand .dot{width:10px;height:10px;border-radius:999px;background:linear-gradient(90deg,var(--brand),var(--brand-2))}
    .brand .name{letter-spacing:.06em;text-transform:uppercase;font-size:12px}
    header nav a{font-weight:800;text-decoration:none;color:#0b1220;margin-left:18px}
    .wrap{max-width:1100px;margin:20px auto;padding:0 24px}
    .grid{display:grid;grid-template-columns:2fr 1fr;gap:16px}
    @media (max-width:980px){ .grid{grid-template-columns:1fr} }
    .card{background:var(--card);border:1px solid var(--line);border-radius:20px;box-shadow:var(--shadow);padding:18px}
    .card h2{margin:0 0 10px;font-size:18px}
    .muted{color:var(--muted)}
    .pill{display:inline-block;padding:2px 8px;border-radius:999px;background:#eef2ff;font-weight:800;font-size:12px}
    table{width:100%;border-collapse:collapse}
    th,td{padding:10px;border-bottom:1px solid #f1f5f9;text-align:left}
    th{background:#f8fafc;position:sticky;top:0;z-index:1}
    .row{display:flex;gap:8px;flex-wrap:wrap}
    input,select{padding:6px;border:1px solid var(--line);border-radius:12px}
    button,.btn{padding:6px 10px;border-radius:12px;border:1px solid var(--line);background:#fff;font-weight:800;cursor:pointer;text-decoration:none;color:#0b1220}
    .btn.primary,button.primary{background:linear-gradient(90deg,var(--brand),var(--brand-2));border:none;color:#fff}
    .btn.danger{border-color:#fecaca;background:#fee2e2;color:#7f1d1d}
    .badge{font-weight:800}
    .b-ok{color:var(--ok)} .b-warn{color:var(--warn)} .b-bad{color:var(--bad)}
    .inline-form{display:flex;gap:6px;align-items:center;flex-wrap:wrap}
    .right{margin-left:auto}
  </style>
</head>
<body>
  <header>
    <div class="sitebar">
      <a class="brand" href="/"><div class="dot"></div><div class="name">Lustra • Director</div></a>
      <nav>
        <a href="/app">App</a>
        <a href="/pricing">Pricing</a>
        <a href="/director/logout">Log out</a>
      </nav>
    </div>
  </header>

  <div class="wrap">
    <div class="grid">
      <div class="col">
        <div class="card" id="orgCard">
          <h2>Organization</h2>
          <div class="row">
            <div><strong>Pool balance:</strong> <span id="orgBalance" class="badge"></span></div>
            <div class="muted">(<span id="orgName"></span>)</div>
          </div>
        </div>

        <div class="card">
          <h2>Users</h2>
          <div class="muted" style="margin:-6px 0 10px">Enable/disable, reset passwords, set monthly caps, or delete users.</div>
          <div style="overflow:auto">
            <table id="usersTable">
              <thead>
                <tr>
                  <th>User</th>
                  <th>ID</th>
                  <th>Status</th>
                  <th>Balance</th>
                  <th>Monthly cap</th>
                  <th class="right">Actions</th>
                </tr>
              </thead>
              <tbody id="usersBody"></tbody>
            </table>
          </div>

          <hr style="border:none;border-top:1px solid #eef2ff;margin:14px 0">
          <h3 style="margin:0 0 8px">Create user</h3>
          <form id="createUserForm" class="inline-form">
            <input name="username" placeholder="username" required>
            <input name="password" type="password" placeholder="password" required>
            <button class="primary" type="submit">Create</button>
            <span id="createUserMsg" class="muted"></span>
          </form>
        </div>

        <div class="card">
          <h2>Recent activity</h2>
          <div class="muted" style="margin:-6px 0 10px">Polishes & credit events (last 50)</div>
          <div style="overflow:auto">
            <table id="eventsTable">
              <thead><tr><th>When</th><th>User</th><th>Event</th><th>Details</th></tr></thead>
              <tbody id="eventsBody"></tbody>
            </table>
          </div>
        </div>
      </div>

      <div class="col">
        <div class="card">
          <h2>Plans & top-ups</h2>
          <p class="muted">Buy packs or subscribe monthly without leaving this console.</p>
          <div class="row">
            <a href="/pricing" class="btn">See plans</a>
            <a href="/start" class="btn primary">Start / talk to sales</a>
          </div>
        </div>

        <div class="card">
          <h2>Change director password</h2>
          <form id="dirPassForm" class="inline-form">
            <input type="password" name="newpass" placeholder="new director password" required>
            <button class="primary" type="submit">Update</button>
            <span id="dirPassMsg" class="muted"></span>
          </form>
        </div>
      </div>
    </div>
  </div>

  <script>
    const $ = (s) => document.querySelector(s);
    const esc = (s) => (s==null?"":String(s).replace(/[&<>"]/g,c=>({ '&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c])));

    async function fetchJSON(url, opts){
      const r = await fetch(url, opts||{});
      if(!r.ok) throw new Error(await r.text().catch(()=>r.statusText));
      return r.json();
    }

    async function loadOrg(){
      const dash = await fetchJSON('/__admin/dashboard');
      $('#orgName').textContent = dash.orgName || '';
      $('#orgBalance').textContent = (dash.pool && dash.pool.balance!=null) ? dash.pool.balance : '-';
      renderEvents(dash.recent||[]);
    }

    async function loadUsers(){
      const data = await fetchJSON('/director/api/users');
      const tb = $('#usersBody'); tb.innerHTML = '';
      for(const u of (data.users||[])){
        const tr = document.createElement('tr');
        const status = u.active ? '<span class="pill">active</span>' : '<span class="pill" style="background:#fee2e2;color:#7f1d1d">disabled</span>';
        const bal = (u.balance==null?'—':u.balance);
        const capCtl = `
          <form class="inline-form" onsubmit="return setCap(${u.id}, this)">
            <input name="cap" type="number" min="0" step="1" placeholder="none" style="width:90px">
            <button type="submit">Set</button>
          </form>
        `;
        const actions = `
          <div class="inline-form right">
            <button onclick="toggleActive(${u.id}, ${u.active?0:1})">${u.active?'Disable':'Enable'}</button>
            <button onclick="resetUserPass(${u.id})">Reset pwd</button>
            <button class="danger" onclick="deleteUser(${u.id}, '${esc(u.username)}')">Delete</button>
          </div>
        `;
        tr.innerHTML = `
          <td>${esc(u.username)}</td>
          <td>${u.id}</td>
          <td>${status}</td>
          <td>${bal}</td>
          <td>${capCtl}</td>
          <td>${actions}</td>
        `;
        tb.appendChild(tr);
      }
    }

    function renderEvents(list){
      const tb = $('#eventsBody'); tb.innerHTML = '';
      for(const e of list){
        const when = e.ts || e.created_at || '';
        const who = e.username || '';
        const what = e.reason || (e.delta!=null?'credits':'polish');
        const details = (e.delta!=null) ? ('Δ ' + e.delta + (e.reason?(' · '+e.reason):'')) : (e.candidate||e.filename||'');
        const tr = document.createElement('tr');
        tr.innerHTML = `<td>${esc(when)}</td><td>${esc(who)}</td><td>${esc(what)}</td><td>${esc(details)}</td>`;
        tb.appendChild(tr);
      }
    }

    async function toggleActive(userId, active){
      await fetchJSON(`/director/api/user/set_active?user_id=${userId}&active=${active}`);
      loadUsers();
    }
    async function resetUserPass(userId){
      const p = prompt('New password for this user:');
      if(!p) return;
      await fetchJSON(`/director/api/user/reset_password?user_id=${userId}&password=${encodeURIComponent(p)}`);
      alert('Password updated.');
    }
    async function deleteUser(userId, uname){
      if(!confirm(`Delete user "${uname}"? This cannot be undone.`)) return;
      await fetchJSON(`/director/api/user/delete?user_id=${userId}`);
      loadUsers();
    }
    function setCap(userId, form){
      const raw = new FormData(form).get('cap');
      const v = (raw===''||raw==null) ? 'null' : String(parseInt(raw,10));
      fetchJSON(`/director/api/user/set-monthly-cap?user_id=${userId}&cap=${encodeURIComponent(v)}`)
        .then(()=>{ form.reset(); })
        .catch(err=>alert(err));
      return false;
    }

    document.querySelector('#createUserForm').addEventListener('submit', async (e)=>{
      e.preventDefault();
      const fd = new FormData(e.currentTarget);
      const r = await fetch('/director/users/create', { method:'POST', body: fd });
      if(r.ok){ document.querySelector('#createUserMsg').textContent = 'Created.'; e.currentTarget.reset(); loadUsers(); }
      else { document.querySelector('#createUserMsg').textContent = 'Error creating user.'; }
      setTimeout(()=>document.querySelector('#createUserMsg').textContent='', 2000);
    });

    document.querySelector('#dirPassForm').addEventListener('submit', async (e)=>{
      e.preventDefault();
      const fd = new FormData(e.currentTarget);
      const newpass = fd.get('newpass');
      if(!newpass) return;
      await fetchJSON(`/director/api/self/reset-password?newpass=${encodeURIComponent(newpass)}`);
      document.querySelector('#dirPassMsg').textContent = 'Updated.';
      e.currentTarget.reset();
      setTimeout(()=>document.querySelector('#dirPassMsg').textContent='', 2000);
    });

    (async function(){
      try{ await loadOrg(); await loadUsers(); }
      catch(e){ alert('Failed to load director console: ' + e.message); }
    })();
  </script>
</body>
</html>
"""

DIRECTOR_FORGOT_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Reset director password</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{--blue:#003366;--ink:#111827;--muted:#6b7280;--line:#e5e7eb;--bg:#f2f6fb;--card:#ffffff}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}
    .wrap{max-width:520px;margin:48px auto;padding:0 18px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:18px}
    label{font-weight:600;font-size:13px}
    input{width:100%;padding:10px;border:1px solid var(--line);border-radius:10px;margin-top:6px}
    button{width:100%;margin-top:12px;background:linear-gradient(90deg,var(--blue),#0a4d8c);color:#fff;border:none;border-radius:10px;padding:10px 16px;font-weight:700;cursor:pointer}
    .muted{color:var(--muted);font-size:12px;text-align:center;margin-top:8px}
    .err{margin-top:8px;color:#b91c1c;font-weight:700;font-size:12px}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <h3>Reset director password</h3>
      <!--RERR-->
      <form method="post" action="/director/forgot">
        <label>Reset code</label>
        <input type="text" name="code" required />
        <label>New director password</label>
        <input type="password" name="newpass" required />
        <button type="submit">Set new password</button>
      </form>
      <div class="muted">Back to <a href="/director">Director login</a></div>
    </div>
  </div>
</body>
</html>
"""

FORGOT_HTML = r"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Reset password</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{--blue:#003366;--ink:#111827;--muted:#6b7280;--line:#e5e7eb;--bg:#f2f6fb;--card:#ffffff}
    body{font-family:Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0}
    .wrap{max-width:520px;margin:48px auto;padding:0 18px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:18px}
    label{font-weight:600;font-size:13px}
    input{width:100%;padding:10px;border:1px solid var(--line);border-radius:10px;margin-top:6px}
    button{width:100%;margin-top:12px;background:linear-gradient(90deg,var(--blue),#0a4d8c);color:#fff;border:none;border-radius:10px;padding:10px 16px;font-weight:700;cursor:pointer}
    .muted{color:var(--muted);font-size:12px;text-align:center;margin-top:8px}
    .err{margin-top:8px;color:#b91c1c;font-weight:700;font-size:12px}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <h3>Reset password</h3>
      <!--FERR-->
      <form method="post" action="/forgot">
        <label>Username</label>
        <input type="text" name="username" required />
        <label>Reset code</label>
        <input type="text" name="code" required />
        <label>New password</label>
        <input type="password" name="newpass" required />
        <button type="submit">Reset</button>
      </form>
      <div class="muted">Back to <a href="/login">Sign in</a></div>
    </div>
  </div>
</body>
</html>
"""

app = Flask(__name__)
# Create DB tables on boot (no-op if DATABASE_URL is missing)
init_db()
# Ensure env admin exists in DB (idempotent)
seed_admin_user()


# ------------------------ session secret + default creds (unchanged) ------------------------
app.secret_key = os.getenv("APP_SECRET_KEY", "dev-secret-change-me")
APP_ADMIN_USER = os.getenv("APP_ADMIN_USER", "admin")
APP_ADMIN_PASS = os.getenv("APP_ADMIN_PASS", "hamilton")

# ------------------------ Gate protected routes (/app, /polish, /stats, /director*) ------------------------
@app.before_request
def gate_protected_routes():
    protected_prefixes = ["/app", "/polish", "/stats", "/director", "/skills"]
    p = request.path or "/"
    if any(p.startswith(x) for x in protected_prefixes):
        if not session.get("authed"):
            return redirect(url_for("login"))

# ------------------------ Public routes ------------------------
@app.get("/")
def home():
    resp = make_response(render_template_string(HOMEPAGE_HTML))
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.get("/about")
def about():
    resp = make_response(render_template_string(ABOUT_HTML))
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.get("/pricing")
def pricing():
    resp = make_response(render_template_string(PRICING_HTML))
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.get("/trial")
def start_trial():
    # Keep buttons working: send them to the new Start Free Trial form
    return redirect("/start")

# at top of file (once)
from secrets import token_hex
import os, re

@app.get("/start")
def contact_get():
    # CSRF token for the form
    if not session.get("csrf"):
        session["csrf"] = token_hex(16)
    html = CONTACT_HTML.replace("{{csrf}}", session["csrf"])
    resp = make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.post("/start")
def contact_post():
    # CSRF + honeypot (supports old and new field names)
    if (request.form.get("csrf") or "") != (session.get("csrf") or ""):
        return "invalid csrf", 400
    if (request.form.get("website") or request.form.get("company_website") or "").strip():
        return "ok", 200  # silent drop for bots

    company   = (request.form.get("company") or "").strip()
    email     = (request.form.get("email") or "").strip()
    name      = (request.form.get("name") or "").strip()
    volume    = (request.form.get("volume") or "").strip()
    users     = (request.form.get("users") or "").strip()
    templates = (request.form.get("templates") or "").strip()
    need_sso  = (request.form.get("need_sso") or "").strip()
    message   = (request.form.get("message") or "").strip()
    ip        = request.headers.get("X-Forwarded-For","").split(",")[0].strip() or request.remote_addr or ""
    ua        = request.headers.get("User-Agent","").strip()

    # optional file
    upname = ""
    try:
        f = request.files.get("file")
        if f and (f.filename or "").strip():
            fn = re.sub(r"[^A-Za-z0-9_.-]+","_", f.filename.strip())
            if len(fn) > 60: fn = fn[-60:]
            path = os.path.join("/tmp", f"lead_" + token_hex(4) + "_" + fn)
            f.save(path)
            upname = os.path.basename(path)
    except Exception:
        pass

    try:
        db_execute("""
          INSERT INTO leads (company,email,name,volume,users,templates,need_sso,message,filename,ip,user_agent)
          VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """, (company,email,name,volume,users,templates,need_sso,message,upname,ip,ua))
    except Exception as e:
        print("lead insert failed:", e)

    # thank-you view
    thanks = """
    <div class="pagebox">
      <h1>Thanks — we’ve got it!</h1>
      <p class="sub">We’ll reach out within 1 business day. You can also email <a href="mailto:hello@lustra.uk">hello@lustra.uk</a>.</p>
      <p><a class="btn" href="/login" style="width:auto;padding-left:20px;padding-right:20px">Sign in to start your 5-CV trial</a></p>
    </div>
    """
    html = CONTACT_HTML.replace(
        '<form method="POST" enctype="multipart/form-data" class="grid">', thanks
    ).replace("</form>", "")
    resp = make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})
    resp.headers["Cache-Control"] = "no-store"
    return resp
    
# ------------------------ Auth routes (unchanged, plus user DB support) ------------------------
@app.get("/login")
def login():
    if session.get("authed"):
        return redirect(url_for("app_page"))  # goes to /app
    try:
        resp = make_response(render_template_string(LOGIN_HTML))
        resp.headers["Cache-Control"] = "no-store"
        return resp
    except Exception as e:
        # This prints the real template error to your Render logs
        print("LOGIN_HTML render failed:", repr(e))
        return "Login template error. Check service logs for details.", 500

@app.post("/login")
def do_login():
    user = (request.form.get("username") or "").strip()
    pw = (request.form.get("password") or "").strip()

    # 1) Try Postgres first (preferred)
    try:
        rec = get_user_db(user)
        if rec and rec["active"] and check_password_hash(rec["password_hash"], pw):
            session["authed"] = True
            session["user"] = rec["username"]
            session["user_id"] = rec["id"]           # <-- store DB user id
            if is_admin():
                return redirect("/owner/console")
            return redirect(url_for("app_page"))

    except Exception as e:
        # non-fatal: fall through to legacy methods
        print("DB login check failed:", e)

    # 2) Legacy env-admin fallback (kept for compatibility)
    if user == APP_ADMIN_USER and pw == APP_ADMIN_PASS:
        session["authed"] = True
        session["user"] = user
        # assign stable numeric user_id for legacy admin
        try:
            import hashlib
            uname = (session.get("user") or "").strip().lower()
            session["user_id"] = int(hashlib.sha1(uname.encode("utf-8")).hexdigest()[:8], 16)
        except Exception:
            session["user_id"] = 0
        return redirect("/owner/console")


    # 3) Legacy users.json fallback (until we migrate)
    u = _get_user(user)
    if u and u.get("active", True) and pw == u.get("password", ""):
        session["authed"] = True
        session["user"] = user
        # assign stable numeric user_id (prefer id in users.json, else hash of username)
        try:
            import hashlib
            uid = u.get("id")
            if uid is None:
                uname = (session.get("user") or "").strip().lower()
                uid = int(hashlib.sha1(uname.encode("utf-8")).hexdigest()[:8], 16)
            session["user_id"] = int(uid)
        except Exception:
            session["user_id"] = 0
        if is_admin():
            return redirect("/owner/console")
        return redirect(url_for("app_page"))


    # Fail
    html = LOGIN_HTML.replace("<!--ERROR-->", "<div class='err'>Invalid credentials</div>")
    resp = make_response(render_template_string(html))
    resp.headers["Cache-Control"] = "no-store"
    return resp, 401

@app.get("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ---------- forgot password (for recruiter users in users.json) ----------
@app.get("/forgot")
def forgot_get():
    resp = make_response(render_template_string(FORGOT_HTML))
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.post("/forgot")
def forgot_post():
    username = (request.form.get("username") or "").strip()
    code = (request.form.get("code") or "").strip()
    newpass = (request.form.get("newpass") or "")
    if code != RESET_CODE:
        html = FORGOT_HTML.replace("<!--FERR-->", "<div class='err'>Invalid reset code</div>")
        return render_template_string(html), 400
    u = _get_user(username)
    if not u:
        html = FORGOT_HTML.replace("<!--FERR-->", "<div class='err'>User not found</div>")
        return render_template_string(html), 404
    u["password"] = newpass
    _save_users()
    return redirect(url_for("login"))

# ---------- serve the logo ----------
@app.get("/logo")
def logo():
    for name in ["Imagem1.png", "hamilton_logo.png", "logo.png"]:
        p = PROJECT_DIR / name
        if p.exists():
            resp = make_response(send_file(str(p)))
            resp.headers["Cache-Control"] = "no-store"
            return resp
    return ("", 204)

# ---------- helper: Word field ----------
def _add_field(paragraph, instr_text: str):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr_text)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = ""
    r.append(t); fld.append(r)
    paragraph._p.append(fld)

# ---------- Extraction ----------
def extract_text_any(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        if fitz is not None:
            try:
                parts = []
                with fitz.open(str(path)) as doc:
                    for page in doc:
                        parts.append(page.get_text("text"))
                return "\n".join(parts) or ""
            except Exception:
                pass
        return pdf_extract_text(str(path)) or ""
    elif ext == ".docx":
        d = Docx(str(path))
        parts = []
        for p in d.paragraphs:
            if p.text: parts.append(p.text)
        for table in d.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    else:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except:
            return ""

# ---------- AI structuring ----------
SCHEMA_PROMPT = """
You are a CV structuring assistant for recruiters. Extract ONLY what exists in the CV and return STRICT JSON:

{
  "personal_info": {"full_name":"","email":"","phone":"","location":"","links":[]},
  "summary":"",
  "experience":[{"job_title":"","company":"","location":"","start_date":"","end_date":"","currently_employed":false,"bullets":[],"raw_text":""}],
  "education":[{"degree":"","institution":"","location":"","start_date":"","end_date":"","bullets":[]}],
  "skills":[],
  "certifications":[],
  "languages":[],
  "awards":[],
  "other":[{"section_title":"","items":[]}]
}

Rules:
- Do NOT invent or embellish content.
- Preserve wording verbatim where possible.
- Use "currently_employed": true when the role is ongoing; leave "end_date" empty in that case.
"""

def ai_or_heuristic_structuring(cv_text: str) -> dict:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if api_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            resp = client.chatCompletions.create(
                model=MODEL,
                messages=[{"role":"system","content":SCHEMA_PROMPT},
                          {"role":"user","content":cv_text}],
                temperature=0
            )
            out = resp.choices[0].message.content.strip()
            if out.startswith("```"):
                out = out.strip("`")
                if out.lower().startswith("json"):
                    out = out[4:].strip()
            import json as _json
            return _json.loads(out)
        except Exception:
            try:
                from openai import OpenAI
                client = OpenAI()
                resp = client.chat.completions.create(
                    model=MODEL,
                    messages=[{"role":"system","content":SCHEMA_PROMPT},
                              {"role":"user","content":cv_text}],
                    temperature=0
                )
                out = resp.choices[0].message.content.strip()
                if out.startswith("```"):
                    out = out.strip("`")
                    if out.lower().startswith("json"):
                        out = out[4:].strip()
                import json as _json
                return _json.loads(out)
            except Exception as e:
                print("OpenAI failed, falling back to heuristics:", e)
                print(traceback.format_exc())

    # Fallback heuristic
    blocks = {"summary":[],"experience":[],"education":[],"skills":[],"certifications":[],"languages":[],"awards":[]}
    current = "summary"
    for ln in cv_text.splitlines():
        U = ln.upper()
        if re.search(r'\b(EXPERIENCE|EMPLOYMENT|CAREER|PROFESSIONAL EXPERIENCE)\b', U):
            current = "experience"; continue
        if "EDUCATION" in U or "QUALIFICATIONS" in U:
            current = "education"; continue
        if "SKILLS" in U:
            current = "skills"; continue
        if "CERTIFICATION" in U or "QUALIFICATION" in U:
            current = "certifications"; continue
        if "LANGUAGE" in U:
            current = "languages"; continue
        if "AWARD" in U or "HONOR" in U:
            current = "awards"; continue
        blocks[current].append(ln.strip())

    def pack(bl): return [x for x in bl if x][:30]
    return {
        "personal_info":{"full_name":"","email":"","phone":"","location":"","links":[]},
        "summary":" ".join(blocks["summary"])[:2000],
        "experience":[{"job_title":"","company":"","location":"","start_date":"","end_date":"","currently_employed":False,"bullets":pack(blocks["experience"]), "raw_text":""}],
        "education":[{"degree":"","institution":"","location":"","start_date":"","end_date":"","bullets":pack(blocks["education"])}],
        "skills":pack(blocks["skills"]),
        "certifications":pack(blocks["certifications"]),
        "languages":pack(blocks["languages"]),
        "awards":pack(blocks["awards"]),
        "other":[]
    }

# ---------- Skills booster (keywords only) ----------
SKILL_CANON = [
    "Ability to Work Independently",
    "Administration",
    "Advanced Excel",
    "AICPA",
    "Alteryx",
    "Anaplan",
    "Annuity",
    "Asset Management",
    "Assurance",
    "Audit",
    "Audit & External Reporting",
    "Audit Completion",
    "Audit Engagement Management",
    "Audit Execution",
    "Audit Planning",
    "AXIS",
    "BMA",
    "BMA EBS Reporting",
    "BMA Regulatory Reporting",
    "BMA Reporting",
    "Big Four",
    "Bermuda Statutory Accounting",
    "BSCR",
    "CALM",
    "Capital Adequacy Analysis",
    "Capital and Risk Management",
    "Captive Insurance",
    "Catastrophe Modelling",
    "Ceded Reinsurance",
    "Claims Liabilities",
    "Clearwater",
    "Client Relationships",
    "Commercial Insurance",
    "Communication (Written and Verbal)",
    "Consolidations & Group Reporting",
    "Corporate Insolvency",
    "Cost and Management Accounting",
    "Credit Risk",
    "Due Diligence",
    "EBS",
    "Emblem Modelling",
    "Excel",
    "External Audit",
    "External Audit of Quarterly and Annual Financial Statements",
    "Federal, State, and International Tax Compliance",
    "Fiduciary Services",
    "Financial Audit",
    "Financial Disclosures",
    "Financial Modelling",
    "Financial Modeling",
    "Financial Reporting",
    "Financial Services",
    "Financial Statement Consolidation",
    "Financial Statement Review",
    "Financial Statements Preparation and Analysis",
    "FRS 102",
    "Fund of Funds",
    "Funds",
    "GAAP",
    "GAAS",
    "Governance, Risk and Regulatory Reviews",
    "Hedge Funds",
    "High Level of Business Acumen",
    "IFRS",
    "IFRS 4",
    "IFRS 17",
    "IFRS Reporting",
    "ILS",
    "ILS Accounting",
    "IND AS",
    "Insurance",
    "Insurance Accounting",
    "Insurance Pricing",
    "Internal Audit Liaison",
    "Internal Control Testing",
    "Internal Controls",
    "Investment",
    "Investment & Treasury Accounting",
    "Investment Audits",
    "Investment Consulting",
    "ISA",
    "ISAE 3402",
    "LICAT",
    "Liquidations",
    "Life and Health Actuary",
    "Life Insurance",
    "Loss Reserves",
    "LDTI",
    "MFMA",
    "Microsoft Excel",
    "Microsoft Outlook",
    "Microsoft PowerPoint",
    "Microsoft Word",
    "NatCat",
    "Offshore Experience",
    "Operational Leadership",
    "Oracle",
    "ORSA",
    "PCAOB",
    "PCAOB Audit",
    "Pensions",
    "PFMA",
    "PL/SQL",
    "Power BI",
    "Power Query",
    "Premium Calculations",
    "Private Equity",
    "Process Automation",
    "Product Development",
    "Progress Report Creation",
    "Project Management",
    "Prophet",
    "Public Sector Accounting Standards",
    "Python",
    "R",
    "Regulatory Compliance",
    "Regulatory Compliance & Filings",
    "Reporting",
    "Research Skills",
    "Reserving",
    "Restructuring",
    "Reinsurance",
    "Reinsurance & ILS Accounting",
    "Reinsurance Audit",
    "Reinsurance Contracts",
    "Reinsurance Finance & Reporting",
    "Reinsurance Pricing",
    "Reinsurance Tender Negotiations",
    "Risk Assessment",
    "Risk Management",
    "Risk-Based Auditing",
    "Risk-Based Audits",
    "SAP",
    "SEC",
    "SICS",
    "SOC 1",
    "SOC 2",
    "Solvency II",
    "Solvency Margin Review",
    "SQL",
    "Stakeholder & Investor Relations Support",
    "Statutory Reporting",
    "Strong Analytical and Problem-Solving",
    "Strong Organizational and Interpersonal Skills",
    "Tax Audits",
    "Tax Consulting",
    "Tax Planning",
    "Tax Research",
    "Taxation for Business and Individuals",
    "Top 10 Audit",
    "Trust & Company Administration",
    "UK GAAP",
    "Underwriting",
    "US GAAP",
    "US GAAP Reporting",
    "US STAT",
    "US Tax",
    "Valuation of Privately-Held Investments",
    "Valuations",
    "Variance Analysis",
    "Variance Investigation",
    "VBA",
    "Wealth and Asset Management",
    "Wealth Structuring",
    "Workiva",
]
# --- Per-client skills config (custom skills + disabled built-ins) ---
SKILLS_FILE = PROJECT_DIR / "skills.json"

def _load_skills_config():
    try:
        if SKILLS_FILE.exists():
            return json.loads(SKILLS_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    # default: no custom, nothing disabled
    return {"custom": [], "base_disabled": []}

SKILLS_CFG = _load_skills_config()

def _save_skills_config():
    SKILLS_FILE.write_text(json.dumps(SKILLS_CFG, indent=2), encoding="utf-8")

def _effective_skills():
    """Built-ins minus disabled + custom (dedup, case-insensitive)."""
    disabled = {s.lower() for s in SKILLS_CFG.get("base_disabled", [])}
    base = [s for s in SKILL_CANON if s.lower() not in disabled]
    custom = [s.strip() for s in SKILLS_CFG.get("custom", []) if isinstance(s, str) and s.strip()]
    eff, seen = [], set()
    for s in base + custom:
        k = s.lower()
        if k not in seen:
            seen.add(k); eff.append(s)
    return eff
def extract_top_skills(text: str):
    tokens = re.findall(r"[A-Za-z0-9\-\&\./+]+", text)
    txt_up = " ".join(tokens).upper()
    canon = _effective_skills()
    found, seen = [], set()
    for s in canon:
        if s.upper() in txt_up:
            k = s.lower()
            if k not in seen:
                seen.add(k)
                found.append(s)
    return found[:25]

# ---------- Word helpers ----------
SOFT_BLACK = RGBColor(64, 64, 64)

def _tone_runs(paragraph, size=11, bold=False, color=SOFT_BLACK, name="Calibri"):
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bool(bold)
        run.font.color.rgb = color

def _add_center_line(doc: Docx, text: str, size=11, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(size); r.bold = bool(bold); r.font.color.rgb = SOFT_BLACK
    return p

def _add_section_heading(doc: Docx, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text.upper().strip())
    r.font.name = "Calibri"; r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = SOFT_BLACK
    return p

def _remove_all_body_content(doc: Docx):
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

# ---------- Post-save zip scrub of header XML ----------
def _zip_scrub_header_labels(docx_path: Path):
    pat_one = re.compile(
        r'<w:p\b[^>]*>.*?(?:professional).*?(?:experience).*?(?:continued).*?</w:p>',
        re.I | re.S
    )
    pat_two = re.compile(
        r'(<w:p\b[^>]*>.*?(?:professional).*?(?:experience).*?</w:p>)\s*(<w:p\b[^>]*>.*?(?:continued).*?</w:p>)',
        re.I | re.S
    )
    blank_p = '<w:p><w:r><w:t> </w:t></w:r></w:p>'

    src = str(docx_path)
    tmp = str(docx_path.with_suffix('.tmp.docx'))

    with zipfile.ZipFile(src, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith('word/header') and item.filename.endswith('.xml'):
                xml = data.decode('utf-8', errors='ignore')
                xml = pat_two.sub(blank_p, xml)
                xml = pat_one.sub(blank_p, xml)
                data = xml.encode('utf-8')
            zout.writestr(item, data)

    os.replace(tmp, src)

# ---------- Ensure a spacer paragraph in primary headers (pages 2+) ----------
def _ensure_primary_header_spacer(doc: Docx):
    try:
        for sec in doc.sections:
            hdr = sec.header
            needs = True
            try:
                if hdr.paragraphs:
                    if (hdr.paragraphs[-1].text or "").strip() == "":
                        needs = False
            except Exception:
                pass
            if needs:
                p = hdr.add_paragraph()
                r = p.add_run(" ")
    except Exception:
        pass

# ---------- Compose CV ----------
def build_cv_document(cv: dict, template_override: str | None = None) -> Path:
    # Prefer an explicit override (per-org), otherwise fall back to bundled templates
    tpath = Path(template_override) if template_override else None
    if tpath and tpath.exists():
        template_path = tpath
    else:
        template_path = None
        for pth in [PROJECT_DIR / "hamilton_template.docx",
                    PROJECT_DIR / "HAMILTON TEMPLATE.docx",
                    PROJECT_DIR / "master_template.docx"]:
            if pth.exists():
                template_path = pth
                break

    if template_path:
        doc = Docx(str(template_path))
    else:
        doc = Docx()

    _remove_all_body_content(doc)
    # Profile-based labels (optional, per-org)
    labels = {
        "summary": "EXECUTIVE SUMMARY",
        "certifications": "PROFESSIONAL QUALIFICATIONS",
        "skills": "PROFESSIONAL SKILLS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "references": "REFERENCES",
    }
    try:
        oid = _current_user_org_id()
        if oid:
            row = db_query_one("SELECT profile_json FROM orgs WHERE id=%s", (oid,))
            import json
            prof = json.loads(row[0]) if row and row[0] else None
            if prof and prof.get("enable_profile") and isinstance(prof.get("labels"), dict):
                for k, v in prof["labels"].items():
                    if isinstance(v, str) and v.strip() and k in labels:
                        labels[k] = v.strip()
    except Exception as e:
        print("profile labels failed:", e)

    spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(6)

    pi = (cv or {}).get("personal_info") or {}
    full_name = (pi.get("full_name") or "Candidate").strip()
    name_p = doc.add_paragraph(); name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; name_p.paragraph_format.space_after = Pt(2)
    name_r = name_p.add_run(full_name)
    name_r.font.name="Calibri"; name_r.font.size=Pt(18); name_r.bold=True; name_r.font.color.rgb=SOFT_BLACK

    tel = (pi.get("phone") or "").strip()
    email = (pi.get("email") or "").strip()
    location = (pi.get("location") or "").strip()
    bits = []
    if tel: bits.append(f"Tel: {tel}")
    if email: bits.append(f"Email: {email}")
    if bits:
        _add_center_line(doc, " | ".join(bits), size=11, bold=False, space_after=0)
    if location:
        _add_center_line(doc, f"Location: {location}", size=11, bold=False, space_after=6)
    links = [s for s in (pi.get("links") or []) if s]
    if links: _add_center_line(doc, " | ".join(links), size=11, bold=False, space_after=6)

    if cv.get("summary"):
        _add_section_heading(doc, labels["summary"])
        p = doc.add_paragraph(cv["summary"]); p.paragraph_format.space_after = Pt(8); _tone_runs(p, size=11, bold=False)

    quals = []
    if cv.get("certifications"): quals += [q for q in cv["certifications"] if q]
    edu = cv.get("education") or []
    for ed in edu:
        deg = (ed.get("degree") or "").strip()
        inst = (ed.get("institution") or "").strip()
        date_span = " – ".join([x for x in [(ed.get("start_date") or "").strip(), (ed.get("end_date") or "").strip()] if x]).strip(" –")
        line = " | ".join([s for s in [deg, inst, date_span] if s])
        if line: quals.append(line)
    if quals:
        _add_section_heading(doc, labels["certifications"])
        for q in quals:
            p = doc.add_paragraph(q)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            _tone_runs(p, size=11, bold=False)

    skills = cv.get("skills") or []
    if skills:
        _add_section_heading(doc, labels["skills"])
        line = " | ".join(skills)
        p = doc.add_paragraph(line); p.paragraph_format.space_after = Pt(8); _tone_runs(p, size=11, bold=False)

    exp = cv.get("experience") or []
    if exp:
        _add_section_heading(doc, labels["experience"])
        first = True
        for role in exp:
            if not first:
                g = doc.add_paragraph()
                g.paragraph_format.space_after = Pt(8)
                _tone_runs(g, size=11, bold=False)
            first = False

            title_company = " — ".join([x for x in [role.get("job_title",""), role.get("company","")] if x]).strip()
            p = doc.add_paragraph(); r = p.add_run(title_company or "Role")
            r.font.name="Calibri"; r.font.size=Pt(11); r.bold=True; r.font.color.rgb=SOFT_BLACK
            p.paragraph_format.space_after = Pt(0)

            sd = (role.get("start_date") or "").strip()
            edd = (role.get("end_date") or "").strip()
            if role.get("currently_employed") and not edd: edd = "Present"
            dates = f"{sd} – {edd}".strip(" –")
            loc = (role.get("location") or "").strip()
            meta = " | ".join([x for x in [dates, loc] if x])
            if meta:
                meta_p = doc.add_paragraph(meta)
                meta_p.paragraph_format.space_after = Pt(6)
                _tone_runs(meta_p, size=11, bold=False)

            if role.get("bullets"):
                for b in role["bullets"]:
                    bp = doc.add_paragraph(b, style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(0)
                    _tone_runs(bp, size=11, bold=False)
            elif role.get("raw_text"):
                rp = doc.add_paragraph(role["raw_text"])
                rp.paragraph_format.space_after = Pt(0)
                _tone_runs(rp, size=11, bold=False)

    # --- Education ---
    if edu:
        _add_section_heading(doc, labels["education"])
        for ed in edu:
            line = " — ".join([x for x in [ed.get("degree",""), ed.get("institution","")] if x]).strip()
            p = doc.add_paragraph(); rr = p.add_run(line or "Education")
            rr.font.name="Calibri"; rr.font.size=Pt(11); rr.bold=True; rr.font.color.rgb=SOFT_BLACK
            p.paragraph_format.space_after = Pt(0)

            sd = (ed.get("start_date") or "").strip()
            ee = (ed.get("end_date") or "").strip()
            dates = f"{sd} – {ee}".strip(" –")
            loc = (ed.get("location") or "").strip()
            meta = " | ".join([x for x in [dates, loc] if x])
            if meta:
                meta_p = doc.add_paragraph(meta)
                meta_p.paragraph_format.space_after = Pt(2)
                _tone_runs(meta_p, size=11, bold=False)

            if ed.get("bullets"):
                for b in ed["bullets"]:
                    bp = doc.add_paragraph(b, style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(0)
                    _tone_runs(bp, size=11, bold=False)

    # --- References (fixed text) ---
    _add_section_heading(doc, labels["references"])
    p = doc.add_paragraph("Full references are available on request")
    p.paragraph_format.space_after = Pt(0)
    _tone_runs(p, size=11, bold=False)

    _ensure_primary_header_spacer(doc)

    out = PROJECT_DIR / "polished_cv.docx"
    doc.save(str(out))
    _zip_scrub_header_labels(out)
    return out

# ---------- helpers ----------
def _downloads_this_month():
    try:
        now = datetime.now()
        ym = (now.year, now.month)
        count = 0
        for item in STATS.get("history", []):
            ts = item.get("ts")
            if not ts: continue
            d = datetime.strptime(ts, "%Y-%m-%d %H:%M:%S")
            if (d.year, d.month) == ym:
                count += 1
        return count
    except Exception:
        return STATS.get("downloads", 0)

def _count_since(months: int) -> int:
    try:
        cutoff = datetime.now() - timedelta(days=30*months)
        c = 0
        for item in STATS.get("history", []):
            ts = item.get("ts")
            if not ts: continue
            d = datetime.strptime(ts, "%Y-%m-%d %H:%M:%S")
            if d >= cutoff:
                c += 1
        return c
    except Exception:
        return 0

# ---------- App + API ----------
APP_HTML = HTML

@app.get("/app")
def app_page():
    if is_admin():
        return redirect("/owner/console")
    # Render template
    html = render_template_string(
        HTML,
        show_director_link=bool(is_admin() or session.get("director"))
    )

    # Inject Director link (if admin/director)
    if is_admin() or session.get("director"):
        html = html.replace(
            "</body>",
            (
                '<a href="/director" class="dir-link" title="Director dashboard">Director dashboard</a>'
                '<style>'
                '.dir-link{position:fixed;right:16px;bottom:16px;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;'
                'background:#fff;color:#0f172a;text-decoration:none;box-shadow:0 1px 2px rgba(0,0,0,0.06);'
                'font:14px/1.2 -apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Arial,sans-serif}'
                '.dir-link:hover{box-shadow:0 2px 6px rgba(0,0,0,0.12)}'
                '</style>'
                '</body>'
            )
        )

        # Inject Owner link (admins only)
    if is_admin():
        html = html.replace(
            "</body>",
            (
                '<a href="/owner/console" class="owner-link" title="Owner console">Owner</a>'
                '<style>'
                '.owner-link{position:fixed;right:100px;bottom:16px;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;'
                'background:#fff;color:#0f172a;text-decoration:none;box-shadow:0 1px 2px rgba(0,0,0,0.06);'
                'font:14px/1.2 -apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Arial,sans-serif}'
                '.owner-link:hover{box-shadow:0 2px 6px rgba(0,0,0,0.12)}'
                '</style>'
                '</body>'
            )
        )
    # Inject Skills toggle + lazy loader
    html = html.replace(
        "</body>",
        (
            '<script>(function(){'
            'var btn=document.getElementById("skillsToggle");'
            'var panel=document.getElementById("skillsCard");'
            'var loaded=false;'
            'async function loadSkills(){try{'
            ' const r=await fetch("/skills",{cache:"no-store"});'
            ' const j=await r.json();'
            ' var all=(j.effective||[]).slice().sort(function(a,b){return a.localeCompare(b)});'
            ' var allEl=document.getElementById("skillsAll"); var hdr=document.getElementById("skillsAllHeader");'
            ' if(allEl){allEl.style.display="block";allEl.innerHTML=all.map(function(s){var esc=s.replace(/"/g,"&quot;");return "<span class=\\"pill\\">"+s+" <button type=\\"button\\" class=\\"x\\" data-skill=\\""+esc+"\\" data-src=\\"all\\">×</button></span>";}).join("")}'
            ' if(hdr){hdr.style.display="block"}'
            ' var cust=document.getElementById("customSkills");'
            ' if(cust){var c=(j.custom||[]).slice().sort(function(a,b){return a.localeCompare(b)});'
            '  cust.innerHTML=c.length?c.map(function(s){var esc=s.replace(/"/g,"&quot;");return "<span class=\\"pill\\">"+s+" <button type=\\"button\\" class=\\"x\\" data-skill=\\""+esc+"\\" data-src=\\"custom\\">×</button></span>";}).join(""):"<span class=\\"muted\\">(none)</span>"}'
            ' var base=document.getElementById("baseSkills");'
            ' if(base){var dis=new Set(j.base_disabled||[]);'
            '  var b=(j.base||[]).filter(function(s){return !dis.has(s)}).sort(function(a,b){return a.localeCompare(b)});'
            '  base.innerHTML=b.length?b.map(function(s){var esc=s.replace(/"/g,"&quot;");return "<span class=\\"pill\\">"+s+" <button type=\\"button\\" class=\\"x\\" data-skill=\\""+esc+"\\" data-src=\\"base\\">×</button></span>";}).join(""):"<span class=\\"muted\\">(none)</span>"}'
            ' loaded=true;'
            ' window.__skillsState=j;'
            ' var skillsCardEl=document.getElementById("skillsCard");'
            ' if(skillsCardEl){'
            '   skillsCardEl.addEventListener("click", async function(ev){'
            '     var btn = ev.target && ev.target.closest(".pill .x");'
            '     if(!btn) return;'
            '     var skill = btn.getAttribute("data-skill") || "";'
            '     var src   = btn.getAttribute("data-src")   || "";'
            '     try{'
            '       if(src==="custom"){'
            '         await fetch("/skills/custom/remove",{method:"POST",headers:{"Content-Type":"application/x-www-form-urlencoded"},body:new URLSearchParams({skill})});'
            '       }else if(src==="base"){'
            '         var fd=new FormData(); fd.append("skill",skill); fd.append("action","disable");'
            '         await fetch("/skills/base/toggle",{method:"POST",body:fd});'
            '       }else{'
            '         var customSet = new Set(((window.__skillsState&&window.__skillsState.custom)||[]).map(function(s){return s.toLowerCase()}));'
            '         if(customSet.has(skill.toLowerCase())){'
            '           await fetch("/skills/custom/remove",{method:"POST",headers:{"Content-Type":"application/x-www-form-urlencoded"},body:new URLSearchParams({skill})});'
            '         }else{'
            '           var fd2=new FormData(); fd2.append("skill",skill); fd2.append("action","disable");'
            '           await fetch("/skills/base/toggle",{method:"POST",body:fd2});'
            '         }'
            '       }'
            '       loaded=false; await loadSkills();'
            '     }catch(e){ console.log("skill remove failed", e); }'
            '   });'
            ' }'
            '}catch(e){var allEl=document.getElementById("skillsAll");if(allEl)allEl.innerHTML="<span class=\\"muted\\">Could not load skills.</span>";}}'
            'if(btn&&panel){btn.addEventListener("click",async function(){'
            ' var show=(panel.style.display==="none"||panel.style.display==="");'
            ' panel.style.display=show?"block":"none";'
            ' btn.textContent=show?"Hide":"Show";'
            ' if(show && !loaded) await loadSkills();'
            '});}'
            'var addForm=document.getElementById("skillAddForm");'
            'if(addForm){addForm.addEventListener("submit",async function(ev){ev.preventDefault();'
            ' var inp=document.getElementById("skillInput"); var v=(inp&&inp.value||"").trim(); if(!v)return;'
            ' try{await fetch("/skills/custom/add",{method:"POST",headers:{"Content-Type":"application/x-www-form-urlencoded"},body:new URLSearchParams({skill:v})});'
            ' if(inp) inp.value=""; loaded=false; await loadSkills();}catch(e){console.log("add skill failed",e);}'
            '});}'
            '})();</script></body>'
        )
)

                # Tweak Session Stats fonts (smaller values + history; keep titles bigger)
    html = html.replace(
        "</body>",
        (
            '<style id="sessionStatsCSS">'
            '#sessionStats *{font-size:12px !important;}'
            '#sessionStats h2{font-size:20px !important;}'
            '#sessionStats .btn, #sessionStats button{font-size:12px !important;}'
            '</style>'
            '<script>(function(){try{'
            ' var hs=document.querySelectorAll("h2"), box=null;'
            ' for(var i=0;i<hs.length;i++){'
            '   var t=(hs[i].textContent||"").trim().toLowerCase();'
            '   if(t==="session stats"){ box=hs[i].closest(".card")||hs[i].parentElement; break; }'
            ' }'
            ' if(box){ box.id="sessionStats"; }'
            '}catch(e){console.log("stats css enforce failed",e);} })();</script></body>'
        )
    )
        # Force smaller values + history inside Session Stats (titles untouched)
    html = html.replace(
        "</body>",
        (
            '<script>(function(){try{'
            '  var hs=document.querySelectorAll("h2"), card=null;'
            '  for(var i=0;i<hs.length;i++){'
            '    var t=(hs[i].textContent||"").trim().toLowerCase();'
            '    if(t==="session stats"){ card=hs[i].closest(".card")||hs[i].parentElement; break; }'
            '  }'
            '  if(!card){return;}'
            '  function shrink(el){ try{el.style.fontSize="13px"; el.style.lineHeight="1.3";}catch(e){} }'
            '  var titles=["Downloads this month","Last Candidate","Last Polished","Credits Used"];'
            '  titles.forEach(function(title){'
            '    var nodes=card.querySelectorAll("*"), label=null;'
            '    for(var i=0;i<nodes.length;i++){'
            '      var n=nodes[i];'
            '      if(n.children.length===0){'
            '        var txt=(n.textContent||"").trim();'
            '        if(txt.toLowerCase()===title.toLowerCase()){ label=n; break; }'
            '      }'
            '    }'
            '    if(label){'
            '      var parent=label.parentElement;'
            '      if(parent){'
            '        var kids=parent.children;'
            '        for(var k=0;k<kids.length;k++){ if(kids[k]!==label){ shrink(kids[k]); } }'
            '      }'
            '    }'
            '  });'
            '  var hist=document.getElementById("history");'
            '  if(hist){'
            '    shrink(hist);'
            '    var items=hist.querySelectorAll("*");'
            '    for(var i=0;i<items.length;i++){ try{items[i].style.fontSize="13px";}catch(e){} }'
            '  }'
            '}catch(e){console.log("stats font tweak failed",e);} })();</script></body>'
        )
    )

        # Force smaller values + history inside Session Stats (titles untouched)
    html = html.replace(
        "</body>",
        (
            '<script>(function(){try{'
            '  var hs=document.querySelectorAll("h2"), card=null;'
            '  for(var i=0;i<hs.length;i++){'
            '    var t=(hs[i].textContent||"").trim().toLowerCase();'
            '    if(t==="session stats"){ card=hs[i].closest(".card")||hs[i].parentElement; break; }'
            '  }'
            '  if(!card){return;}'
            '  function shrink(el){ try{el.style.fontSize="13px"; el.style.lineHeight="1.3";}catch(e){} }'
            '  var titles=["Downloads this month","Last Candidate","Last Polished","Credits Used"];'
            '  titles.forEach(function(title){'
            '    var nodes=card.querySelectorAll("*"), label=null;'
            '    for(var i=0;i<nodes.length;i++){'
            '      var n=nodes[i];'
            '      if(n.children.length===0){'
            '        var txt=(n.textContent||"").trim();'
            '        if(txt.toLowerCase()===title.toLowerCase()){ label=n; break; }'
            '      }'
            '    }'
            '    if(label){'
            '      var parent=label.parentElement;'
            '      if(parent){'
            '        var kids=parent.children;'
            '        for(var k=0;k<kids.length;k++){ if(kids[k]!==label){ shrink(kids[k]); } }'
            '      }'
            '    }'
            '  });'
            '  var hist=document.getElementById("history");'
            '  if(hist){'
            '    shrink(hist);'
            '    var items=hist.querySelectorAll("*");'
            '    for(var i=0;i<items.length;i++){ try{items[i].style.fontSize="13px";}catch(e){} }'
            '  }'
            '}catch(e){console.log("stats font tweak failed",e);} })();</script></body>'
        )
    )

    # Inject Uploading/Processing/Downloading overlay + XHR downloader
    html = html.replace(
        "</body>",
        (
            '<style>'
            '#busyOverlay{position:fixed;inset:0;background:rgba(15,23,42,.55);display:none;z-index:9999}'
            '#busyBox{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);background:#fff;padding:16px 18px;border-radius:12px;'
            ' box-shadow:0 10px 30px rgba(0,0,0,.2);font:14px/1.4 -apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Arial,sans-serif}'
            '#busyMsg{font-weight:600}.busyBar{height:3px;background:#e5e7eb;overflow:hidden;border-radius:2px;margin-top:8px}'
            '.busyBar>div{width:30%;height:100%;animation:busy 1s linear infinite;background:#0ea5e9}'
            '@keyframes busy{0%{transform:translateX(-100%)}100%{transform:translateX(400%)}}'
            '</style>'
            '<div id="busyOverlay"><div id="busyBox"><div id="busyMsg">Preparing…</div>'
            '<div class="busyBar"><div></div></div></div></div>'
            '<script>(function(){'
            'var form=document.querySelector(\'form[action="/polish"]\')||document.querySelector(\'form[action^="/polish"]\');'
            'if(!form)return;'
            'var fileInput=form.querySelector(\'input[type="file"][name="cv"]\');'
            'var overlay=document.getElementById("busyOverlay");var msg=document.getElementById("busyMsg");'
            'function show(t){if(overlay)overlay.style.display="block";if(msg)msg.textContent=t;}'
            'function hide(){if(overlay)overlay.style.display="none";}'
            'form.addEventListener("submit",function(ev){'
            ' try{'
            '  if(!fileInput||!fileInput.files||!fileInput.files[0])return;'
            '  ev.preventDefault();'
            '  var fd=new FormData(form);'
            '  var xhr=new XMLHttpRequest();'
            '  xhr.open("POST",form.getAttribute("action")||"/polish",true);'
            '  xhr.responseType="blob";'
            '  var sawDownload=false;'
            '  show("Uploading…");'
            '  if(xhr.upload){xhr.upload.onprogress=function(e){if(e.lengthComputable&&e.loaded>=e.total){show("Processing…");}}}'
            '  xhr.onprogress=function(){if(!sawDownload){show("Downloading…");sawDownload=true;}};'
            '  xhr.onerror=function(){hide();form.submit();};'
            '  xhr.onload=function(){try{'
            '    if(xhr.status!==200){hide();form.submit();return;}'
            '    var disp=xhr.getResponseHeader("Content-Disposition")||"";'
            '    var m=/filename\\*=UTF-8\\\'\\\'([^;]+)|filename=\\"?([^\\"]+)\\"?/i.exec(disp);'
            '    var name=(m&&decodeURIComponent(m[1]||m[2]||"polished_cv.docx"))||"polished_cv.docx";'
            '    var blob=xhr.response;'
            '    var a=document.createElement("a");a.href=URL.createObjectURL(blob);a.download=name;'
            '    document.body.appendChild(a);a.click();setTimeout(function(){URL.revokeObjectURL(a.href);a.remove();hide();},100);'
            '    if(window.refreshStats) setTimeout(window.refreshStats, 300);'
            '  }catch(_){hide();form.submit();}};'
            '  xhr.send(fd);'
            ' }catch(_){hide();form.submit();}'
            '});'
            '})();</script></body>'
        )
    )
    # Inject Full History data loader (fires on first click)
    html = html.replace(
        "</body>",
        (
            '<script>(function(){'
            'var t=document.getElementById("historyToggle");'
            'var h=document.getElementById("history");'
            'var loaded=false;'
            'async function load(){'
            '  try{'
            '    const r=await fetch("/me/history",{cache:"no-store"});'
            '    const j=await r.json();'
            '    var rows=j.history||[];'
            '    if(!h) return;'
            '    h.innerHTML = rows.length'
            '      ? rows.map(function(it){'
            '          return "<div class=\\"row\\" style=\\"padding:6px 0;border-bottom:1px solid var(--line)\\">" +'
            '                 "<span class=\\"muted\\">"+(it.ts||"-")+"</span> — " +'
            '                 "<strong>"+(it.candidate||"-")+"</strong> " +'
            '                 "<span class=\\"muted\\">("+(it.filename||"-")+")</span>" +'
            '                 "</div>";'
            '        }).join("")'
            '      : "<div class=\\"muted\\">(no history yet)</div>";'
            '    loaded=true;'
            '  }catch(e){ if(h) h.innerHTML="<div class=\\"muted\\">Could not load history.</div>"; }'
            '}'
            'if(t){ t.addEventListener("click", function(){ if(!loaded) load(); }); }'
            '})();</script></body>'
        )
    )
    

    # --- Session Stats tiles: refresh on load and on demand ---
    html = html.replace("</body>", """
<script>
  // Fills: #downloadsMonth, #lastCandidate, #lastTime, #creditsUsed (and #creditsBalance if present)
  window.refreshStats = async function(){
    try {
      const r = await fetch('/me/dashboard', { cache: 'no-store' });
      if (!r.ok) return;
      const d = await r.json();
      const set = (sel, val) => { const el = document.querySelector(sel); if (el) el.textContent = (val ?? '').toString(); };

      set('#downloadsMonth', d.downloadsMonth);
      set('#lastCandidate', d.lastCandidate || '');

      if (d.lastTime) {
        const dt = new Date(d.lastTime);
        set('#lastTime', isNaN(dt.getTime()) ? d.lastTime : dt.toLocaleString());
      } else {
        set('#lastTime','');
      }

      // Credits used and (optional) balance
      // Credits left: prefer org/user remaining from /me/credits; fall back to balance from /me/dashboard
try {
  const mc = await fetch('/me/credits', { cache: 'no-store' });
  if (mc.ok) {
    const j = await mc.json();
    if (j && j.ok) {
      const left = (j.myRemainingThisMonth != null) ? j.myRemainingThisMonth
                 : (j.balance != null) ? j.balance
                 : null;
      const el = document.querySelector('#creditsLeft') || document.querySelector('#creditsUsed'); // fallback
      if (el) el.textContent = (left == null) ? '—' : String(left);
    }
  }
} catch(e) { /* ignore */ }

// Also keep the old dashboard call working (already present):
if (typeof d.creditsBalance === 'number') {
  const el = document.querySelector('#creditsLeft') || document.querySelector('#creditsUsed');
  if (el && el.textContent === '—') el.textContent = d.creditsBalance;
}

    } catch (e) {
      console.log('refreshStats failed', e);
    }
  };

  // Auto-run once when the page loads
  document.addEventListener('DOMContentLoaded', () => {
    if (window.refreshStats) window.refreshStats();
  });
</script>
</body>""")

    resp = make_response(html)
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.get("/stats")
def stats():
    # Default values (fallback if DB is not set)
    downloads_month = _downloads_this_month()
    last_candidate = STATS.get("last_candidate", "")
    last_time = STATS.get("last_time", "")
    paid_left = int((STATS.get("credits", {}) or {}).get("balance", 0))
    trial_left = int(session.get("trial_credits", 0))
    total_left = paid_left + trial_left

    # If DB is available, prefer DB usage counts
    if DB_POOL:
        row = db_query_one("SELECT COUNT(*) FROM usage_events WHERE ts >= (NOW() - interval '30 days')")
        if row:
            downloads_month = row[0]
        row2 = db_query_one("SELECT candidate, ts FROM usage_events ORDER BY ts DESC LIMIT 1")
        if row2:
            last_candidate, last_time = row2[0], row2[1].strftime("%Y-%m-%d %H:%M:%S")

    return jsonify({
    "ok": True,
    # your other fields…
    "downloads_this_month": downloads_month,   # NEW: what the front-end expects
    "downloads_month": downloads_month,        # OLD: kept for backward compatibility
    # if you previously had "downloads": … and other fields, keep them as they were
})
    # --- Per-user usage (for the app JS) ---

@app.get("/me/last-event")
def me_last_event():
    # last candidate + timestamp for this user; safe if missing
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    cand, ts = (None, None)
    if uid:
        try:
            cand, ts = last_event_for_user(uid)
        except Exception as e:
            print("me_last_event error:", e)

    return jsonify({"ok": True, "candidate": cand or "", "ts": ts or ""})


# ---- Quick diag for your account (legacy; no secrets) ----
@app.get("/__me/diag-legacy")
def me_diag_legacy():
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    try:
        month_cnt = int(get_user_month_usage(uid)) if uid else 0
    except Exception:
        month_cnt = 0

    try:
        cand, ts = last_event_for_user(uid) if uid else (None, None)
    except Exception:
        cand, ts = None, None

    return jsonify({
        "ok": True,
        "logged_in": bool(uid),
        "user_id": uid or None,
        "username": session.get("user") or None,
        "db_pool": bool(DB_POOL),
        "month_usage": month_cnt,
        "last_event": {"candidate": cand or "", "ts": ts or ""},
    })
    # ---------- Skills API (view/add/remove/toggle) ----------
@app.get("/skills")
def skills_get():
    data = {
        "base": sorted(SKILL_CANON, key=lambda s: s.lower()),
        "custom": sorted(SKILLS_CFG.get("custom", []), key=lambda s: s.lower()),
        "base_disabled": sorted(SKILLS_CFG.get("base_disabled", []), key=lambda s: s.lower()),
        "effective": sorted(_effective_skills(), key=lambda s: s.lower()),
    }
    resp = jsonify(data)
    resp.headers["Cache-Control"] = "no-store"
    return resp

@app.post("/skills/custom/add")
def skills_custom_add():
    # accept form or JSON
    skill = (request.form.get("skill") or (request.json.get("skill") if request.is_json else "")).strip()
    if not skill or len(skill) > 60 or re.search(r"[<>]", skill):
        abort(400, "Invalid skill")
    custom = SKILLS_CFG.setdefault("custom", [])
    if skill.lower() not in {x.lower() for x in custom} and skill.lower() not in {x.lower() for x in SKILL_CANON}:
        custom.append(skill)
        custom.sort(key=lambda s: s.lower())
        _save_skills_config()
    return jsonify({
        "ok": True,
        "custom": sorted(SKILLS_CFG["custom"], key=lambda s: s.lower()),
        "effective": sorted(_effective_skills(), key=lambda s: s.lower())
    })

@app.post("/skills/custom/remove")
def skills_custom_remove():
    skill = (request.form.get("skill") or (request.json.get("skill") if request.is_json else "")).strip()
    custom = SKILLS_CFG.setdefault("custom", [])
    SKILLS_CFG["custom"] = [x for x in custom if x.lower() != skill.lower()]
    _save_skills_config()
    return jsonify({
        "ok": True,
        "custom": sorted(SKILLS_CFG["custom"], key=lambda s: s.lower()),
        "effective": sorted(_effective_skills(), key=lambda s: s.lower())
    })

@app.post("/skills/base/toggle")
def skills_base_toggle():
    skill = (request.form.get("skill") or (request.json.get("skill") if request.is_json else "")).strip()
    action = (request.form.get("action") or (request.json.get("action") if request.is_json else "")).strip().lower()
    if skill not in SKILL_CANON:
        abort(400, "Unknown built-in skill")
    disabled = {*(SKILLS_CFG.setdefault("base_disabled", []))}
    if action == "disable":
        disabled.add(skill)
    elif action == "enable":
        disabled.discard(skill)
    else:
        abort(400, "Bad action")
    SKILLS_CFG["base_disabled"] = sorted(disabled, key=lambda s: s.lower())
    _save_skills_config()
    return jsonify({
        "ok": True,
        "base_disabled": sorted(SKILLS_CFG["base_disabled"], key=lambda s: s.lower()),
        "effective": sorted(_effective_skills(), key=lambda s: s.lower())
    })
# ---------- Me (per-user) endpoints — conflict-free versions ----------
@app.get("/x/me-usage")
def me_usage_x():
    """Polishes this month for the logged-in user (DB if available, else 0)."""
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    count = 0
    try:
        if DB_POOL and uid:
            count = int(count_usage_month_db(uid))
    except Exception as e:
        print("me_usage_x error:", e)

    return jsonify({"ok": True, "user_id": (uid or None), "month_usage": count})


@app.get("/x/me-last-event")
def me_last_event_x():
    """Last candidate + timestamp for the logged-in user (DB preferred, legacy fallback)."""
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    cand = ""
    ts = ""
    try:
        if DB_POOL and uid:
            c, t = last_event_for_user(uid)
            cand = c or ""
            ts = t or ""
        if not cand:
            cand = STATS.get("last_candidate", "") or ""
        if not ts:
            ts = STATS.get("last_time", "") or ""
    except Exception as e:
        print("me_last_event_x error:", e)

    return jsonify({"ok": True, "candidate": cand, "ts": ts})


@app.get("/x/me-history")
def me_history_x():
    """
    Recent usage rows for this user.
    Returns: {"ok": True, "history": [{"ts": "...", "candidate": "...", "filename": "..."}]}
    """
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    out = []

    # Prefer Postgres
    if DB_POOL and uid:
        conn = db_conn()
        if conn:
            try:
                with conn:
                    with conn.cursor() as cur:
                        cur.execute("""
                            SELECT to_char(e.ts, 'YYYY-MM-DD HH24:MI:SS') AS ts,
                                   COALESCE(e.candidate, '') AS candidate,
                                   COALESCE(e.filename, '')  AS filename
                              FROM usage_events e
                             WHERE e.user_id = %s
                             ORDER BY e.ts DESC
                             LIMIT 100
                        """, (uid,))
                        for ts, cand, fn in cur.fetchall():
                            out.append({"ts": ts, "candidate": cand, "filename": fn})
            except Exception as e:
                print("me_history_x DB error:", e)
            finally:
                try:
                    db_put(conn)
                except Exception:
                    pass

    # Fallback to legacy JSON
    if not out:
        for it in (STATS.get("history", []) or [])[-100:][::-1]:
            out.append({
                "ts": it.get("ts", ""),
                "candidate": it.get("candidate", ""),
                "filename": it.get("filename", ""),
            })

    return jsonify({"ok": True, "history": out})

# --- Canonical per-user endpoints expected by the UI ---
@app.get("/me/usage")
def me_usage():
    # Delegate to the existing implementation
    return me_usage_x()

@app.get("/me/history")
def me_history():
    # Delegate to the existing implementation
    return me_history_x()
    
@app.get("/me/credits")
def me_credits():
    # Identify user
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    # Compute 'used' (month usage) for backward compatibility
    used = 0
    try:
        if DB_POOL and uid:
            try:
                used = int(count_usage_month_db(uid))
            except Exception:
                # optional legacy fallback
                try:
                    used = int(get_user_month_usage(uid)) if uid else 0
                except Exception:
                    used = 0
    except Exception:
        used = 0

    # Org-aware balance & caps
    org = _user_org_id(uid)
    if org:
        bal = org_balance(org)
        cap = get_user_monthly_cap(org, uid)
        spent = org_user_spent_this_month(org, uid)
        return jsonify({
            "ok": True,
            "scope": "org",
            "org_id": org,
            "balance": bal,
            # cap info
            "myMonthlyCap": cap,
            "mySpentThisMonth": spent,
            "myRemainingThisMonth": (None if cap is None else max(0, cap - spent)),
            # backward-compat fields
            "user_id": uid,
            "used": used,
            "total": None
        })

    # Fallback for users without an org: show personal balance (as before)
    balance = None
    if DB_POOL and uid:
        try:
            row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
            balance = int(row[0]) if row else 0
        except Exception:
            balance = None
    else:
        # legacy session fallback
        try:
            tmp = session.get("trial_credits")
            balance = int(tmp) if tmp is not None else None
        except Exception:
            balance = None

    return jsonify({
        "ok": True,
        "scope": "user",
        "user_id": uid,
        "used": used,
        "balance": balance,
        "total": None
    })
    
# --- Admin utility: ensure the usage_events table exists (safe to run anytime) ---
@app.get("/__admin/ensure-usage-events")
def ensure_usage_events():
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403
        
    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    sql = """
    CREATE TABLE IF NOT EXISTS usage_events (
        id SERIAL PRIMARY KEY,
        user_id INTEGER,
        ts TIMESTAMPTZ DEFAULT now(),
        candidate TEXT,
        filename TEXT
    )
    """
    conn = None
    try:
        conn = DB_POOL.getconn()
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql)
        return jsonify({"ok": True, "created_or_exists": True})
    except Exception as e:
        # Return the error message for quick diagnosis
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass
# --- Admin utility: ensure the credits_ledger table exists ---
@app.get("/__admin/ensure-credits-ledger")
def ensure_credits_ledger():
    # Access guard: only admin/director
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    sql = """
    CREATE TABLE IF NOT EXISTS credits_ledger (
        id SERIAL PRIMARY KEY,
        user_id INTEGER NOT NULL,
        delta INTEGER NOT NULL,
        reason TEXT,
        ext_ref TEXT,
        ts TIMESTAMPTZ DEFAULT now()
    )
    """
    conn = None
    try:
        conn = DB_POOL.getconn()
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql)
        return jsonify({"ok": True, "created_or_exists": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass


# --- Admin utility: grant credits to a user (positive delta) ---
@app.get("/__admin/grant-credits")
def admin_grant_credits():
    # Access guard: only admin/director
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    # Params: user_id (int), delta (int>0), reason (optional), ext_ref (optional)
    try:
        uid = int(request.args.get("user_id") or "0")
        delta = int(request.args.get("delta") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id or delta"}), 400

    if uid <= 0 or delta <= 0:
        return jsonify({"ok": False, "error": "user_id>0 and delta>0 required"}), 400

    reason = (request.args.get("reason") or "grant").strip()
    ext_ref = (request.args.get("ext_ref") or "").strip()

    sql = "INSERT INTO credits_ledger (user_id, delta, reason, ext_ref) VALUES (%s,%s,%s,%s)"
    conn = None
    try:
        conn = DB_POOL.getconn()
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, (uid, delta, reason, ext_ref))
        return jsonify({"ok": True, "granted": {"user_id": uid, "delta": delta, "reason": reason, "ext_ref": ext_ref}})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass


# --- Admin utility: quick check of a user's ledger + balance ---
@app.get("/__admin/credits-summary")
def admin_credits_summary():
    # Access guard: only admin/director
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    try:
        uid = int(request.args.get("user_id") or "0")
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    rows = db_query_all(
        "SELECT id, delta, reason, ext_ref, ts FROM credits_ledger WHERE user_id=%s ORDER BY ts DESC LIMIT 200",
        (uid,)
    )
    balance_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
    balance = int(balance_row[0]) if balance_row else 0

    out = [{"id": r[0], "delta": int(r[1]), "reason": r[2] or "", "ext_ref": r[3] or "", "ts": (r[4].isoformat() if r[4] else None)} for r in rows]
    return jsonify({"ok": True, "user_id": uid, "balance": balance, "rows": out})            
# --- Admin utility: insert a mock usage event for the current user (for testing only) ---
@app.get("/__admin/mock-usage")
def admin_mock_usage():
    """
    Inserts a single usage_events row for the currently logged-in user.

    Query params (optional):
      - candidate: defaults to 'Mock Candidate'
      - filename:  defaults to 'mock.docx'

    Example:
      /__admin/mock-usage?candidate=John%20Doe&filename=demo.docx
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    if not uid:
        return jsonify({"ok": False, "error": "No user_id in session (log in first)"}), 400

    candidate = request.args.get("candidate", "Mock Candidate")
    filename  = request.args.get("filename", "mock.docx")

    sql = "INSERT INTO usage_events (user_id, candidate, filename) VALUES (%s, %s, %s)"
    conn = None
    try:
        conn = DB_POOL.getconn()
        with conn:
            with conn.cursor() as cur:
                cur.execute(sql, (uid, candidate, filename))
        return jsonify({"ok": True, "inserted": {"user_id": uid, "candidate": candidate, "filename": filename}})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass

# --- Admin utility: set a user's credits balance to an exact value ---
@app.get("/__admin/set-credits")
def admin_set_credits():
    # guard: only admin/director
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    # Params
    try:
        uid = int(request.args.get("user_id") or "0")
        target = int(request.args.get("balance") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id or balance"}), 400
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    reason = (request.args.get("reason") or "adjust").strip()

    # Compute delta = target - current
    cur_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
    current = int(cur_row[0]) if cur_row else 0
    diff = target - current
    if diff == 0:
        return jsonify({"ok": True, "user_id": uid, "balance": current, "changed": False})

    # Apply adjustment
    try:
        ok = db_execute(
            "INSERT INTO credits_ledger (user_id, delta, reason, ext_ref) VALUES (%s,%s,%s,%s)",
            (uid, diff, reason, "set-credits"),
        )
        if not ok:
            return jsonify({"ok": False, "error": "insert failed"}), 500
        # new balance
        new_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
        new_bal = int(new_row[0]) if new_row else current + diff
        return jsonify({"ok": True, "user_id": uid, "old_balance": current, "new_balance": new_bal, "delta": diff})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: org pool credits summary ---
# --- Admin: org pool credits summary / grant / set (single canonical block) ---

@app.get("/__admin/org/credits-summary")
def admin_org_credits_summary():
    # admin only
    if not (session.get("is_admin")
            or (session.get("username", "").lower() == "admin")
            or (session.get("user", "").lower() == "admin")):
        return jsonify({"ok": False, "error": "forbidden"}), 403
    if not DB_POOL:
        return jsonify({"ok": False, "error": "db_unavailable"}), 500

    try:
        org_id = int(request.args.get("org_id") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad_org_id"}), 400
    if org_id <= 0:
        return jsonify({"ok": False, "error": "org_id required"}), 400

    # balance
    bal_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    balance = int(bal_row[0]) if bal_row else 0

    # rows (avoid columns that might not exist on older schemas)
    rows = db_query_all(
        "SELECT id, delta, reason, created_at "
        "FROM org_credits_ledger WHERE org_id=%s "
        "ORDER BY id DESC LIMIT 200",
        (org_id,),
    ) or []

    out_rows = []
    for r in rows:
        rid, dlt, rsn, ts = r[0], int(r[1] or 0), (r[2] or ""), r[3]
        ts_txt = ts.isoformat(sep=" ", timespec="seconds") if hasattr(ts, "isoformat") else str(ts)
        out_rows.append({"id": rid, "delta": dlt, "reason": rsn, "ts": ts_txt})

    return jsonify({"ok": True, "org_id": org_id, "balance": balance, "rows": out_rows})


@app.get("/__admin/org/grant-credits")
def admin_org_grant_credits():
    # admin only
    if not (session.get("is_admin")
            or (session.get("username", "").lower() == "admin")
            or (session.get("user", "").lower() == "admin")):
        return jsonify({"ok": False, "error": "forbidden"}), 403
    if not DB_POOL:
        return jsonify({"ok": False, "error": "db_unavailable"}), 500

    try:
        org_id = int(request.args.get("org_id") or "0")
        delta = int(request.args.get("delta") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad_org_id/delta"}), 400
    reason = (request.args.get("reason") or "admin_grant").strip()
    if org_id <= 0 or delta == 0:
        return jsonify({"ok": False, "error": "org_id and non-zero delta required"}), 400

    ok = db_execute(
        "INSERT INTO org_credits_ledger (org_id, delta, reason) VALUES (%s,%s,%s)",
        (org_id, delta, reason),
    )
    if not ok:
        return jsonify({"ok": False, "error": "insert_failed"}), 500

    bal_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    new_bal = int(bal_row[0]) if bal_row else 0
    return jsonify({"ok": True, "org_id": org_id, "delta": delta, "new_balance": new_bal, "reason": reason})


@app.get("/__admin/org/set-credits")
def admin_org_set_credits():
    # admin only
    if not (session.get("is_admin")
            or (session.get("username", "").lower() == "admin")
            or (session.get("user", "").lower() == "admin")):
        return jsonify({"ok": False, "error": "forbidden"}), 403
    if not DB_POOL:
        return jsonify({"ok": False, "error": "db_unavailable"}), 500

    try:
        org_id = int(request.args.get("org_id") or "0")
        target = int(request.args.get("balance") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad_org_id/balance"}), 400
    if org_id <= 0:
        return jsonify({"ok": False, "error": "org_id required"}), 400

    bal_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    cur = int(bal_row[0]) if bal_row else 0
    delta = target - cur
    if delta == 0:
        return jsonify({"ok": True, "org_id": org_id, "balance": cur, "note": "no_change"})

    ok = db_execute(
        "INSERT INTO org_credits_ledger (org_id, delta, reason) VALUES (%s,%s,%s)",
        (org_id, delta, "admin_set_balance"),
    )
    if not ok:
        return jsonify({"ok": False, "error": "insert_failed"}), 500

    new_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    new_bal = int(new_row[0]) if new_row else cur + delta
    return jsonify({"ok": True, "org_id": org_id, "new_balance": new_bal})
# --- Admin utility: enable/disable a user (protect 'admin') ---
@app.get("/__admin/set-user-active")
def admin_set_user_active():
    """
    Usage (as admin/director):
      /__admin/set-user-active?user_id=N&active=1   -> enable
      /__admin/set-user-active?user_id=N&active=0   -> disable

    Hard rule: the 'admin' account cannot be enabled/disabled via this route.
    """
    # guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    # params
    try:
        uid = int(request.args.get("user_id") or "0")
        active_raw = request.args.get("active")
        if active_raw is None:
            return jsonify({"ok": False, "error": "missing active (0|1)"}), 400
        active_val = 1 if str(active_raw) in ("1", "true", "True") else 0
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id/active"}), 400
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    # fetch target
    try:
        row = db_query_one("SELECT username FROM users WHERE id=%s", (uid,))
        if not row:
            return jsonify({"ok": False, "error": "user not found"}), 404
        target_username = (row[0] or "").strip().lower()
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

    # hard protect 'admin' (cannot be toggled)
    if target_username == "admin":
        return jsonify({
            "ok": False,
            "error": "cannot modify 'admin' via this route",
            "forbidden_admin_target": True
        }), 403

    # apply
    try:
        ok = db_execute("UPDATE users SET active=%s WHERE id=%s", (active_val, uid))
        if not ok:
            return jsonify({"ok": False, "error": "update failed"}), 500
        return jsonify({"ok": True, "user_id": uid, "username": target_username, "active": bool(active_val)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500    

# --- Admin: ONE-TIME migration to enable org-shared credits ---
@app.get("/__admin/migrate_org_pool")
def admin_migrate_org_pool():
    # admin only
    if not (session.get("is_admin") or session.get("username","").lower()=="admin"):
        return jsonify({"ok": False, "error": "forbidden"}), 403

    stmts = [
        """
        CREATE TABLE IF NOT EXISTS org_credits_ledger (
          id SERIAL PRIMARY KEY,
          org_id INTEGER NOT NULL,
          delta INTEGER NOT NULL,
          reason TEXT,
          user_id INTEGER,
          created_by INTEGER,
          created_at TIMESTAMPTZ DEFAULT NOW()
        )
        """,
        "CREATE INDEX IF NOT EXISTS idx_org_credits_ledger_org ON org_credits_ledger(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_org_credits_ledger_org_user_month ON org_credits_ledger(org_id, user_id, created_at)",
        """
        CREATE TABLE IF NOT EXISTS org_user_limits (
          org_id INTEGER NOT NULL,
          user_id INTEGER NOT NULL,
          monthly_cap INTEGER,
          active BOOLEAN DEFAULT TRUE,
          PRIMARY KEY (org_id, user_id)
        )
        """
    ]
    for s in stmts:
        ok = db_execute(s, tuple())
        if not ok:
            return jsonify({"ok": False, "error": "migration_failed"}), 500
    return jsonify({"ok": True, "migrated": True})
# --- Admin utility: ensure the orgs schema exists (safe to run anytime) ---
@app.get("/__admin/ensure-orgs-schema")
def ensure_orgs_schema():
    """
    Creates the minimal organisation layer:
      - orgs table (id, name UNIQUE NOT NULL)
      - users.org_id column (if missing)
      - usage_events.org_id column (if missing)
      - helpful indexes

    This does NOT assign users to orgs yet (that’s the next steps).
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    sql_statements = [
        # Orgs table
        """
        CREATE TABLE IF NOT EXISTS orgs (
            id   SERIAL PRIMARY KEY,
            name TEXT UNIQUE NOT NULL
        );
        """,
        # Users -> org_id
        "ALTER TABLE IF EXISTS users ADD COLUMN IF NOT EXISTS org_id INTEGER;",
        # usage_events -> org_id
        "ALTER TABLE IF EXISTS usage_events ADD COLUMN IF NOT EXISTS org_id INTEGER;",
        # Indexes (no FKs yet to avoid locking surprises)
        "CREATE INDEX IF NOT EXISTS idx_users_org_id ON users(org_id);",
        "CREATE INDEX IF NOT EXISTS idx_usage_events_org_id ON usage_events(org_id);"
    ]

    conn = None
    try:
        conn = DB_POOL.getconn()
        with conn:
            with conn.cursor() as cur:
                for stmt in sql_statements:
                    cur.execute(stmt)
        return jsonify({"ok": True, "created_or_exists": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass

# --- Admin: ensure org template columns (idempotent) ---
@app.get("/__admin/ensure-template-schema")
def __admin_ensure_template_schema():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403
    try:
        db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS template_path TEXT")
        db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS template_updated_at TIMESTAMPTZ")
        return jsonify({
            "ok": True,
            "orgs_template_path": True,
            "orgs_template_updated_at": True
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500  

# --- Admin: ensure per-org profile column (idempotent) ---
@app.get("/__admin/ensure-org-profile")
def __admin_ensure_org_profile():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403
    try:
        # Prefer JSONB; if not available, fall back to TEXT
        try:
            db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS profile_json JSONB")
        except Exception:
            db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS profile_json TEXT")
        return jsonify({"ok": True, "applied": {"orgs.profile_json": True}})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: upload a DOCX template for an org (GET=form, POST=upload) ---
@app.route("/__admin/upload-org-template", methods=["GET", "POST"])
def __admin_upload_org_template():
    # admin guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin_flag = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin_flag = False
    if not is_admin_flag:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # GET: tiny HTML form
    if request.method == "GET":
        opts = []
        try:
            rows = db_query_all("SELECT id, COALESCE(name,'') FROM orgs ORDER BY id") or []
            for oid, oname in rows:
                opts.append(f'<option value="{int(oid)}">{int(oid)} — {oname or "org "+str(int(oid))}</option>')
        except Exception:
            pass
        org_select = (
            f'<select name="org_id" required>{"".join(opts)}</select>'
            if opts else
            '<input type="number" name="org_id" placeholder="Org ID" required min="1" />'
        )
        html = f"""<!doctype html><html><head><meta charset="utf-8"><title>Upload Org Template</title>
<style>body{{font:14px/1.4 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px}}
form{{display:grid;gap:10px;max-width:520px}}
input,select,button{{padding:8px;border:1px solid #e5e7eb;border-radius:8px}}
.btn{{display:inline-block;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;background:#fff;text-decoration:none;color:#0f172a}}
</style></head><body>
  <h1>Upload DOCX template (per org)</h1>
  <p>Choose an organisation and upload a <strong>.docx</strong> file.</p>
  <form method="POST" enctype="multipart/form-data">
    <label>Organisation: {org_select}</label>
    <label>Template (.docx): <input type="file" name="file" accept=".docx" required></label>
    <button type="submit">Upload</button>
  </form>
  <p style="margin-top:14px">
    <a class="btn" href="/owner/console">Owner</a>
    <a class="btn" href="/app">App</a>
  </p>
</body></html>"""
        return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

    # --- Owner: credits audit (admin-only, read-only) ---
@app.get("/owner/api/credits-ledger")
def owner_api_credits_ledger():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # params
    try:
        org_id = int(request.args.get("org_id", "0"))
    except Exception:
        org_id = 0
    limit = request.args.get("limit", "200").strip()
    try:
        limit = max(1, min(1000, int(limit)))
    except Exception:
        limit = 200

    # org filter
    where = ""
    args = []
    if org_id > 0:
        where = "WHERE ocl.org_id = %s"
        args.append(org_id)

    # Prefer org_credits_ledger (top-ups)
    rows = db_query_all(f"""
        SELECT ocl.id, ocl.org_id, o.name AS org_name, ocl.delta, ocl.reason,
               COALESCE(ocl.created_at, ocl.ts) AS ts
          FROM org_credits_ledger ocl
          JOIN orgs o ON o.id = ocl.org_id
          {where}
         ORDER BY ts DESC
         LIMIT %s
    """, (*args, limit)) or []

    out = [
        {
            "id": r[0],
            "org_id": r[1],
            "org_name": r[2],
            "delta": int(r[3] or 0),
            "reason": (r[4] or "").strip(),
            "ts": (r[5].isoformat() if hasattr(r[5], "isoformat") else str(r[5])),
        }
        for r in rows
    ]

    return jsonify({"ok": True, "items": out})

    # POST: handle upload
    file = request.files.get("file")
    org_id_raw = request.form.get("org_id") or request.args.get("org_id")
    try:
        org_id = int(org_id_raw or "0")
    except Exception:
        org_id = 0
    if not (file and org_id):
        return jsonify({"ok": False, "error": "missing file or org_id"}), 400
    if not db_query_one("SELECT 1 FROM orgs WHERE id=%s", (org_id,)):
        return jsonify({"ok": False, "error": "org not found"}), 404

    filename = getattr(file, "filename", "") or ""
    if not filename.lower().endswith(".docx"):
        return jsonify({"ok": False, "error": "must be a .docx file"}), 400

    from werkzeug.utils import secure_filename
    import os, time
    base_dir = "/mnt/data/org_templates"
    os.makedirs(base_dir, exist_ok=True)
    org_dir = os.path.join(base_dir, str(org_id))
    os.makedirs(org_dir, exist_ok=True)
    ts = int(time.time())
    safe = secure_filename(filename) or f"template_{ts}.docx"
    canonical_path = os.path.join(org_dir, "template.docx")
    file.save(canonical_path)
    db_execute(
        "UPDATE orgs SET template_path=%s, template_updated_at=NOW() WHERE id=%s",
        (canonical_path, org_id),
    )
    size_bytes = os.path.getsize(canonical_path) if os.path.exists(canonical_path) else None
    return jsonify({"ok": True, "org_id": org_id, "template_path": canonical_path, "size": size_bytes})

# --- Admin: simple form to create a user and assign to an org (GET -> calls /__admin/create-user) ---
@app.get("/__admin/new-user")
def __admin_new_user():
    # admin guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin_flag = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin_flag = False
    if not is_admin_flag:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # orgs for dropdown
    opts = []
    try:
        rows = db_query_all("SELECT id, COALESCE(name,'') FROM orgs ORDER BY id") or []
        for oid, oname in rows:
            opts.append(f'<option value="{int(oid)}">{int(oid)} — {oname or "org "+str(int(oid))}</option>')
    except Exception:
        pass
    org_select = (
        f'<select name="org_id" required>{"".join(opts)}</select>'
        if opts else
        '<input type="number" name="org_id" placeholder="Org ID" required min="1" />'
    )

    # tiny form that submits to /__admin/create-user (GET)
    html = f"""
<!doctype html>
<html><head><meta charset="utf-8"><title>Create user</title>
<style>
  body{{font:14px/1.4 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px}}
  form{{display:grid;gap:10px;max-width:520px}}
  input,select,button{{padding:8px;border:1px solid #e5e7eb;border-radius:8px}}
  .btn{{display:inline-block;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;background:#fff;text-decoration:none;color:#0f172a}}
  .row{{display:flex;gap:8px;align-items:center}}
</style></head>
<body>
  <h1>Create user</h1>
  <p>Fill the fields and submit. The form calls <code>/__admin/create-user</code> and shows its JSON.</p>

  <form method="GET" action="/__admin/create-user" target="_blank">
    <label>Username <input type="text" name="u" placeholder="e.g. acme1" required></label>
    <label>Password <input type="text" name="p" placeholder="Temp1234!" required></label>
    <label>Email (optional) <input type="email" name="email" placeholder="user@example.com"></label>
    <label>Organisation {org_select}</label>
    <div class="row">
      <button type="submit">Create user</button>
      <a class="btn" href="/owner/console">Owner</a>
      <a class="btn" href="/app">App</a>
    </div>
  </form>
</body></html>
"""
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

# --- Admin: reset a user's password (GET=form, POST=apply) ---
@app.route("/__admin/reset-password", methods=["GET", "POST"])
def __admin_reset_password():
    # admin guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin_flag = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin_flag = False
    if not is_admin_flag:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if request.method == "GET":
        html = """<!doctype html><html><head><meta charset="utf-8"><title>Reset user password</title>
<style>body{font:14px/1.4 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px}
form{display:grid;gap:10px;max-width:520px}
input,button{padding:8px;border:1px solid #e5e7eb;border-radius:8px}
.hint{color:#64748b;font-size:12px}
.btn{display:inline-block;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;background:#fff;text-decoration:none;color:#0f172a}
.row{display:flex;gap:8px;align-items:center}
</style></head><body>
  <h1>Reset user password (admin)</h1>
  <p class="hint">Enter either a <strong>User ID</strong> <em>or</em> a <strong>Username</strong>, plus the new password.</p>
  <form method="POST" action="/__admin/reset-password" target="_blank">
    <label>User ID (number) <input type="number" name="user_id" min="1" placeholder="e.g. 12"></label>
    <label>Username (text) <input type="text" name="username" placeholder="e.g. hamilton"></label>
    <label>New password <input type="text" name="new_password" required placeholder="Temp1234!"></label>
    <div class="row">
      <button type="submit">Reset password</button>
      <a class="btn" href="/owner/console">Owner</a>
      <a class="btn" href="/app">App</a>
    </div>
  </form>
</body></html>"""
        return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

    # POST: apply reset
    uid_raw = (request.form.get("user_id") or "").strip()
    uname_raw = (request.form.get("username") or "").strip()
    new_pw = (request.form.get("new_password") or "").strip()

    try:
        uid = int(uid_raw) if uid_raw else 0
    except Exception:
        uid = 0

    if not new_pw:
        return jsonify({"ok": False, "error": "new_password required"}), 400

    # resolve user by id or username
    row = None
    if uid > 0:
        row = db_query_one("SELECT id, username FROM users WHERE id=%s", (uid,))
    elif uname_raw:
        row = db_query_one("SELECT id, username FROM users WHERE LOWER(username)=LOWER(%s)", (uname_raw,))
    else:
        return jsonify({"ok": False, "error": "user_id or username required"}), 400

    if not row:
        return jsonify({"ok": False, "error": "user_not_found"}), 404

    target_id = int(row[0])
    target_username = (row[1] or "").strip().lower()
    if target_username == "admin":
        return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403

    try:
        hashed = generate_password_hash(new_pw)
        ok = db_execute("UPDATE users SET password_hash=%s WHERE id=%s", (hashed, target_id))
        if not ok:
            return jsonify({"ok": False, "error": "update_failed"}), 500
        return jsonify({"ok": True, "user_id": target_id, "username": target_username})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: simple form to create a new organisation (GET -> calls /__admin/create-org) ---
@app.get("/__admin/new-org")
def __admin_new_org():
    # admin guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin_flag = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin_flag = False
    if not is_admin_flag:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # tiny form that submits to /__admin/create-org (GET)
    html = """
<!doctype html>
<html><head><meta charset="utf-8"><title>Create organisation</title>
<style>
  body{font:14px/1.4 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px}
  form{display:grid;gap:10px;max-width:520px}
  input,button{padding:8px;border:1px solid #e5e7eb;border-radius:8px}
  .btn{display:inline-block;padding:8px 10px;border:1px solid #e5e7eb;border-radius:8px;background:#fff;text-decoration:none;color:#0f172a}
  .row{display:flex;gap:8px;align-items:center}
</style></head>
<body>
  <h1>Create organisation</h1>
  <p>Fill the name and submit. The form calls <code>/__admin/create-org</code> and shows its JSON.</p>

  <form method="GET" action="/__admin/create-org" target="_blank">
    <label>Name <input type="text" name="name" placeholder="e.g. Acme" required></label>
    <div class="row">
      <button type="submit">Create org</button>
      <a class="btn" href="/owner/console">Owner</a>
      <a class="btn" href="/app">App</a>
    </div>
  </form>
</body></html>
"""
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

# --- Admin: edit per-org profile (JSON) for structure/labels/tone (GET=form, POST=save) ---
@app.route("/__admin/org-profile", methods=["GET", "POST"])
def __admin_org_profile():
    # admin guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin_flag = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin_flag = False
    if not is_admin_flag:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # load orgs for dropdown
    opts = []
    try:
        rows = db_query_all("SELECT id, COALESCE(name,'') FROM orgs ORDER BY id") or []
        for oid, oname in rows:
            opts.append(f'<option value="{int(oid)}">{int(oid)} — {oname or ("org "+str(int(oid)))}</option>')
    except Exception:
        pass

    import json
    # org_id from query or form
    org_id_raw = request.values.get("org_id") or ""
    try:
        org_id = int(org_id_raw) if org_id_raw else 0
    except Exception:
        org_id = 0

    if request.method == "POST":
        if not org_id:
            return jsonify({"ok": False, "error": "missing org_id"}), 400
        profile_text = request.form.get("profile") or ""
        try:
            # validate JSON, then store canonical JSON string
            obj = json.loads(profile_text)
            canon = json.dumps(obj, ensure_ascii=False)
            db_execute("UPDATE orgs SET profile_json=%s WHERE id=%s", (canon, org_id))
            return jsonify({"ok": True, "org_id": org_id, "saved": True, "bytes": len(canon)})
        except Exception as e:
            return jsonify({"ok": False, "error": f"invalid json or save failed: {e}"}), 400

    # GET: render small editor
    selected = ""
    if org_id and opts:
        # re-render options with selected
        new_opts = []
        for o in opts:
            val = o.split('value="',1)[1].split('"',1)[0] if 'value="' in o else ""
            if val and int(val) == org_id:
                new_opts.append(o.replace('value="'+val+'"', 'value="'+val+'" selected'))
            else:
                new_opts.append(o)
        opts = new_opts

    # fetch current profile (if org selected)
    current_json = ""
    if org_id:
        try:
            row = db_query_one("SELECT profile_json FROM orgs WHERE id=%s", (org_id,))
            if row and row[0]:
                # row may be text or json; stringify nicely
                if isinstance(row[0], (dict, list)):
                    current_json = json.dumps(row[0], indent=2, ensure_ascii=False)
                else:
                    # already a string in DB
                    current_json = str(row[0])
        except Exception:
            pass

    # simple example profile (you can tweak later per client)
    example_profile = json.dumps({
        "sections_order": [
            "name_and_contact",
            "executive_summary",
            "skills",
            "experience",
            "education",
            "certifications"
        ],
        "labels": {
            "executive_summary": "Executive Summary",
            "skills": "Key Skills",
            "experience": "Professional Experience",
            "education": "Education"
        },
        "content": {
            "summary_tone": "concise, outcome-focused",
            "bullet_style": "impact-first",
            "date_format": "MMM yyyy",
            "experience_bullets_max": 6
        }
    }, indent=2)

    org_select = (
        f'<select name="org_id" onchange="location.search=`?org_id=`+this.value" required>{"".join(opts)}</select>'
        if opts else
        '<input type="number" name="org_id" placeholder="Org ID" required min="1" />'
    )

    # build HTML
    html = f"""
<!doctype html>
<html><head><meta charset="utf-8"><title>Org profile (JSON)</title>
<style>
  body{{font:14px/1.45 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px}}
  textarea,input,select,button{{padding:8px;border:1px solid #e5e7eb;border-radius:8px;font-family:ui-monospace,Menlo,Consolas,monospace}}
  .row{{display:flex;gap:8px;align-items:center;flex-wrap:wrap}}
  .cols{{display:grid;grid-template-columns:1fr 1fr;gap:16px}}
  .card{{border:1px solid #e5e7eb;border-radius:10px;padding:14px}}
  .btn{{display:inline-block;padding:8px 12px;border:1px solid #e5e7eb;border-radius:8px;background:#fff;text-decoration:none;color:#0f172a}}
  .muted{{color:#64748b}}
</style></head>
<body>
  <h1>Organisation profile (JSON)</h1>
  <div class="row" style="margin-bottom:12px">
    <label>Organisation: {org_select}</label>
    <a class="btn" href="/owner/console">Owner</a>
    <a class="btn" href="/app">App</a>
  </div>

  <div class="cols">
    <div class="card">
      <h3 style="margin-top:0">Current profile</h3>
      <form method="POST">
        <input type="hidden" name="org_id" value="{org_id or ''}">
        <textarea name="profile" rows="22" style="width:100%" placeholder='{{}}'>{json.dumps(current_json, ensure_ascii=False)[1:-1]}</textarea>
        <div class="row" style="margin-top:10px">
          <button type="submit">Save</button>
          <button type="button" onclick="document.querySelector('textarea[name=profile]').value={json.dumps(example_profile)}">Load example</button>
        </div>
      </form>
      <p class="muted" style="margin-top:10px">Tip: leave empty to use the default Hamilton structure. Saving any JSON here will override structure/labels for this org when we wire it in.</p>
    </div>
    <div class="card">
      <h3 style="margin-top:0">Example (starter)</h3>
      <pre style="white-space:pre-wrap">{example_profile}</pre>
    </div>
  </div>
</body></html>
"""
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

            # --- Helper: org of the current session user (or None) ---
def _current_user_org_id():
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if not (DB_POOL and uid):
        return None
    try:
        row = db_query_one("SELECT org_id FROM users WHERE id=%s", (uid,))
        if row and row[0]:
            return int(row[0])
    except Exception as e:
        print("org lookup failed:", e)
    return None

def _month_bounds_utc():
    now = datetime.utcnow()
    start = datetime(now.year, now.month, 1)
    if now.month == 12:
        next_start = datetime(now.year + 1, 1, 1)
    else:
        next_start = datetime(now.year, now.month + 1, 1)
    return start, next_start

def org_balance(org_id: int) -> int:
    row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    return int(row[0]) if row else 0

def org_user_spent_this_month(org_id: int, user_id: int) -> int:
    start, next_start = _month_bounds_utc()
    row = db_query_one("""
        SELECT COALESCE(-SUM(delta),0) FROM org_credits_ledger
        WHERE org_id=%s AND user_id=%s AND delta < 0
          AND created_at >= %s AND created_at < %s
    """, (org_id, user_id, start, next_start))
    return int(row[0]) if row else 0

def get_user_monthly_cap(org_id: int, user_id: int):
    row = db_query_one("""
        SELECT COALESCE(monthly_cap, month_cap)
        FROM org_user_limits
        WHERE org_id=%s AND user_id=%s AND active
        LIMIT 1
    """, (org_id, user_id))
    if not row:
        return None
    return None if row[0] is None else int(row[0])

def _user_org_id(user_id: int):
    row = db_query_one("SELECT org_id FROM users WHERE id=%s", (user_id,))
    return int(row[0]) if row and row[0] is not None else None

def charge_credit_for_polish(user_id: int, cost: int = 1, candidate: str = "", filename: str = ""):
    """
    Returns (ok: bool, err: Optional[str])
      err in {"insufficient_org_credits","user_monthly_cap_reached","insufficient_user_credits","charge_failed"}
    """
    try:
        me_is_admin = bool(session.get("is_admin")) or (session.get("username","").strip().lower() == "admin")
    except Exception:
        me_is_admin = False
    if me_is_admin:
        return True, None

    org_id = _user_org_id(user_id)

    if org_id:
        bal = org_balance(org_id)
        if bal < cost:
            return False, "insufficient_org_credits"

        cap = get_user_monthly_cap(org_id, user_id)
        if cap is not None:
            spent = org_user_spent_this_month(org_id, user_id)
            if spent + cost > cap:
                return False, "user_monthly_cap_reached"

        ok = db_execute(
            "INSERT INTO org_credits_ledger (org_id, delta, reason, user_id, created_by) VALUES (%s,%s,%s,%s,%s)",
            (org_id, -cost, f"polish:{candidate}:{filename}", user_id, user_id)
        )
        return (True, None) if ok else (False, "charge_failed")

    # fallback: personal ledger
    row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (user_id,))
    ubal = int(row[0]) if row else 0
    if ubal < cost:
        return False, "insufficient_user_credits"
    ok = db_execute(
        "INSERT INTO credits_ledger (user_id, delta, reason, created_by) VALUES (%s,%s,%s,%s)",
        (user_id, -cost, f"polish:{candidate}:{filename}", user_id)
    )
    return (True, None) if ok else (False, "charge_failed")
# --- Director (org-scoped): one-call dashboard payload for this org ---
@app.get("/director/api/dashboard")
def director_api_dashboard():
    """
    Org-scoped dashboard for the currently logged-in user.
    Returns: {
      ok, source, orgId, orgName,
      pool: { balance },
      month: { total, rows:[{user_id, username, count}] },
      recent: [ { ts, user_id, username, candidate, filename } ]
    }
    """
    # must be logged in
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    if not DB_POOL:
        return jsonify({"ok": True, "source": "legacy", "month": {"total": 0, "rows": []}, "recent": []})

    # org scope
    org_id = _current_user_org_id()
    if not org_id:
        return jsonify({"ok": True, "source": "no_org", "month": {"total": 0, "rows": []}, "recent": []})

    # inputs
    try:
        limit = int(request.args.get("limit", "50"))
    except Exception:
        limit = 50
    limit = max(1, min(limit, 200))

    # org name
    row = db_query_one("SELECT name FROM orgs WHERE id=%s", (org_id,))
    org_name = (row[0] if row and row[0] else None)

    # ORG POOL BALANCE (sum org_credits_ledger.delta for this org)
    bal_row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    pool_balance = int(bal_row[0]) if bal_row else 0

    # This month per-user counts
    per_user = db_query_all("""
        SELECT u.id AS user_id, u.username, COUNT(e.*) AS cnt
        FROM users u
        LEFT JOIN usage_events e
               ON e.user_id = u.id
              AND e.ts >= date_trunc('month', now())
        WHERE u.org_id = %s
        GROUP BY u.id, u.username
        ORDER BY cnt DESC, u.username ASC
    """, (org_id,)) or []

    month_total_row = db_query_one("""
        SELECT COUNT(*) FROM usage_events
        WHERE org_id = %s AND ts >= date_trunc('month', now())
    """, (org_id,))
    month_total = int(month_total_row[0]) if month_total_row else 0

    # Recent org events
    rec = db_query_all("""
        SELECT e.ts, e.user_id, u.username, e.candidate, e.filename
        FROM usage_events e
        LEFT JOIN users u ON u.id = e.user_id
        WHERE e.org_id = %s
        ORDER BY e.ts DESC
        LIMIT %s
    """, (org_id, limit)) or []

    recent = [{
        "ts": (r[0].isoformat(sep=" ", timespec="seconds") if hasattr(r[0], "isoformat") else str(r[0])),
        "user_id": r[1],
        "username": r[2],
        "candidate": r[3],
        "filename": r[4],
    } for r in rec]

    month_rows = [{"user_id": r[0], "username": r[1], "count": int(r[2])} for r in per_user]

    return jsonify({
        "ok": True,
        "source": "db-org",
        "orgId": org_id,
        "orgName": org_name,
        "pool": {"balance": pool_balance},
        "month": {"total": month_total, "rows": month_rows},
        "recent": recent
    })
            # --- Director (org-scoped): list users in my org with balances ---
@app.get("/director/api/users")
def director_api_users():
    """
    Returns users in the same org as the current session user.
    Shape:
      {
        ok: true,
        org_id: <int|null>,
        users: [
          { id, username, active, balance }
        ]
      }
    """
    # must be logged in
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    if not DB_POOL:
        return jsonify({"ok": True, "org_id": None, "users": []})

    # find my org
    org_id = _current_user_org_id()
    if not org_id:
        return jsonify({"ok": True, "org_id": None, "users": []})

    conn = None
    try:
        conn = DB_POOL.getconn()
        users, bal_map = [], {}

        with conn:
            with conn.cursor() as cur:
                # balances for this org
                cur.execute("""
                    SELECT user_id, COALESCE(SUM(delta),0) AS balance
                    FROM credits_ledger
                    WHERE org_id = %s
                    GROUP BY user_id
                """, (org_id,))
                bal_map = {int(r[0]): int(r[1]) for r in cur.fetchall()}

                # users in this org
                cur.execute("""
                    SELECT id, username, COALESCE(active, TRUE) AS active
                    FROM users
                    WHERE org_id = %s
                    ORDER BY username ASC
                """, (org_id,))
                for uid2, uname, act in cur.fetchall():
                    users.append({
                        "id": int(uid2),
                        "username": uname or "",
                        "active": bool(act),
                        "balance": bal_map.get(int(uid2))
                    })

        return jsonify({"ok": True, "org_id": org_id, "users": users})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass

def _require_logged_in():
    try:
        uid = int(session.get("user_id") or 0)
        return uid if uid > 0 else None
    except Exception:
        return None

# --- Director: set per-user monthly cap within my org ---
@app.get("/director/api/user/set-monthly-cap")
def director_set_monthly_cap():
    me_uid = _require_logged_in()
    if not me_uid:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    my_org = _current_user_org_id()
    if not my_org:
        return jsonify({"ok": False, "error": "no_org"}), 400

    # params: user_id (required), cap (int or "null")
    try:
        target_id = int(request.args.get("user_id") or "0")
        cap_raw = request.args.get("cap")  # string: int or "null"
    except Exception:
        return jsonify({"ok": False, "error": "bad_params"}), 400
    if target_id <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    row = db_query_one("SELECT username, org_id FROM users WHERE id=%s", (target_id,))
    if not row:
        return jsonify({"ok": False, "error": "user_not_found"}), 404

    target_username = (row[0] or "").strip().lower()
    if int(row[1] or 0) != my_org:
        return jsonify({"ok": False, "error": "not_in_my_org"}), 403
    if target_username == "admin":
        return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403

    # cap parsing
    cap_val = None
    if cap_raw is not None and str(cap_raw).lower() != "null":
        try:
            cap_val = max(0, int(cap_raw))
        except Exception:
            return jsonify({"ok": False, "error": "bad_cap"}), 400

    # upsert into org_user_limits
    existing = db_query_one("SELECT 1 FROM org_user_limits WHERE org_id=%s AND user_id=%s", (my_org, target_id))
    if existing:
        ok = db_execute(
            "UPDATE org_user_limits SET monthly_cap=%s, active=TRUE WHERE org_id=%s AND user_id=%s",
            (cap_val, my_org, target_id)
        )
    else:
        ok = db_execute(
            "INSERT INTO org_user_limits (org_id, user_id, monthly_cap, active) VALUES (%s,%s,%s,TRUE)",
            (my_org, target_id, cap_val)
        )
    if not ok:
        return jsonify({"ok": False, "error": "update_failed"}), 500

    spent = org_user_spent_this_month(my_org, target_id)
    return jsonify({"ok": True, "user_id": target_id, "monthly_cap": cap_val, "spent_this_month": spent})

# --- Director: enable/disable a user in my org (protect 'admin') ---
@app.get("/director/api/user/set-active")
def director_set_active():
    # must be logged in and be director/admin
    if not (session.get("user_id") and (session.get("director") or is_admin())):
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # parse inputs
    try:
        uid = int(request.args.get("user_id") or "0")
        active_raw = request.args.get("active")
        if active_raw is None:
            return jsonify({"ok": False, "error": "missing active (0|1)"}), 400
        active_val = 1 if str(active_raw).lower() in ("1", "true", "yes") else 0
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id/active"}), 400
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    # verify same org + protect 'admin'
    row = db_query_one("SELECT username, org_id FROM users WHERE id=%s", (uid,))
    if not row:
        return jsonify({"ok": False, "error": "user not found"}), 404
    if (row[0] or "").strip().lower() == "admin":
        return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403
    my_org = _current_user_org_id()
    if my_org and int(row[1] or 0) != my_org and not is_admin():
        return jsonify({"ok": False, "error": "not_in_my_org"}), 403

    ok = db_execute("UPDATE users SET active=%s WHERE id=%s", (active_val, uid))
    if not ok:
        return jsonify({"ok": False, "error": "update_failed"}), 500
    return jsonify({"ok": True, "user_id": uid, "active": bool(active_val)})

# --- Director: delete a user in my org (protect 'admin') ---
@app.get("/director/api/user/delete")
def director_delete_user():
    # must be logged in and be director/admin
    if not (session.get("user_id") and (session.get("director") or is_admin())):
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # read user_id
    try:
        uid = int(request.args.get("user_id") or "0")
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    # verify same org + protect 'admin'
    row = db_query_one("SELECT username, org_id FROM users WHERE id=%s", (uid,))
    if not row:
        return jsonify({"ok": False, "error": "user not found"}), 404

    target_username = (row[0] or "").strip().lower()
    if target_username == "admin":
        return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403

    my_org = _current_user_org_id()
    if my_org and int(row[1] or 0) != my_org and not is_admin():
        return jsonify({"ok": False, "error": "not_in_my_org"}), 403

    # delete (related rows removed via ON DELETE CASCADE if set)
    ok = db_execute("DELETE FROM users WHERE id=%s", (uid,))
    if not ok:
        return jsonify({"ok": False, "error": "delete_failed"}), 500
    return jsonify({"ok": True, "deleted_user_id": uid})

@app.post("/director/api/change-password")
def change_director_password():
    if not session.get("user_id"):
        return jsonify({"ok": False, "error": "not_logged_in"}), 403

    new_pass = (request.json or {}).get("password")
    if not new_pass:
        return jsonify({"ok": False, "error": "missing_password"}), 400

    hashed = generate_password_hash(new_pass)
    ok = db_execute("UPDATE users SET password_hash=%s WHERE id=%s", (hashed, session["user_id"]))
    if not ok:
        return jsonify({"ok": False, "error": "update_failed"}), 500

    return jsonify({"ok": True})

# Aliases to cover legacy front-ends
@app.get("/director/api/activate")
def _alias_activate():
    return director_set_active()

@app.get("/director/api/user/activate")
def _alias_user_activate():
    return director_set_active()

@app.get("/__admin/set-user-active")
def _alias_admin_activate():
    return director_set_active()

# --- Director: reset a user's password (same org) ---
@app.post("/director/api/user/reset-password")
def director_reset_password():
    me_uid = _require_logged_in()
    if not me_uid:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    my_org = _current_user_org_id()
    if not my_org:
        return jsonify({"ok": False, "error": "no_org"}), 400

    # JSON body: { "user_id": N, "new_password": "..." }
    data = request.get_json(silent=True) or {}
    try:
        target_id = int(data.get("user_id") or 0)
    except Exception:
        target_id = 0
    new_pw = (data.get("new_password") or "").strip()

    if target_id <= 0 or not new_pw:
        return jsonify({"ok": False, "error": "user_id and new_password required"}), 400

    row = db_query_one("SELECT username, org_id FROM users WHERE id=%s", (target_id,))
    if not row:
        return jsonify({"ok": False, "error": "user_not_found"}), 404
    target_username = (row[0] or "").strip().lower()
    if int(row[1] or 0) != my_org:
        return jsonify({"ok": False, "error": "not_in_my_org"}), 403
    if target_username == "admin":
        return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403

    try:
        hashed = generate_password_hash(new_pw)
        ok = db_execute("UPDATE users SET password_hash=%s WHERE id=%s", (hashed, target_id))
        if not ok:
            return jsonify({"ok": False, "error": "update_failed"}), 500
        return jsonify({"ok": True, "user_id": target_id, "username": target_username})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
# --- Director: org credits summary (my org) ---
@app.get("/director/api/org/credits-summary")
def director_org_credits_summary():
    me_uid = _require_logged_in()
    if not me_uid:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    my_org = _current_user_org_id()
    if not my_org:
        return jsonify({"ok": False, "error": "no_org"}), 400

    balance = org_balance(my_org)

    rows = db_query_all("""
        SELECT id, delta, reason, user_id, created_by, created_at
        FROM org_credits_ledger
        WHERE org_id=%s
        ORDER BY id DESC
        LIMIT 200
    """, (my_org,))

    # list users in this org with their caps
    caps = db_query_all("""
        SELECT u.id AS user_id, u.username, l.monthly_cap
        FROM users u
        LEFT JOIN org_user_limits l
          ON l.org_id = u.org_id AND l.user_id = u.id AND l.active
        WHERE u.org_id=%s
        ORDER BY u.username ASC
    """, (my_org,))

    return jsonify({"ok": True, "org_id": my_org, "balance": balance, "rows": rows or [], "limits": caps or []})
# --- Director (org-scoped): create a user in my org (optional seed credits) ---
@app.get("/director/api/create-user")
def director_api_create_user():
    """
    Creates a new active user in the current director's org.
    Query params:
      - u: username (required)
      - p: password (required)
      - seed: optional integer to grant starting credits
    Returns: { ok, id, username, seed_granted }
    """
    # must be logged in
    try:
        me_uid = int(session.get("user_id") or 0)
    except Exception:
        me_uid = 0
    if me_uid <= 0:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

    if not DB_POOL:
        return jsonify({"ok": False, "error": "db_unavailable"}), 500

    # find my org
    org_id = _current_user_org_id()
    if not org_id:
        return jsonify({"ok": False, "error": "no_org"}), 400

    # read inputs
    u = (request.args.get("u") or "").strip()
    p = request.args.get("p") or ""
    seed_raw = request.args.get("seed")
    try:
        seed = int(seed_raw) if seed_raw not in (None, "") else 0
    except Exception:
        seed = 0

    if not u or not p:
        return jsonify({"ok": False, "error": "missing u or p"}), 400

    # basic guards
    uname = u.lower()
    if uname in ("admin", "director"):
        return jsonify({"ok": False, "error": "reserved_username"}), 400

    # create
    try:
        # global uniqueness check (safer)
        row = db_query_one("SELECT id FROM users WHERE username=%s", (u,))
        if row:
            return jsonify({"ok": False, "error": "user_exists", "id": int(row[0])}), 409

        pw_hash = generate_password_hash(p)
        ok = db_execute(
            "INSERT INTO users (username, password_hash, active, org_id) VALUES (%s, %s, %s, %s)",
            (u, pw_hash, True, org_id),
        )
        if not ok:
            return jsonify({"ok": False, "error": "insert_failed"}), 500

        row2 = db_query_one("SELECT id FROM users WHERE username=%s", (u,))
        new_id = int(row2[0]) if row2 else None

        # optionally grant seed credits
        granted = 0
        if new_id and seed > 0:
            if credits_add(new_id, seed, reason="seed", ext_ref="director-create"):
                granted = seed

        return jsonify({"ok": True, "id": new_id, "username": u, "seed_granted": granted})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
# --- Canonical per-user dashboard payload (feeds the four tiles in one call) ---

@app.get("/me/dashboard")
def me_dashboard():
    """
    One-call payload for the Session Stats tiles on the client page.
    Returns only the current user's numbers.
    """
    # must be logged in
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if uid <= 0:
        return jsonify({"ok": False, "error": "not_logged_in"}), 401

        # --- org-aware credits balance + cap info for tiles ---
    org = _user_org_id(uid)
    if org:
        credits_balance = org_balance(org)
        cap = get_user_monthly_cap(org, uid)
        spent = org_user_spent_this_month(org, uid)
        cap_info = {
            "cap": cap,
            "spent": spent,
            "remaining": (None if cap is None else max(0, cap - spent))
        }
    else:
        row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
        credits_balance = int(row[0]) if row else 0
        cap_info = None

    downloads_month = 0
    last_cand = ""
    last_ts_iso = None
    credits_used = 0
    credits_balance = None

    if DB_POOL:
        try:
            # Downloads this month
            row = db_query_one(
                "SELECT COUNT(*) FROM usage_events WHERE user_id=%s AND ts >= date_trunc('month', now())",
                (uid,),
            )
            downloads_month = int(row[0]) if row else 0
        except Exception as e:
            print("me_dashboard count failed:", e)

        try:
            # Last event
            row = db_query_one(
                "SELECT candidate, ts FROM usage_events WHERE user_id=%s ORDER BY ts DESC LIMIT 1",
                (uid,),
            )
            if row:
                last_cand = row[0] or ""
                ts = row[1]
                last_ts_iso = ts.isoformat() if ts else None
        except Exception as e:
            print("me_dashboard last-event failed:", e)

        try:
            # Credits: balance and used (sum of negative deltas as positive number)
            row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
            if row:
                credits_balance = int(row[0])

            row = db_query_one("SELECT COALESCE(SUM(-delta),0) FROM credits_ledger WHERE user_id=%s AND delta < 0", (uid,))
            if row:
                credits_used = int(row[0])
        except Exception as e:
            print("me_dashboard credits failed:", e)

    else:
        # Legacy fallback (very limited)
        try:
            last_cand = STATS.get("last_candidate") or ""
            last_ts_iso = STATS.get("last_time") or None
        except Exception:
            pass

    return jsonify({
        "ok": True,
        "downloadsMonth": downloads_month,
        "lastCandidate": last_cand,
        "lastTime": last_ts_iso,
        "creditsUsed": credits_used,
        "creditsBalance": credits_balance,
    })

# --- Admin: month usage grouped by user (for Director dashboard) ---
@app.get("/__admin/usage-month")
def admin_usage_month():
    """
    Returns counts of usage_events for the current calendar month, grouped by user_id.
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403
    
    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    sql = """
        SELECT user_id, COUNT(*) AS cnt
        FROM usage_events
        WHERE ts >= date_trunc('month', now())
        GROUP BY user_id
        ORDER BY cnt DESC
    """
    sql_total = """
        SELECT COUNT(*) AS total
        FROM usage_events
        WHERE ts >= date_trunc('month', now())
    """
    conn = None
    try:
        conn = DB_POOL.getconn()
        rows = []
        total = 0
        month_start = None
        with conn:
            with conn.cursor() as cur:
                cur.execute("SELECT date_trunc('month', now())::timestamptz")
                month_start = cur.fetchone()[0].isoformat()

                cur.execute(sql)
                for user_id, cnt in cur.fetchall():
                    rows.append({"user_id": user_id, "count": int(cnt)})

                cur.execute(sql_total)
                total = int(cur.fetchone()[0])

        return jsonify({"ok": True, "month_start": month_start, "total": total, "rows": rows})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass

# --- Admin: create an organisation (e.g., "Hamilton") ---
@app.get("/__admin/create-org")
def admin_create_org():
    """
    Usage (admin only):
      /__admin/create-org?name=Hamilton
    Returns: { ok, org_id, already? }
    """
    # guard: admin only
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin = False
    if not is_admin:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    name = (request.args.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "missing name"}), 400

    # create if missing; return existing id if present
    try:
        row = db_query_one("SELECT id FROM orgs WHERE name=%s", (name,))
        if row:
            return jsonify({"ok": True, "org_id": int(row[0]), "already": True})

        ok = db_execute("INSERT INTO orgs (name) VALUES (%s)", (name,))
        if not ok:
            return jsonify({"ok": False, "error": "insert failed"}), 500

        row = db_query_one("SELECT id FROM orgs WHERE name=%s", (name,))
        return jsonify({"ok": True, "org_id": int(row[0]) if row else None})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: set a user's org_id ---
@app.get("/__admin/set-user-org")
def admin_set_user_org():
    """
    Usage (admin only):
      /__admin/set-user-org?user_id=2&org_id=1
    """
    # guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin = False
    if not is_admin:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    # params
    try:
        uid = int(request.args.get("user_id") or "0")
        oid = int(request.args.get("org_id") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id/org_id"}), 400
    if uid <= 0 or oid <= 0:
        return jsonify({"ok": False, "error": "user_id and org_id required"}), 400

    # validate org exists
    try:
        if not db_query_one("SELECT 1 FROM orgs WHERE id=%s", (oid,)):
            return jsonify({"ok": False, "error": "org not found"}), 404
        ok = db_execute("UPDATE users SET org_id=%s WHERE id=%s", (oid, uid))
        if not ok:
            return jsonify({"ok": False, "error": "update failed"}), 500
        return jsonify({"ok": True, "user_id": uid, "org_id": oid})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: create a user (optionally into a specific org) ---
@app.get("/__admin/create-user")
def admin_create_user():
    """
    Usage (admin only):
      /__admin/create-user?u=jane&p=Temp1234!&org_id=2&email=jane@acme.com

    - Creates an active user with the given username/password.
    - If org_id is provided (>0), assigns the user to that org on creation.
    - Idempotent on username: returns {already:true} if it exists.
    """
    # guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin = False
    if not is_admin:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    u = (request.args.get("u") or "").strip()
    p = request.args.get("p") or ""
    email = (request.args.get("email") or "").strip()
    try:
        org_id = int(request.args.get("org_id") or "0")
    except Exception:
        org_id = 0

    if not u or not p:
        return jsonify({"ok": False, "error": "username and password required"}), 400
    if u.lower() == "admin":
        return jsonify({"ok": False, "error": "cannot create/modify 'admin' this way"}), 400

    # if username exists, return its id
    row = db_query_one("SELECT id FROM users WHERE LOWER(username)=LOWER(%s)", (u,))
    if row and row[0]:
        return jsonify({"ok": True, "already": True, "id": int(row[0])})

    # optional org validation
    if org_id > 0:
        if not db_query_one("SELECT 1 FROM orgs WHERE id=%s", (org_id,)):
            return jsonify({"ok": False, "error": "org not found"}), 404

    # create user
    try:
        hashed = generate_password_hash(p)
        if org_id > 0:
            ok = db_execute(
                "INSERT INTO users (username, password_hash, email, active, org_id) VALUES (%s,%s,%s,TRUE,%s)",
                (u, hashed, (email or None), org_id),
            )
        else:
            ok = db_execute(
                "INSERT INTO users (username, password_hash, email, active) VALUES (%s,%s,%s,TRUE)",
                (u, hashed, (email or None)),
            )
        if not ok:
            return jsonify({"ok": False, "error": "insert failed"}), 500

        row = db_query_one("SELECT id, org_id FROM users WHERE LOWER(username)=LOWER(%s)", (u,))
        new_id = int(row[0]) if row and row[0] else None
        new_org = (int(row[1]) if row and row[1] is not None else None)
        return jsonify({"ok": True, "id": new_id, "username": u, "org_id": new_org})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
# --- Admin: backfill org_id onto historical rows for that user ---
@app.get("/__admin/backfill-user-org-data")
def admin_backfill_user_org_data():
    """
    Usage (admin only):
      /__admin/backfill-user-org-data?user_id=2

    Copies users.org_id onto that user's existing usage_events and credits_ledger rows.
    Safe to run multiple times.
    """
    # guard
    try:
        uname = (session.get("user") or "").strip().lower()
        is_admin = bool(session.get("is_admin")) or (uname == "admin")
    except Exception:
        is_admin = False
    if not is_admin:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    if not DB_POOL:
        return jsonify({"ok": False, "error": "DB pool not initialized"}), 500

    # which user?
    try:
        uid = int(request.args.get("user_id") or "0")
    except Exception:
        return jsonify({"ok": False, "error": "bad user_id"}), 400
    if uid <= 0:
        return jsonify({"ok": False, "error": "user_id required"}), 400

    conn = None
    try:
        # ensure the credits_ledger has org_id column too (idempotent safeguard)
        db_execute("ALTER TABLE IF EXISTS credits_ledger ADD COLUMN IF NOT EXISTS org_id INTEGER")
        db_execute("CREATE INDEX IF NOT EXISTS idx_credits_ledger_org_id ON credits_ledger(org_id)")

        # get user's org_id
        row = db_query_one("SELECT org_id FROM users WHERE id=%s", (uid,))
        if not row or not row[0]:
            return jsonify({"ok": False, "error": "user has no org_id set"}), 400
        oid = int(row[0])

        # set org_id where missing on historical rows
        a = db_execute("UPDATE usage_events   SET org_id=%s WHERE user_id=%s AND org_id IS NULL", (oid, uid))
        b = db_execute("UPDATE credits_ledger SET org_id=%s WHERE user_id=%s AND org_id IS NULL", (oid, uid))

        return jsonify({"ok": True, "user_id": uid, "org_id": oid})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# --- Admin: recent usage events (for Director dashboard) ---
@app.get("/__admin/recent-usage")
def admin_recent_usage():
    """
    Returns the most recent usage events.
    Query params:
      - limit (int, optional): number of rows to return, default 50, max 200.
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # Parse & clamp limit
    try:
        limit = int(request.args.get("limit", "50"))
    except Exception:
        limit = 50
    limit = max(1, min(limit, 200))

    # If we have a DB, read from usage_events
    if DB_POOL:
        sql = """
            SELECT id, user_id, ts, candidate, filename
            FROM usage_events
            ORDER BY ts DESC
            LIMIT %s
        """
        conn = None
        try:
            conn = DB_POOL.getconn()
            rows = []
            with conn:
                with conn.cursor() as cur:
                    cur.execute(sql, (limit,))
                    for _id, uid, ts, cand, fname in cur.fetchall():
                        rows.append({
                            "id": int(_id),
                            "user_id": uid,
                            "ts": (ts.isoformat() if ts else None),
                            "candidate": cand or "",
                            "filename": fname or ""
                        })
            return jsonify({"ok": True, "rows": rows, "source": "db"})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
        finally:
            try:
                if conn:
                    DB_POOL.putconn(conn)
            except Exception:
                pass

    # Fallback: legacy JSON history (if DB not initialized)
    out = []
    try:
        for it in (STATS.get("history", []) or [])[::-1][:limit]:
            out.append({
                "id": None,
                "user_id": None,
                "ts": it.get("ts", ""),
                "candidate": it.get("candidate", ""),
                "filename": it.get("filename", "")
            })
    except Exception:
        out = []

    return jsonify({"ok": True, "rows": out, "source": "legacy"})

    # --- Admin: combined dashboard payload (month summary + recent events) ---
@app.get("/__admin/dashboard")
def admin_dashboard():
    """
    Returns:
      {
        ok: true,
        source: "db" | "legacy",
        month: { total: int, rows: [{ user_id, username, count, balance }] },
        recent: [{ ts, user_id, username, candidate, filename }]
      }
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # Parse & clamp limit
    try:
        limit = int(request.args.get("limit", "50"))
    except Exception:
        limit = 50
    limit = max(1, min(limit, 200))

    # Legacy path if no DB
    if not DB_POOL:
        out = []
        try:
            for it in (STATS.get("history", []) or [])[::-1][:limit]:
                out.append({
                    "ts": it.get("ts", ""),
                    "user_id": None,
                    "username": "",
                    "candidate": it.get("candidate", ""),
                    "filename": it.get("filename", "")
                })
        except Exception:
            out = []
        return jsonify({
            "ok": True,
            "source": "legacy",
            "month": {"total": int(STATS.get("downloads", 0)), "rows": []},
            "recent": out
        })

    # DB path
    conn = None
    try:
        conn = DB_POOL.getconn()

        # Month-by-user counts
        with conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT user_id, COUNT(*) AS cnt
                    FROM usage_events
                    WHERE date_trunc('month', ts) = date_trunc('month', now())
                    GROUP BY user_id
                    ORDER BY cnt DESC
                """)
                raw_month = cur.fetchall()
                month_total = sum(int(r[1]) for r in raw_month) if raw_month else 0

        # Recent events
        with conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT ts, user_id, candidate, filename
                    FROM usage_events
                    ORDER BY ts DESC
                    LIMIT %s
                """, (limit,))
                raw_recent = cur.fetchall()

        # Collect referenced user_ids
        uids = set()
        for r in raw_month or []:
            if r[0] is not None:
                uids.add(int(r[0]))
        for r in raw_recent or []:
            if r[1] is not None:
                uids.add(int(r[1]))
        uid_list = list(uids)

        # Map user_id -> username
        name_map = {}
        if uid_list:
            with conn:
                with conn.cursor() as cur:
                    cur.execute("SELECT id, username FROM users WHERE id = ANY(%s)", (uid_list,))
                    for row in cur.fetchall():
                        name_map[int(row[0])] = row[1] or ""

        # Map user_id -> balance
        bal_map = {}
        if uid_list:
            with conn:
                with conn.cursor() as cur:
                    cur.execute("""
                        SELECT user_id, COALESCE(SUM(delta),0)
                        FROM credits_ledger
                        WHERE user_id = ANY(%s)
                        GROUP BY user_id
                    """, (uid_list,))
                    for row in cur.fetchall():
                        bal_map[int(row[0])] = int(row[1])

        # Build outputs
        month_rows = []
        for r in raw_month or []:
            uid = int(r[0]) if r[0] is not None else 0
            cnt = int(r[1])
            month_rows.append({
                "user_id": uid or None,
                "username": name_map.get(uid, ""),
                "count": cnt,
                "balance": bal_map.get(uid) if uid else None,
            })

        recent_rows = []
        for ts, uid, cand, fname in (raw_recent or []):
            uid_int = int(uid) if uid is not None else 0
            recent_rows.append({
                "ts": ts.isoformat() if ts else None,
                "user_id": uid_int or None,
                "username": name_map.get(uid_int, ""),
                "candidate": cand or "",
                "filename": fname or "",
            })

        return jsonify({
            "ok": True,
            "source": "db",
            "month": {"total": month_total, "rows": month_rows},
            "recent": recent_rows
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if conn:
                DB_POOL.putconn(conn)
        except Exception:
            pass
# --- Admin: minimal UI to view the dashboard data (no styling, just tables) ---
@app.get("/__admin/ui")
def admin_ui():
    """
    Simple HTML page for directors to view month summary and recent events.
    Uses /__admin/dashboard under the hood.
    """
    # Access guard: allow only director/admin sessions
    try:
        uname = (session.get("user") or "").strip().lower()
        is_dir = bool(session.get("is_director")) or bool(session.get("is_admin")) or (uname in ("admin", "director"))
    except Exception:
        is_dir = False
    if not is_dir:
        return jsonify({"ok": False, "error": "forbidden"}), 403

    return """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Director Admin UI</title>
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <style>
    body { font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; padding: 16px; }
    h1 { margin: 0 0 8px 0; }
    h2 { margin: 24px 0 8px 0; }
    table { border-collapse: collapse; width: 100%; }
    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
    th { background: #f6f6f6; }
    .muted { color: #666; }
    .badge { display: inline-block; padding: 2px 8px; border: 1px solid #ddd; border-radius: 12px; font-size: 12px; margin-left: 6px; }
    .balance-ok { color: #0a7; font-weight: 600; }
    .balance-low { color: #d9822b; font-weight: 600; }
    .balance-zero { color: #d33; font-weight: 700; }
  </style>
</head>
<body>
  <h1>Director Dashboard <span id="src" class="badge muted"></span></h1>
  <div class="muted">Tip: append <code>?limit=20</code> to the URL to change how many recent rows you load.</div>

  <h2>This Month (by user)</h2>
  <div id="monthBox">Loading…</div>

  <h2>Recent Events</h2>
  <div id="recentBox">Loading…</div>

<section style="margin:24px 0; padding:16px; border:1px solid #ddd; border-radius:8px">
    <h2>Credits Tools</h2>
    <div style="display:flex; gap:24px; flex-wrap:wrap;">
      <form id="grantForm" style="display:flex; gap:8px; align-items:flex-end;">
        <div>
          <label>User ID</label><br>
          <input id="grantUid" type="number" min="1" required style="padding:6px;">
        </div>
        <div>
          <label>Grant (delta)</label><br>
          <input id="grantDelta" type="number" min="1" value="10" required style="padding:6px;">
        </div>
        <div>
          <label>Reason</label><br>
          <input id="grantReason" type="text" value="grant" style="padding:6px;">
        </div>
        <button type="submit" style="padding:8px 12px;">Grant credits</button>
      </form>

      <form id="setForm" style="display:flex; gap:8px; align-items:flex-end;">
        <div>
          <label>User ID</label><br>
          <input id="setUid" type="number" min="1" required style="padding:6px;">
        </div>
        <div>
          <label>Target balance</label><br>
          <input id="setBalance" type="number" min="0" value="0" required style="padding:6px;">
        </div>
        <div>
          <label>Reason</label><br>
          <input id="setReason" type="text" value="adjust" style="padding:6px;">
        </div>
        <button type="submit" style="padding:8px 12px;">Set exact balance</button>
      </form>
    </div>

    <pre id="creditsOut" style="margin-top:12px; padding:12px; background:#f7f7f7; border-radius:6px; max-width:100%; overflow:auto;"></pre>
  </section>

  <script>
  (async function(){
    const out = document.getElementById('creditsOut');
    function print(obj){ out.textContent = JSON.stringify(obj, null, 2); }

    // Grant credits
    const grant = document.getElementById('grantForm');
    grant?.addEventListener('submit', async (e) => {
      e.preventDefault();
      const uid = document.getElementById('grantUid').value;
      const delta = document.getElementById('grantDelta').value;
      const reason = document.getElementById('grantReason').value || 'grant';
      try {
        const r = await fetch(`/__admin/grant-credits?user_id=${encodeURIComponent(uid)}&delta=${encodeURIComponent(delta)}&reason=${encodeURIComponent(reason)}`);
        print(await r.json());
      } catch(err){ print({ ok:false, error:String(err) }); }
    });

    // Set exact balance
    const setf = document.getElementById('setForm');
    setf?.addEventListener('submit', async (e) => {
      e.preventDefault();
      const uid = document.getElementById('setUid').value;
      const bal = document.getElementById('setBalance').value;
      const reason = document.getElementById('setReason').value || 'adjust';
      try {
        const r = await fetch(`/__admin/set-credits?user_id=${encodeURIComponent(uid)}&balance=${encodeURIComponent(bal)}&reason=${encodeURIComponent(reason)}`);
        print(await r.json());
      } catch(err){ print({ ok:false, error:String(err) }); }
    });
  })();
  </script>

  <script>
    (async () => {
      // pass through any ?limit=… query param to the API
      const qs = window.location.search || "";
      const res = await fetch("/__admin/dashboard" + qs);
      if (!res.ok) {
        document.body.innerHTML = "<p>Failed to load dashboard ("+res.status+"). Are you logged in as director/admin?</p>";
        return;
      }
      const d = await res.json();
      const $ = (sel) => document.querySelector(sel);
      const esc = (s) => (s == null ? "" : String(s).replace(/[&<>"]/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}[c])));

      $("#src").textContent = d.source || "";

      // Month table
      const monthRows = (d.month && d.month.rows) || [];
const monthTotal = (d.month && d.month.total) || 0;

if (!monthRows.length) {
  $("#monthBox").textContent = "No usage yet this month.";
} else {
  let html = '<table><thead><tr><th>User</th><th>User ID</th><th>Count</th><th>Balance</th></tr></thead><tbody>';
  for (const r of monthRows) {
    const uname = r.username || '';
    const balNum = (typeof r.balance === 'number') ? r.balance : null;
    const balClass = (balNum === null) ? '' : (balNum <= 0 ? 'balance-zero' : (balNum <= 3 ? 'balance-low' : 'balance-ok'));
    const balCell = (balNum === null) ? '' : `<span class="${balClass}">${esc(balNum)}</span>`;
    html += `<tr><td>${esc(uname)}</td><td>${esc(r.user_id)}</td><td>${esc(r.count)}</td><td>${balCell}</td></tr>`;
  }
  html += `</tbody></table><div class="muted" style="margin-top:6px">Total this month: <strong>${esc(monthTotal)}</strong></div>`;
  $("#monthBox").innerHTML = html;
}

// Recent table
      const recent = d.recent || [];
      if (!recent.length) {
        $("#recentBox").textContent = "No recent events.";
      } else {
         let html = '<table><thead><tr><th>When</th><th>User</th><th>User ID</th><th>Candidate</th><th>Filename</th></tr></thead><tbody>';
      for (const r of recent) {
        const when = r.ts ? new Date(r.ts) : null;
        const whenTxt = when && !isNaN(when.getTime()) ? when.toLocaleString() : (r.ts || "");
        const uname = r.username || '';
        html += `<tr><td>${esc(whenTxt)}</td><td>${esc(uname)}</td><td>${esc(r.user_id)}</td><td>${esc(r.candidate)}</td><td>${esc(r.filename)}</td></tr>`;
      }
        html += '</tbody></table>';
        $("#recentBox").innerHTML = html;
      }
    })().catch(err => {
      document.body.innerHTML = "<p>Unexpected error loading dashboard.</p>";
    });
  </script>
</body>
</html>
    """
# --- Director: minimal UI for org-scoped dashboard (read-only) ---
# --- Director: minimal UI for org-scoped dashboard (read-only + enable/disable) ---
# --- Director UI (fixed: triple quotes + ASCII only) ---


@app.get("/director/ui")
def director_ui():
    """
    Director Console (polished UI).
    - Shows org balance, users, recent events
    - Set monthly cap, enable/disable, delete user
    - Create user, reset password
    """
    # Must be logged in
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0
    if uid <= 0:
        return redirect("/login")

    # Must be director or admin
    try:
        am_admin = bool(session.get("is_admin")) or (session.get("user","").strip().lower() == "admin")
    except Exception:
        am_admin = False
    if not (session.get("director") or am_admin):
        return make_response("forbidden", 403)

    # Resolve org
    org_id = _current_user_org_id()
    if not org_id and am_admin:
        # allow admin to pass ?org_id=...
        try:
            org_id = int(request.args.get("org_id") or "0")
        except Exception:
            org_id = 0
    if not org_id:
        return make_response("No organization assigned to this account.", 403)

    org_name = None
    if DB_POOL:
        r = db_query_one("SELECT name FROM orgs WHERE id=%s", (org_id,))
        org_name = (r[0] if r and r[0] else None)

    #  always define this, regardless of DB_POOL
    org_label = org_name or f"Org #{org_id}"

    # Inline HTML (ASCII only). JS braces are doubled to survive the Python f-string.
    html = f"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Director — Console</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root {{
      --bg:#f7faff; --panel:#ffffff; --ink:#0f172a; --muted:#64748b;
      --brand:#2563eb; --brand2:#22d3ee; --line:#e5e7eb; --ok:#065f46; --off:#b91c1c;
      --radius:16px;
    }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; padding:24px; font:14px/1.5 ui-sans-serif,system-ui,Segoe UI,Roboto,Arial;
           color:var(--ink); background:linear-gradient(180deg,#f7fbff 0%,#f4f8ff 70%); }}
    a {{ color:#2563eb; text-decoration:none; }}
    .wrap {{ max-width:1100px; margin:0 auto; }}
    header {{ display:flex; align-items:center; justify-content:space-between; gap:12px; }}
    h1 {{ margin:0 0 4px 0; font-size:28px; letter-spacing:.2px; }}
    .kicker {{ color:var(--muted); }}
    .grid-metrics {{ display:grid; grid-template-columns:repeat(3,1fr); gap:16px; margin:18px 0; }}
    .metric {{ background:var(--panel); border:1px solid var(--line); border-radius:var(--radius);
               padding:16px; box-shadow:0 8px 24px rgba(2,6,23,.08); }}
    .metric .label {{ color:var(--muted); font-size:12px; }}
    .metric .value {{ font-weight:800; font-size:26px; margin-top:6px; }}

    .grid {{ display:grid; grid-template-columns:1.2fr .8fr; gap:16px; }}
    .card {{ background:var(--panel); border:1px solid var(--line); border-radius:var(--radius);
             padding:16px; box-shadow:0 8px 24px rgba(2,6,23,.08); }}

    table {{ width:100%; border-collapse:collapse; }}
    th,td {{ padding:10px 8px; text-align:left; border-bottom:1px solid var(--line); font-size:13px; }}
    th {{ color:var(--muted); font-weight:600; background:#f8fafc; position:sticky; top:0; z-index:1; }}
    tr:hover td {{ background:#f8fbff; }}
    .pill {{ display:inline-block; padding:2px 8px; border-radius:999px; font-size:12px; border:1px solid var(--line); }}
    .pill.ok {{ background:#ecfdf5; color:#065f46; }}
    .pill.off {{ background:#fef2f2; color:#b91c1c; }}
    .balance-ok {{ color:#065f46; }}
    .balance-low {{ color:#92400e; }}
    .balance-zero {{ color:#b91c1c; }}

    input,button {{ padding:10px 12px; border:1px solid var(--line); border-radius:12px; font-size:14px; }}
    button {{ cursor:pointer; background:#fff; }}
    .btn {{ background:linear-gradient(135deg,var(--brand),var(--brand2)); color:#fff; border:0; }}
    .btn.small {{ padding:8px 12px; border-radius:10px; }}
    .btn.danger {{ background:linear-gradient(135deg,#ef4444,#f97316); }}

    .row {{ display:grid; gap:10px; }}
    @media (min-width:700px) {{ .row2 {{ grid-template-columns:1fr 1fr; }} .row3 {{ grid-template-columns:1fr 1fr 1fr; }} }}

    .hidden {{ display:none; }}
  </style>
</head>
<body>
  <div class="wrap">
    <header>
      <div>
        <h1>Director Console</h1>
        <div class="kicker">Org tools and audit. Org: {org_label}</div>
      </div>
      <div><a href="/app">Back to app</a></div>
    </header>

    <section class="grid-metrics">
      <div class="metric"><div class="label">Org balance</div><div id="poolBox" class="value">—</div></div>
      <div class="metric"><div class="label">Users</div><div id="usersCount" class="value">—</div></div>
      <div class="metric"><div class="label">Recent events</div><div id="recentCount" class="value">—</div></div>
    </section>

    <section class="grid">
      <div style="display:grid; gap:16px;">
        <div class="card">
          <h3 style="margin:0 0 8px">Users</h3>
          <div class="kicker" style="margin:-4px 0 8px">Monthly cap, enable/disable, delete.</div>
          <div class="table-wrap">
            <table>
              <thead>
                <tr>
                  <th>User ID</th>
                  <th>Username</th>
                  <th>Monthly Cap</th>
                  <th>Active</th>
                  <th>Actions</th>
                </tr>
              </thead>
              <tbody id="usersBody"><tr><td colspan="5" class="kicker">Loading…</td></tr></tbody>
            </table>
          </div>
        </div>

        <div class="card">
          <div style="display:flex; align-items:center; justify-content:space-between;">
            <h3 style="margin:0">Recent Activity</h3>
            <button id="ra_toggle" class="btn small" aria-controls="ra_panel" aria-expanded="true">Hide</button>
          </div>
          <div id="ra_panel" style="margin-top:8px;">
            <div id="recentBox" class="kicker">Loading…</div>
          </div>
        </div>
      </div>

      <div style="display:grid; gap:16px;">
        <div class="card">
          <h3 style="margin:0 0 8px">Create User</h3>
          <div class="row row3">
            <input id="cu_u" type="text" placeholder="Username" />
            <input id="cu_p" type="password" placeholder="Password" />
            <input id="cu_seed" type="number" inputmode="numeric" placeholder="Seed credits (optional)" />
          </div>
          <div style="display:flex; gap:10px; align-items:center; margin-top:10px">
            <button id="cu_btn" class="btn">Create</button>
            <div id="cu_msg" class="kicker"></div>
          </div>
        </div>

        <div class="card">
          <h3 style="margin:0 0 8px">Reset User Password</h3>
          <div class="row row2">
            <input id="rp_uid" type="number" inputmode="numeric" placeholder="User ID" />
            <input id="rp_pw" type="password" placeholder="New password" />
          </div>
          <div style="display:flex; gap:10px; align-items:center; margin-top:10px">
            <button id="rp_btn" class="btn">Reset</button>
            <div id="rp_msg" class="kicker"></div>
          </div>
        </div>
      </div>
    </section>
  </div>

  <script>
  // helpers
  const $  = (q) => document.querySelector(q);
  const $$ = (q) => Array.from(document.querySelectorAll(q));
  function esc(s) {{ return String(s ?? '').replace(/[&<>"]/g, c => ({{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;'}})[c]); }}
  async function json(url) {{ try {{ const r = await fetch(url); return await r.json(); }} catch {{ return {{}}; }} }}

  // main loader
  async function loadDashboard() {{
    // optimistic UI
    if ($('#poolBox'))     $('#poolBox').textContent = '—';
    if ($('#usersCount'))  $('#usersCount').textContent = '—';
    if ($('#recentCount')) $('#recentCount').textContent = '—';
    if ($('#usersBody'))   $('#usersBody').innerHTML = '<tr><td colspan="5" class="kicker">Loading…</td></tr>';
    if ($('#recentBox'))   $('#recentBox').textContent = 'Loading…';

    const [d, u] = await Promise.all([
      json('/director/api/dashboard'),
      json('/director/api/users')
    ]);

    // cards
    if ($('#poolBox'))
      $('#poolBox').textContent = (d.pool && typeof d.pool.balance === 'number') ? (d.pool.balance + ' credits') : '—';
    if ($('#usersCount'))
      $('#usersCount').textContent = (u.users && Array.isArray(u.users)) ? u.users.length : '—';
    if ($('#recentCount'))
      $('#recentCount').textContent = (d.recent && Array.isArray(d.recent)) ? d.recent.length : '—';

    // users table
    const rows = Array.isArray(u.users) ? u.users : [];
    if (!rows.length) {{
      if ($('#usersBody')) $('#usersBody').innerHTML = '<tr><td colspan="5" class="kicker">No users yet.</td></tr>';
    }} else {{
      let html = '';
      for (const usr of rows) {{
        const id     = usr.id ?? usr.user_id;
        const uname  = usr.username ?? '';
        const active = Boolean(usr.active ?? true);
        const pill   = `<span class="pill ${{active ? 'ok' : 'off'}}">${{active ? 'Active' : 'Disabled'}}</span>`;
        const next   = active ? 0 : 1;
        html += `
          <tr data-uid="${{id}}">
            <td>${{id}}</td>
            <td>${{esc(uname)}}</td>
            <td>
              <input class="cap" type="number" inputmode="numeric" placeholder="(none)" />
              <button class="btn small setcap">Save</button>
            </td>
            <td>${{pill}}</td>
            <td style="display:flex; gap:8px; flex-wrap:wrap">
              <button class="btn small toggle" data-next="${{next}}">${{active ? 'Disable' : 'Enable'}}</button>
              <button class="btn small danger delete">Delete</button>
            </td>
          </tr>`;
      }}
      if ($('#usersBody')) $('#usersBody').innerHTML = html;
    }}

    // recent table
    const recent = Array.isArray(d.recent) ? d.recent : [];
    if (!recent.length) {{
      if ($('#recentBox')) $('#recentBox').textContent = 'No recent events.';
    }} else {{
      let html = '<table><thead><tr><th>When</th><th>User</th><th>User ID</th><th>Candidate</th><th>Filename</th></tr></thead><tbody>';
      for (const r of recent) {{
        const when    = r.ts ? new Date(r.ts) : null;
        const whenTxt = (when && !isNaN(when.getTime())) ? when.toLocaleString() : (r.ts || '');
        html += `<tr>
          <td>${{esc(whenTxt)}}</td>
          <td>${{esc(r.username || '')}}</td>
          <td>${{esc(r.user_id)}}</td>
          <td>${{esc(r.candidate || '')}}</td>
          <td>${{esc(r.filename  || '')}}</td>
        </tr>`;
      }}
      html += '</tbody></table>';
      if ($('#recentBox')) $('#recentBox').innerHTML = html;
    }}
  }}

  // keep old calls working
  window.loadDash = loadDashboard;

  // event delegation: users table actions + quick actions
  document.addEventListener('click', async (e) => {{
    const tr = e.target.closest('tr[data-uid]');

    // table: set cap
    if (tr && e.target.classList.contains('setcap')) {{
      e.preventDefault();
      const uid    = Number(tr.dataset.uid);
      const capStr = (tr.querySelector('input.cap')?.value || '').trim();
      const url    = new URL('/director/api/setcap', location.origin);
      url.searchParams.set('user_id', String(uid));
      url.searchParams.set('cap', capStr === '' ? 'null' : String(Number(capStr)));
      await fetch(url.toString());
      await loadDashboard();
      return;
    }}

    // table: toggle active
    if (tr && e.target.classList.contains('toggle')) {{
      e.preventDefault();
      const uid  = Number(tr.dataset.uid);
      const next = Number(e.target.getAttribute('data-next'));
      const url  = new URL('/director/api/user/set-active', location.origin);
      url.searchParams.set('user_id', String(uid));
      url.searchParams.set('active', String(next));
      await fetch(url.toString());
      await loadDashboard();
      return;
    }}

    // table: delete user
    if (tr && e.target.classList.contains('delete')) {{
      e.preventDefault();
      if (!confirm('Delete this user permanently?')) return;
      const uid = Number(tr.dataset.uid);
      const url = new URL('/director/api/user/delete', location.origin);
      url.searchParams.set('user_id', String(uid));
      await fetch(url.toString());
      await loadDashboard();
      return;
    }}

    // quick actions: create user
    if (e.target && e.target.id === 'cu_btn') {{
      e.preventDefault();
      const u = $('#cu_u')?.value.trim() || '';
      const p = $('#cu_p')?.value || '';
      const s = $('#cu_seed')?.value.trim() || '';
      const url = new URL('/director/api/create-user', location.origin);
      if (u) url.searchParams.set('u', u);
      if (p) url.searchParams.set('p', p);
      if (s !== '') url.searchParams.set('seed', String(Number(s || 0)));
      const r  = await fetch(url.toString());
      const js = await r.json().catch(() => ({{}}));
      if ($('#cu_msg')) $('#cu_msg').textContent = js.ok ? 'Created.' : (js.error || 'Failed.');
      if (js.ok) {{
        if ($('#cu_u'))    $('#cu_u').value = '';
        if ($('#cu_p'))    $('#cu_p').value = '';
        if ($('#cu_seed')) $('#cu_seed').value = '';
        await loadDashboard();
      }}
      return;
    }}

    // quick actions: reset password
    if (e.target && e.target.id === 'rp_btn') {{
      e.preventDefault();
      const id = Number($('#rp_uid')?.value || '');
      const pw = $('#rp_pw')?.value || '';
      if (!id || !pw) {{
        if ($('#rp_msg')) $('#rp_msg').textContent = 'User ID and new password required.';
        return;
      }}
      const url = '/director/api/user/reset-password?user_id=' + id + '&password=' + encodeURIComponent(pw);
      const r   = await fetch(url);
      const js  = await r.json().catch(() => ({{}}));
      if ($('#rp_msg')) $('#rp_msg').textContent = js.ok ? 'Password reset.' : (js.error || 'Failed.');
      return;
    }}
  }});

  // Recent Activity show/hide + first load (works pre/post DOMContentLoaded)
  (function () {{
    function initDirectorUI() {{
      const btn   = document.getElementById('ra_toggle');
      const panel = document.getElementById('ra_panel');
      if (btn && panel) {{
        const hidden = localStorage.getItem('director_ra_hidden') === '1';
        panel.classList.toggle('hidden', hidden);
        btn.textContent = hidden ? 'Show' : 'Hide';
        btn.setAttribute('aria-expanded', (!hidden).toString());
        btn.addEventListener('click', (e) => {{
          e.preventDefault();
          const nowHidden = panel.classList.toggle('hidden');
          btn.textContent = nowHidden ? 'Show' : 'Hide';
          btn.setAttribute('aria-expanded', (!nowHidden).toString());
          localStorage.setItem('director_ra_hidden', nowHidden ? '1' : '');
        }});
      }}
      loadDashboard().catch(() => {{}});
    }}

    if (document.readyState === 'loading') {{
      document.addEventListener('DOMContentLoaded', initDirectorUI, {{ once: true }});
    }} else {{
      initDirectorUI();
    }}
  }})();
</script>
<h2>Change My Password</h2>
<form onsubmit="return changeMyPass(this)">
  <input type="password" name="password" placeholder="New password" required>
  <button type="submit">Update</button>
</form>

<script>
async function changeMyPass(form) {{
  const pw = form.password.value.trim();
  if(!pw) return false;
  const res = await fetch("/director/api/change-password", {{
    method: "POST",
    headers: {{"Content-Type":"application/json"}},
    body: JSON.stringify({{ password: pw }})
  }});
  const data = await res.json();
  if (data.ok) {{
    alert("Password updated!");
  }} else {{
    alert("Error: " + data.error);
  }}
  return false;
}}
</script>
</body>
</html>
"""
    resp = make_response(html, 200, { "Content-Type": "text/html; charset=utf-8" })
    resp.headers["Cache-Control"] = "no-store"
    return resp

# --- Friendly 402 page (Out of credits) ---
def _render_out_of_credits(reason_text=None):
    # who am I
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    # compute balance (org-aware)
    scope = "anon"
    balance = None
    org_id = None
    try:
        if uid > 0:
            org_id = _user_org_id(uid)
            if org_id:
                scope = "org"
                balance = org_balance(org_id)
            else:
                scope = "user"
                row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid,))
                balance = int(row[0]) if row else 0
    except Exception:
        pass

    msg = reason_text or "You’ve run out of credits."
    bal_str = "" if balance is None else f"{balance}"
    scope_label = {"org":"Your organization pool", "user":"Your account", "anon":"Your account"}[scope]

    html = f"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8"/>
  <title>Out of credits</title>
  <meta name="viewport" content="width=device-width,initial-scale=1"/>
  <style>
    body {{ font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; margin: 40px; }}
    .card {{ max-width: 640px; border: 1px solid #eee; border-radius: 12px; padding: 20px; box-shadow: 0 2px 6px rgba(0,0,0,.06); }}
    h1 {{ margin: 0 0 10px; }}
    .muted {{ color:#666; }}
    .links a {{ display:inline-block; margin-right:12px; }}
    .balance {{ font-size: 18px; margin: 10px 0 16px; }}
    .tag {{ display:inline-block; padding:2px 8px; border:1px solid #ddd; border-radius:12px; font-size:12px; color:#666; margin-left:8px; }}
  </style>
</head>
<body>
  <div class="card">
    <h1>Out of credits <span class="tag">{scope_label}</span></h1>
    <div class="muted">{msg}</div>
    <div class="balance">Current balance: <strong>{bal_str if bal_str!='' else '—'}</strong></div>
    <div class="links">
      <a href="/me/credits" target="_blank">View my credits</a>
      <a href="/director/ui" target="_blank">Director dashboard</a>
      <a href="/">Back to upload</a>
    </div>
  </div>
</body>
</html>
"""
    return make_response(html, 402, {"Content-Type": "text/html; charset=utf-8"})

class PaymentRequired(HTTPException):
    code = 402
    description = "Payment Required"

@app.errorhandler(PaymentRequired)
def on_payment_required(e):
    reason = getattr(e, "description", None)
    return _render_out_of_credits(reason)


# Optional: direct route to preview the page
@app.get("/out-of-credits")
def out_of_credits_preview():
    reason = request.args.get("msg") or "Preview: this is how the page looks when credits run out."
    return _render_out_of_credits(reason)   

# --- Admin: create org tables if missing (safe to run anytime) ---
@app.get("/__admin/ensure-org-schema")
def admin_ensure_org_schema():
    # admin only
    if not (session.get("is_admin")
            or (session.get("username","").lower() == "admin")
            or (session.get("user","").lower() == "admin")):
        return jsonify({"ok": False, "error": "forbidden"}), 403
    if not DB_POOL:
        return jsonify({"ok": False, "error": "db_unavailable"}), 500

    # Create minimal tables used by org credits + per-user caps
    ddl = [
        # org credits ledger
        """
        CREATE TABLE IF NOT EXISTS org_credits_ledger (
          id         SERIAL PRIMARY KEY,
          org_id     INTEGER NOT NULL,
          user_id    INTEGER,
          delta      INTEGER NOT NULL,
          reason     TEXT,
          created_by INTEGER,
          created_at TIMESTAMP DEFAULT NOW()
        )
        """,
        "CREATE INDEX IF NOT EXISTS idx_org_credits_ledger_org ON org_credits_ledger(org_id)",

        # optional per-user monthly caps within an org
        """
        CREATE TABLE IF NOT EXISTS org_user_limits (
          id         SERIAL PRIMARY KEY,
          org_id     INTEGER NOT NULL,
          user_id    INTEGER NOT NULL,
          month_cap  INTEGER,
          updated_at TIMESTAMP DEFAULT NOW()
        )
        """,

        # (optional) orgs table — harmless if you already have one
        """
        CREATE TABLE IF NOT EXISTS orgs (
          id         SERIAL PRIMARY KEY,
          name       TEXT UNIQUE,
          active     BOOLEAN DEFAULT TRUE,
          created_at TIMESTAMP DEFAULT NOW()
        )
        """
    ]

    created = []
    for stmt in ddl:
        ok = db_execute(stmt)
        created.append(bool(ok))

    return jsonify({"ok": True, "created_or_exists": created})

# --- one-time DB column fixer (safe to call anytime) ---
@app.get("/__admin/ensure-core-columns")
def __admin_ensure_core_columns():
    if not (session.get("is_admin")
            or (session.get("username","").lower() == "admin")
            or (session.get("user","").lower() == "admin")):
        return jsonify({"ok": False, "error": "forbidden"}), 403

    results = {}
    # orgs + users need 'active'
    results["orgs_active"]  = bool(db_execute("ALTER TABLE orgs  ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE"))
    results["users_active"] = bool(db_execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE"))

    # orgs need plan fields
    results["orgs_plan_credits_month"] = bool(db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS plan_credits_month INTEGER"))
    results["orgs_plan_name"]          = bool(db_execute("ALTER TABLE orgs ADD COLUMN IF NOT EXISTS plan_name TEXT"))

    # if you use a 'plans' table anywhere, make sure it has these too (harmless if table missing)
    try:
        results["plans_active"] = bool(db_execute("ALTER TABLE plans ADD COLUMN IF NOT EXISTS active BOOLEAN DEFAULT TRUE"))
        results["plans_monthly_credits"] = bool(db_execute("ALTER TABLE plans ADD COLUMN IF NOT EXISTS monthly_credits INTEGER"))
        results["plans_overage_rate"]    = bool(db_execute("ALTER TABLE plans ADD COLUMN IF NOT EXISTS overage_rate NUMERIC"))
    except Exception:
        results["plans_table_present"] = False

    return jsonify({"ok": True, "applied": results})
# ---- Quick diagnostic (no secrets) ----
# ---------- Owner (admin) console ----------
@app.get("/owner/console")
def owner_console():
    if not is_admin():
        return redirect("/login")

    html = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <title>Owner Console — Lustra</title>
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <style>
    :root{--ink:#0f172a;--muted:#64748b;--line:#e5e7eb;--bg:#f6f8fb;--card:#fff;--brand:#2563eb}
    *{box-sizing:border-box}
    body{font:14px/1.45 Inter,system-ui,-apple-system,"Segoe UI",Roboto,Arial,sans-serif;background:var(--bg);color:var(--ink);margin:18px}
    .top{display:flex;justify-content:space-between;align-items:center;margin:0 0 14px}
    .top h1{margin:0;font-size:20px}
    a.btn{display:inline-block;padding:8px 10px;border:1px solid var(--line);border-radius:8px;text-decoration:none;color:var(--ink);background:#fff}
    .grid{display:grid;grid-template-columns:1fr;gap:12px}
    .kpis{display:grid;grid-template-columns:repeat(5,1fr);gap:10px}
    .card{background:var(--card);border:1px solid var(--line);border-radius:10px;padding:12px}
    .k .t{font-size:11px;color:var(--muted);font-weight:700}
    .k .v{font-size:18px;font-weight:900}
    table{width:100%;border-collapse:collapse}
    th,td{padding:8px;border-bottom:1px solid #f1f5f9;text-align:left;vertical-align:middle}
    th{background:#f8fafc;font-size:12px;color:#334155;position:sticky;top:0}
    input[type="text"],input[type="number"]{width:100%;padding:8px;border:1px solid var(--line);border-radius:8px;background:#fff}
    .row{display:flex;gap:8px}
    .small{font-size:12px;color:var(--muted)}
    .pill{display:inline-block;padding:3px 7px;border:1px solid var(--line);border-radius:999px;font-weight:700;font-size:11px;background:#fff}
    .grant{display:flex;gap:6px}
    .grant input{max-width:110px}
    .saveState{font-size:12px;color:var(--muted)}
    /* cap badges */
    .badge{display:inline-block;font-size:12px;padding:2px 6px;border-radius:999px;border:1px solid #e5e7eb;margin-left:6px}
    .badge.ok{background:#f0fdf4;border-color:#dcfce7}
    .badge.bad{background:#fef2f2;border-color:#fee2e2}
  </style>
</head>
<body>
  <div class="top">
    <h1>Owner Console</h1>
    <div class="row">
  <a class="btn" href="/">Home</a>
  <a class="btn" href="/owner/api/export" title="Download last 30 days (all orgs)">Export CSV</a>
  <a class="btn" id="exportOrgBtn" href="#" title="Export a single org">Export CSV (org)</a>
  <a class="btn" id="exportRangeBtn" href="#" title="Export by date range">Export CSV (range)</a>
  <a class="btn" href="/owner/new-client" title="Create org + user (+ optional template/profile)">New Client</a>
  <a class="btn" href="/logout">Log out</a>
</div>
  </div>

  <div class="card" id="usageCard" style="margin:10px 0">
    <div style="display:flex;justify-content:space-between;align-items:center">
      <div style="font-weight:600">Usage (last 30 days)</div>
      <a id="usageAllBtn" class="btn" href="#" title="Show all orgs">All</a>
    </div>
    <svg id="usageSpark" width="100%" height="60"></svg>
  </div>

  <div class="card" id="auditCard" style="margin:10px 0">
    <div style="display:flex;justify-content:space-between;align-items:center">
      <div style="font-weight:600">Credit Audit</div>
      <div>
        <input id="auditOrgId" type="number" min="0" placeholder="org_id (blank=all)" style="width:140px;margin-right:8px" />
        <button id="auditReload" class="btn">Reload</button>
      </div>
    </div>
    <div style="overflow:auto">
      <table id="auditTable" class="table" style="width:100%;margin-top:8px">
        <thead><tr>
          <th style="text-align:left">When</th>
          <th style="text-align:left">Org</th>
          <th style="text-align:right">Δ Credits</th>
          <th style="text-align:left">Reason</th>
        </tr></thead>
        <tbody></tbody>
      </table>
    </div>
  </div>

<script>
document.addEventListener('click', function(e){
  var btn = e.target.closest('#exportOrgBtn'); if(!btn) return;
  e.preventDefault();
  var id = prompt('Org ID to export? (e.g., 1)');
  if(id){ window.location.href = '/owner/api/export?org_id='+encodeURIComponent(id); }
});
</script>

<script>
(async function(){
  try{
    const r = await fetch('/owner/api/overview', {cache:'no-store'});
    const j = await r.json();
    if(!j.ok) return;

    const orgs  = j.orgs || [];
    const tbody = document.getElementById('orgs_tbody');
    if(!tbody) return;

    orgs.forEach(o=>{
      const tr = tbody.querySelector(`tr[data-oid="${o.id}"]`);
      if(!tr) return;

      // ---- Cap badge in the 4th cell (index 3) ----
      const capCell = tr.children[3]; // 0:id, 1:name, 2:plan_name, 3:plan_credits_month
      if(capCell){
        const old = capCell.querySelector('.badge'); if(old) old.remove();
        if (o.cap && o.cap > 0){
          const span = document.createElement('span');
          span.className = 'badge ' + (o.cap_exceeded ? 'bad' : 'ok');
          span.title = `Monthly cap is ${o.cap}; used ${o.usage_month}`;
          span.textContent = o.cap_exceeded ? 'Cap exceeded' : `${(o.cap_remaining ?? 0)} left`;
          span.style.marginLeft = '6px';
          capCell.appendChild(span);
        }
      }

      // ---- Per-row Export button in the actions column (last <td>) ----
      const actions = tr.querySelector('td:last-child');
      if (actions && !actions.querySelector('.export-org-row')) {
        const a = document.createElement('a');
        a.className = 'btn export-org-row';
        a.href = `/owner/api/export?org_id=${o.id}`;
        a.textContent = 'Export';
        a.title = `Download CSV for org ${o.id}`;
        a.style.marginLeft = '6px';
        actions.appendChild(a);
      }

            // ---- Show a green "Template set" badge if this org has a DOCX uploaded
      if (o.has_template && !actions.querySelector('.tpl-badge')) {
        const b = document.createElement('span');
        b.className = 'badge ok tpl-badge';
        b.textContent = 'Template set';
        b.title = o.template_updated_at ? ('Updated ' + o.template_updated_at) : 'Template present';
        b.style.marginRight = '6px';
        actions.insertBefore(b, actions.firstChild);
      }
      
      // ---- Per-row Profile button (opens JSON profile editor)
if (actions && !actions.querySelector('.profile-org-row')) {
  const p = document.createElement('a');
  p.className = 'btn profile-org-row';
  p.href = `/__admin/org-profile?org_id=${o.id}`;
  p.textContent = 'Profile';
  p.title = `Open org ${o.id} profile`;
  p.style.marginLeft = '6px';
  actions.appendChild(p);
}

// ---- Per-row Template button (opens DOCX upload form)
if (actions && !actions.querySelector('.template-org-row')) {
  const t = document.createElement('a');
  t.className = 'btn template-org-row';
  t.href = `/__admin/upload-org-template?org_id=${o.id}`;
  t.textContent = 'Template';
  t.title = `Upload/replace DOCX template for org ${o.id}`;
  t.style.marginLeft = '6px';
  actions.appendChild(t);
}

// ---- Low-credit badge in the Credits column (index 7) ----
// columns: 0:id 1:name 2:plan 3:cap 4:users 5:usage_m 6:usage_total 7:balance 8:actions
const balCell = tr.children[7];
if (balCell) {
  const old = balCell.querySelector('.low-badge');
  if (old) old.remove();
  const bal = Number(o.credits_balance || 0);
  if (!Number.isNaN(bal) && bal <= 5) {
    const b = document.createElement('span');
    b.className = 'badge bad low-badge';
    b.textContent = 'Low';
    b.title = `Only ${bal} credits left`;
    b.style.marginLeft = '6px';
    balCell.appendChild(b);
  }
}
    });
  }catch(e){
    console.log('cap badge / export add failed', e);
  }
})();
</script>

<script>
document.addEventListener('click', function(e){
  const btn = e.target.closest('#exportRangeBtn'); if(!btn) return;
  e.preventDefault();

  const org   = prompt('Org ID (leave empty for ALL)');
  const start = prompt('Start date (YYYY-MM-DD), e.g. 2025-09-01');
  const end   = prompt('End date (YYYY-MM-DD), e.g. 2025-09-16');

  function ok(d){ return /^\d{4}-\d{2}-\d{2}$/.test(d||''); }
  if(!ok(start) || !ok(end)){
    alert('Please enter dates as YYYY-MM-DD.');
    return;
  }

  const params = new URLSearchParams();
  params.set('start', start);
  params.set('end', end);
  if((org||'').trim()) params.set('org_id', org.trim());

  window.location.href = '/owner/api/export?' + params.toString();
});
</script>

  <div class="grid">
    <div class="kpis">
      <div class="card k"><div class="t">Orgs</div><div class="v" id="k_orgs">—</div></div>
      <div class="card k"><div class="t">Active orgs</div><div class="v" id="k_orgs_active">—</div></div>
      <div class="card k"><div class="t">Users</div><div class="v" id="k_users">—</div></div>
      <div class="card k"><div class="t">Usage (30d)</div><div class="v" id="k_usage30">—</div></div>
      <div class="card k"><div class="t">Credits (sum)</div><div class="v" id="k_creds">—</div></div>
    </div>

    <div class="card">
      <div class="row" style="justify-content:space-between;align-items:center;margin-bottom:8px">
        <h2 style="margin:0;font-size:16px">Organisations</h2>
        <div class="saveState" id="saveState"></div>
      </div>
      <div class="small" style="margin-bottom:8px">Edit <strong>Name</strong>, <strong>Plan</strong> (plan_name) and <strong>Plan credits</strong> (plan_credits_month). Use <em>Grant</em> to top up org credits now. Changes auto-save on blur.</div>
      <div style="overflow:auto;max-height:70vh;border:1px solid var(--line);border-radius:10px">
        <table>
          <thead>
            <tr>
              <th>ID</th>
              <th style="min-width:200px">Name</th>
              <th>Active</th>
              <th style="min-width:160px">Plan</th>
              <th>Plan credits/mo</th>
              <th>Balance</th>
              <th>Usage (month)</th>
              <th>Usage (total)</th>
              <th>Users</th>
              <th>Grant</th>
              <th>Director</th>
            </tr>
          </thead>
          <tbody id="tbody"></tbody>
        </table>
      </div>
    </div>
  </div>

<script>
function qs(el,sel){return el.querySelector(sel)}
function fmt(n){return new Intl.NumberFormat('en-GB').format(n||0)}

let data=null, saveTimer=null;

async function load(){
  const r = await fetch('/owner/api/overview', {cache:'no-store'});
  if(!r.ok){ alert('Failed to load overview'); return; }
  data = await r.json();
  if(!data.ok){ alert(data.error||'Overview error'); return; }

  // KPIs
  document.getElementById('k_orgs').textContent = fmt(data.kpis.total_orgs);
  document.getElementById('k_orgs_active').textContent = fmt(data.kpis.active_orgs);
  document.getElementById('k_users').textContent = fmt(data.kpis.total_users);
  document.getElementById('k_usage30').textContent = fmt(data.kpis.usage_30d);
  document.getElementById('k_creds').textContent = fmt(data.kpis.credits_balance_sum);

  // Table
  const tb = document.getElementById('tbody'); tb.innerHTML='';
  (data.orgs||[]).forEach(o=>{
    const tr = document.createElement('tr');
    tr.innerHTML = `
      <td>${o.id}</td>
      <td><input type="text" value="${(o.name||'').replaceAll('"','&quot;')}" data-k="name"></td>
      <td>${o.active ? 'Yes' : 'No'}</td>
      <td><input type="text" value="${(o.plan_name||'').replaceAll('"','&quot;')}" data-k="plan_name"></td>
      <td><input type="number" value="${o.plan_credits_month||0}" min="0" step="1" data-k="plan_credits_month"></td>
      <td><span class="pill" data-k="credits_balance">${fmt(o.credits_balance||0)}</span></td>
      <td>${fmt(o.usage_month||0)}</td>
      <td>${fmt(o.usage_total||0)}</td>
      <td>${fmt(o.users_count||0)}</td>
      <td>
        <div class="grant">
          <input type="number" placeholder="+100" step="1" />
          <button type="button">Grant</button>
        </div>
      </td>
      <td>
  <div class="row">
    <a class="btn" href="/owner/console">Owner</a>
    <a class="btn" href="/director">Usage</a>
    <a class="btn" href="/director?org_id=${o.id}">Director</a>
    <a class="btn" href="/owner/api/export?org_id=${o.id}">Export</a>
    <a class="btn" href="/__admin/org-profile?org_id=${o.id}">Profile</a>
    <a class="btn" href="/__admin/upload-org-template">Template</a>
  </div>
</td>
    `;
    // auto-save on blur
    tr.querySelectorAll('input[data-k]').forEach(inp=>{
      inp.addEventListener('blur', ()=> saveRow(o.id, tr));
    });
    // grant handler
    const gBtn = tr.querySelector('.grant button');
    const gInp = tr.querySelector('.grant input');
    gBtn.addEventListener('click', async ()=>{
      const delta = parseInt(gInp.value||'0',10);
      if(!delta) return;
      await saveRow(o.id, tr, delta);
      gInp.value='';
    });

    tb.appendChild(tr);
  });
}

async function saveRow(id, tr, grantDelta){
  const name  = qs(tr,'input[data-k="name"]').value;
  const plan  = qs(tr,'input[data-k="plan_name"]').value;
  const creds = parseInt(qs(tr,'input[data-k="plan_credits_month"]').value||'0',10);
  const qsParams = new URLSearchParams({ id: String(id), name, plan_name: plan, plan_credits_month: String(creds) });
  if (grantDelta) qsParams.set('grant', String(grantDelta));

  setSaveState('Saving…');
  const r = await fetch('/owner/api/set-org-plan?' + qsParams.toString(), { method:'GET' });
  const j = await r.json().catch(()=>({ok:false}));
  if(j && j.ok){
    const bal = j.credits_balance;
    if (typeof bal === 'number') {
      const chip = qs(tr,'[data-k="credits_balance"]');
      if (chip) chip.textContent = new Intl.NumberFormat('en-GB').format(bal);
    }
    setSaveState('Saved');
  }else{
    setSaveState('Error');
    alert(j && j.error ? j.error : 'Save failed');
  }
  if (saveTimer) clearTimeout(saveTimer);
  saveTimer = setTimeout(()=>setSaveState(''), 1200);
}

function setSaveState(t){
  const el = document.getElementById('saveState');
  if(el) el.textContent = t||'';
}

document.addEventListener('DOMContentLoaded', load);
</script>
</body>
</html>
    """
    
    html += """
    <script>
    (async function(){
      const svg = document.getElementById('usageSpark'); if(!svg) return;
      const btnAll = document.getElementById('usageAllBtn');
      const q = new URLSearchParams({days:'30'}); // add org_id later for per-org sparkline

      try{
        const r = await fetch('/owner/api/usage-series?'+q.toString(), {cache:'no-store'});
        const j = await r.json();
        if(!j.ok) return;

        const s = j.series || [];
        const w = svg.clientWidth || 600, h = svg.clientHeight || 60;
        if (!s.length){ svg.innerHTML=''; return; }

        const xs = s.map((_,i)=>i), ys = s.map(o=>o.count||0);
        const xmin=0, xmax=xs.length-1, ymin=0, ymax=Math.max(1, ...ys);
        const x = i => (w-8) * (i - xmin) / Math.max(1,(xmax-xmin)) + 4;
        const y = v => h - 6 - (h-12) * (v - ymin) / Math.max(1,(ymax - ymin));
        const d = xs.map((i)=>`${x(i)},${y(ys[i])}`).join(' ');

        svg.setAttribute('viewBox', `0 0 ${w} ${h}`);
        svg.innerHTML = `
          <polyline points="${d}" fill="none" stroke="currentColor" stroke-width="2" opacity="0.9"></polyline>
          <line x1="0" y1="${y(0)}" x2="${w}" y2="${y(0)}" stroke="currentColor" stroke-width="0.5" opacity="0.15"></line>
        `;
      }catch(e){
        console.log('usage spark failed', e);
      }
    })();
    </script>
    """

    html += """
    <script>
    (async function(){
      const reloadBtn = document.getElementById('auditReload');
      const orgInput = document.getElementById('auditOrgId');
      const tbody = document.querySelector('#auditTable tbody');
      if(!reloadBtn || !tbody) return;

      async function loadAudit(){
        const q = new URLSearchParams({limit: '200'});
        const orgVal = (orgInput && orgInput.value || '').trim();
        if(orgVal) q.set('org_id', orgVal);

        tbody.innerHTML = '<tr><td colspan="4">Loading…</td></tr>';
        try{
          const r = await fetch('/owner/api/credits-ledger?' + q.toString(), {cache:'no-store'});
          const j = await r.json();
          if(!j.ok) { tbody.innerHTML = '<tr><td colspan="4">Forbidden</td></tr>'; return; }
          const items = j.items || [];
          if(!items.length){ tbody.innerHTML = '<tr><td colspan="4">No entries</td></tr>'; return; }
          tbody.innerHTML = items.map(it => `
            <tr>
              <td>${new Date(it.ts).toLocaleString()}</td>
              <td>${it.org_name || ('#'+it.org_id)}</td>
              <td style="text-align:right">${(it.delta>0?'+':'') + it.delta}</td>
              <td>${(it.reason||'').replace(/</g,'&lt;')}</td>
            </tr>`).join('');
        }catch(e){
          console.log('audit load failed', e);
          tbody.innerHTML = '<tr><td colspan="4">Error loading</td></tr>';
        }
      }

      reloadBtn.addEventListener('click', loadAudit);
      loadAudit();
    })();
    </script>
    """
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

# --- Owner: New Client wizard (admin-only; orchestrates existing admin endpoints) ---
@app.get("/owner/new-client")
def owner_new_client():
    if not is_admin():
        return redirect("/login")

    html = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>New Client Wizard</title>
  <style>
    :root{--muted:#64748b}
    body{font:14px/1.45 system-ui,Segoe UI,Roboto,Arial,sans-serif;padding:20px;color:#0f172a}
    .row{display:flex;gap:8px;flex-wrap:wrap;align-items:center}
    .btn{display:inline-block;padding:8px 12px;border:1px solid #e5e7eb;border-radius:10px;background:#fff;text-decoration:none;color:#0f172a}
    .card{border:1px solid #e5e7eb;border-radius:12px;padding:16px;margin-top:12px}
    label{display:block;margin:8px 0 4px;color:#334155}
    input,textarea,select{width:100%;padding:8px;border:1px solid #e5e7eb;border-radius:10px}
    .grid{display:grid;grid-template-columns:1fr 1fr;gap:16px}
    .muted{color:var(--muted)}
    .ok{color:#15803d}
    .err{color:#b91c1c}
    .mono{font-family:ui-monospace,Menlo,Consolas,monospace;font-size:12px;white-space:pre-wrap}
  </style>
</head>
<body>
  <div class="row" style="margin-bottom:12px">
    <a class="btn" href="/owner/console">Owner</a>
    <a class="btn" href="/__admin/new-org">New org (manual)</a>
    <a class="btn" href="/__admin/new-user">Create user (manual)</a>
    <a class="btn" href="/__admin/upload-org-template">Upload template (manual)</a>
    <a class="btn" href="/__admin/org-profile">Org profile (manual)</a>
  </div>

  <h1 style="margin:0 0 8px">New Client Wizard</h1>
  <p class="muted" style="margin-top:0">Creates an organisation, a user attached to it, and optionally uploads a DOCX template and saves a per-org profile JSON. Uses your existing admin endpoints.</p>

  <form id="wiz" class="card">
    <div class="grid">
      <div class="card">
        <h3 style="margin-top:0">1) Organisation</h3>
        <label>Organisation name</label>
        <input name="org_name" placeholder="e.g., Hamilton Recruitment" required>
      </div>

      <div class="card">
        <h3 style="margin-top:0">2) User</h3>
        <label>Username</label>
        <input name="username" placeholder="e.g., hamilton" required>
        <label>Email</label>
        <input name="email" type="email" placeholder="e.g., ops@client.com" required>
        <label>Password</label>
        <input name="password" type="password" placeholder="Temporary password" required>
      </div>

      <div class="card">
        <h3 style="margin-top:0">3) Optional: Template (.docx)</h3>
        <input name="template" type="file" accept=".docx">
      </div>

      <div class="card">
        <h3 style="margin-top:0">4) Optional: Profile JSON</h3>
        <textarea name="profile" rows="10" class="mono" placeholder='{"sections_order":["name_and_contact","executive_summary","skills","experience","education"]}'></textarea>
        <p class="muted" style="margin:6px 0 0">Leave empty to inherit the default (Hamilton) structure. This does not change polishing unless we wire it in later.</p>
      </div>
    </div>

    <div class="row" style="margin-top:12px">
      <button class="btn" type="submit">Create client</button>
      <span id="status" class="muted"></span>
    </div>
  </form>

  <div id="result" class="card mono" style="display:none"></div>

<script>
(async function(){
  const form = document.getElementById('wiz');
  const statusEl = document.getElementById('status');
  const resultEl = document.getElementById('result');

  function setStatus(t, cls){ statusEl.textContent=t; statusEl.className = cls||'muted'; }
  function showResult(obj){
    resultEl.style.display='block';
    try{ resultEl.textContent = JSON.stringify(obj, null, 2); }
    catch(e){ resultEl.textContent = String(obj); }
  }

  form.addEventListener('submit', async (ev)=>{
    ev.preventDefault();
    setStatus('Working...', 'muted');
    resultEl.style.display='none';

    const fd = new FormData(form);
    const orgName = (fd.get('org_name')||'').trim();
    const username = (fd.get('username')||'').trim();
    const email = (fd.get('email')||'').trim();
    const password = (fd.get('password')||'').trim();
    const template = fd.get('template');
    const profile  = (fd.get('profile')||'').toString().trim();

    const out = {steps:[]};
    try{
      // 1) Create org
      const r1 = await fetch('/__admin/create-org?name='+encodeURIComponent(orgName), {credentials:'same-origin'});
      const j1 = await r1.json();
      out.steps.push({create_org:j1});
      if(!j1.ok){ setStatus('Failed creating org', 'err'); showResult(out); return; }
      const org_id = j1.id || j1.org_id;
      if(!org_id){ setStatus('No org_id returned', 'err'); showResult(out); return; }

      // 2) Create user in that org
      const uURL = new URL('/__admin/create-user', location.origin);
      uURL.searchParams.set('u', username);
      uURL.searchParams.set('p', password);
      uURL.searchParams.set('org_id', org_id);
      uURL.searchParams.set('email', email);
      const r2 = await fetch(uURL.toString(), {credentials:'same-origin'});
      const j2 = await r2.json();
      out.steps.push({create_user:j2});
      if(!j2.ok){ setStatus('Failed creating user', 'err'); showResult(out); return; }

      // 3) Optional: upload template
      if(template && template.name){
        const tfd = new FormData();
        tfd.append('org_id', org_id);
        tfd.append('template', template);
        const r3 = await fetch('/__admin/upload-org-template', {method:'POST', body:tfd, credentials:'same-origin'});
        let j3; try{ j3 = await r3.json(); } catch(e){ j3 = {ok:false, error:'template upload parse fail'}; }
        out.steps.push({upload_template:j3});
        if(!j3.ok){ setStatus('Template upload failed (continuing)', 'err'); }
      }

      // 4) Optional: save profile JSON
      if(profile){
        let obj;
        try{ obj = JSON.parse(profile); }
        catch(e){ out.steps.push({profile:'invalid JSON — skipped'}); obj = null; }
        if(obj){
          const pfd = new URLSearchParams();
          pfd.set('org_id', org_id);
          pfd.set('profile', JSON.stringify(obj));
          const r4 = await fetch('/__admin/org-profile', {
            method:'POST',
            headers:{'Content-Type':'application/x-www-form-urlencoded'},
            body:pfd.toString(),
            credentials:'same-origin'
          });
          const j4 = await r4.json();
          out.steps.push({save_profile:j4});
          if(!j4.ok){ setStatus('Profile save failed (continuing)', 'err'); }
        }
      }

      setStatus('Done', 'ok');
      out.links = {
        owner_console: '/owner/console',
        org_profile: '/__admin/org-profile?org_id='+org_id,
        director: '/director?org_id='+org_id
      };
      showResult(out);
    }catch(e){
      setStatus('Error', 'err');
      out.error = String(e);
      showResult(out);
    }
  });
})();
</script>

</body>
</html>
"""

    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

# --- Owner: daily usage series (admin-only, read-only) ---
@app.get("/owner/api/usage-series")
def owner_api_usage_series():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # params
    try:
        days = int(request.args.get("days", "30"))
    except Exception:
        days = 30
    if days < 1 or days > 365:
        days = 30

    org_id = request.args.get("org_id", "").strip()
    try:
        org_id = int(org_id) if org_id else None
    except Exception:
        org_id = None

    # Build range (uses 'ts', keep consistent)
    if org_id:
        rows = db_query_all("""
            SELECT date_trunc('day', ue.ts)::date AS d, COUNT(*)
              FROM usage_events ue
              JOIN users u ON u.id = ue.user_id
             WHERE ue.ts >= now() - (%s || ' days')::interval
               AND u.org_id = %s
             GROUP BY 1
             ORDER BY 1
        """, (days, org_id)) or []
    else:
        rows = db_query_all("""
            SELECT date_trunc('day', ts)::date AS d, COUNT(*)
              FROM usage_events
             WHERE ts >= now() - (%s || ' days')::interval
             GROUP BY 1
             ORDER BY 1
        """, (days,)) or []

    # Fill gaps with 0s
    from datetime import datetime, timedelta
    today = datetime.utcnow().date()
    start = today - timedelta(days=days - 1)
    by_day = {
        (r[0].isoformat() if hasattr(r[0], "isoformat") else str(r[0])): int(r[1] or 0)
        for r in rows
    }
    out = []
    cur = start
    while cur <= today:
        key = cur.isoformat()
        out.append({"date": key, "count": by_day.get(key, 0)})
        cur += timedelta(days=1)

    return jsonify({"ok": True, "days": days, "series": out})

@app.get("/owner/api/overview")
def owner_api_overview():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # --- Orgs base info (keep columns minimal + stable) ---
    org_rows = db_query_all("""
        SELECT id,
               name,
               COALESCE(active, TRUE) AS active,
               COALESCE(plan_name, '') AS plan_name,
               COALESCE(plan_credits_month, 0) AS plan_credits_month
          FROM orgs
         ORDER BY id
    """) or []

    # --- Aggregates by org ---
    cred_rows  = db_query_all("SELECT org_id, COALESCE(SUM(delta),0) FROM org_credits_ledger GROUP BY org_id") or []
    month_rows = db_query_all("""
        SELECT org_id, COUNT(*)
          FROM usage_events
         WHERE date_trunc('month', ts) = date_trunc('month', now())
         GROUP BY org_id
    """) or []
    total_rows = db_query_all("SELECT org_id, COUNT(*) FROM usage_events GROUP BY org_id") or []
    users_rows = db_query_all("""
        SELECT org_id, COUNT(*)
          FROM users
         WHERE COALESCE(active, TRUE) = TRUE
           AND LOWER(username) <> 'admin'
         GROUP BY org_id
    """) or []

    cred = {r[0]: int(r[1] or 0) for r in cred_rows}
    usem = {r[0]: int(r[1] or 0) for r in month_rows}
    uset = {r[0]: int(r[1] or 0) for r in total_rows}
    ucnt = {r[0]: int(r[1] or 0) for r in users_rows}

    # --- Template status per org (optional UI badges) ---
    tpl_rows = db_query_all("""
        SELECT id,
               (CASE WHEN COALESCE(template_path,'') <> '' THEN TRUE ELSE FALSE END) AS has_template,
               template_updated_at
          FROM orgs
    """) or []
    tpl_has  = {r[0]: bool(r[1]) for r in tpl_rows}
    tpl_when = {r[0]: (r[2].isoformat() if hasattr(r[2], "isoformat") else (str(r[2]) if r[2] else None)) for r in tpl_rows}

    # --- Build response rows (no stray indents, no created_at index mismatch) ---
    orgs = []
    for r in org_rows:
        oid = r[0]
        cap = int(r[4] or 0)
        usage_m = int(usem.get(oid, 0))
        exceeded = (cap > 0 and usage_m > cap)
        remaining = (cap - usage_m) if cap > 0 else None
        if remaining is not None and remaining < 0:
            remaining = 0

        orgs.append({
            "id": oid,
            "name": r[1],
            "active": bool(r[2]),
            "plan_name": r[3],
            "plan_credits_month": cap,
            "credits_balance": int(cred.get(oid, 0)),
            "usage_month": usage_m,
            "usage_total": int(uset.get(oid, 0)),
            "users_count": int(ucnt.get(oid, 0)),
            # extra badges for UI:
            "has_template": bool(tpl_has.get(oid, False)),
            "template_updated_at": tpl_when.get(oid),
            "cap": cap,
            "cap_exceeded": bool(exceeded),
            "cap_remaining": (int(remaining) if remaining is not None else None),
        })

    # --- KPIs ---
    k_total_orgs = len(orgs)
    k_active_orgs = sum(1 for o in orgs if o["active"])
    k_total_users = sum(ucnt.values()) if ucnt else 0
    k_usage_30d   = int((db_query_one("SELECT COUNT(*) FROM usage_events WHERE ts >= now() - interval '30 days'")[0]) or 0)
    k_cred_sum    = int((db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger")[0]) or 0)

    return jsonify({
        "ok": True,
        "kpis": {
            "total_orgs": k_total_orgs,
            "active_orgs": k_active_orgs,
            "total_users": k_total_users,
            "usage_30d": k_usage_30d,
            "credits_balance_sum": k_cred_sum,
        },
        "orgs": orgs,
    })

@app.get("/owner/api/set-org-plan")
def owner_api_set_org_plan():
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # Inputs
    try:
        org_id = int(request.args.get("id", "0"))
    except Exception:
        org_id = 0
    if not org_id:
        return jsonify({"ok": False, "error": "missing id"}), 400

    name  = (request.args.get("name") or "").strip()
    plan  = (request.args.get("plan_name") or "").strip()
    try:
        plan_credits = int(request.args.get("plan_credits_month", "0") or 0)
    except Exception:
        plan_credits = 0

    # Optional grant/top-up to the org credits pool
    try:
        grant = int(request.args.get("grant", "0") or 0)
    except Exception:
        grant = 0

    # Update org fields
    if name:
        db_execute("UPDATE orgs SET name=%s WHERE id=%s", (name, org_id))
    db_execute("UPDATE orgs SET plan_name=%s, plan_credits_month=%s WHERE id=%s",
               (plan or None, plan_credits, org_id))

    # Record credit grant if provided
    if grant:
        admin_user = session.get("user") or ""
        created_by = None
        try:
            u = get_user_db(admin_user)
            if u and u.get("id"):
                created_by = int(u["id"])
        except Exception:
            pass
        db_execute(
            "INSERT INTO org_credits_ledger (org_id, delta, reason, created_by) VALUES (%s,%s,%s,%s)",
            (org_id, grant, "grant", created_by)
        )

    # Return fresh balance
    row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM org_credits_ledger WHERE org_id=%s", (org_id,))
    balance = int(row[0] or 0) if row else 0

    return jsonify({"ok": True, "id": org_id, "credits_balance": balance})

    # --- Owner: export usage CSV (admin-only) ---
@app.get("/owner/api/export")
def owner_api_export():
    # Guard
    if not is_admin():
        return jsonify({"ok": False, "error": "forbidden"}), 403

    # Params: ?org_id=1&start=YYYY-MM-DD&end=YYYY-MM-DD
    import io, csv
    from datetime import datetime, timedelta

    org_id = request.args.get("org_id", "").strip()
    try:
        org_id = int(org_id) if org_id else None
    except Exception:
        org_id = None

    def _parse_date(s):
        try:
            return datetime.strptime(s, "%Y-%m-%d")
        except Exception:
            return None

    start = _parse_date(request.args.get("start", ""))
    end   = _parse_date(request.args.get("end", ""))

    # Default to last 30 days if not provided
    if not end:
        end = datetime.utcnow()
    if not start:
        start = end - timedelta(days=30)

    # Make end exclusive (+1 day if start==end)
    if end <= start:
        end = start + timedelta(days=1)

    rows = []
    try:
        if org_id:
            rows = db_query_all("""
                SELECT
                    ue.created_at,
                    u.org_id,
                    COALESCE(o.name, '') AS org_name,
                    u.id AS user_id,
                    COALESCE(u.username, '') AS username,
                    COALESCE(ue.candidate_name, '') AS candidate_name,
                    COALESCE(ue.filename, '') AS filename
                FROM usage_events ue
                LEFT JOIN users u ON u.id = ue.user_id
                LEFT JOIN orgs  o ON o.id = u.org_id
                WHERE ue.created_at >= %s AND ue.created_at < %s
                  AND (u.org_id = %s)
                ORDER BY ue.created_at DESC
            """, (start, end, org_id)) or []
        else:
            rows = db_query_all("""
                SELECT
                    ue.created_at,
                    u.org_id,
                    COALESCE(o.name, '') AS org_name,
                    u.id AS user_id,
                    COALESCE(u.username, '') AS username,
                    COALESCE(ue.candidate_name, '') AS candidate_name,
                    COALESCE(ue.filename, '') AS filename
                FROM usage_events ue
                LEFT JOIN users u ON u.id = ue.user_id
                LEFT JOIN orgs  o ON o.id = u.org_id
                WHERE ue.created_at >= %s AND ue.created_at < %s
                ORDER BY ue.created_at DESC
            """, (start, end)) or []
    except Exception as e:
        return jsonify({"ok": False, "error": f"query failed: {e}"}), 500

    # Build CSV
    sio = io.StringIO()
    w = csv.writer(sio)
    w.writerow(["timestamp_utc", "org_id", "org_name", "user_id", "username", "candidate", "filename"])
    for r in rows:
        ts = r[0]
        ts_str = ts.isoformat() if hasattr(ts, "isoformat") else str(ts)
        w.writerow([ts_str, r[1], r[2], r[3], r[4], r[5], r[6]])

    csv_bytes = sio.getvalue().encode("utf-8")
    fname = f'usage_export_{datetime.utcnow().strftime("%Y%m%d")}.csv'
    resp = make_response(csv_bytes, 200, {
        "Content-Type": "text/csv; charset=utf-8",
        "Content-Disposition": f'attachment; filename="{fname}"'
    })
    return resp
# --- Hard block: non-admins cannot modify the 'admin' user via any toggle/enable/disable/delete route ---
def _is_admin_session():
    try:
        uname = (session.get("user") or "").strip().lower()
        return uname == "admin" or bool(session.get("is_admin")) or bool(session.get("is_director") and uname == "admin")
    except Exception:
        return False

@app.before_request
def _protect_root_admin_from_mutation():
    """
    Safety net: if a request tries to modify the 'admin' user and the session is NOT admin,
    block it. We look at common mutation endpoints and read the target username from query/form.
    """
    try:
        path = (request.path or "").lower()
        # Only inspect potentially mutating areas to keep overhead tiny
        if not any(seg in path for seg in ("/director", "/admin", "/legacy", "/user", "/users")):
            return

        # target username can arrive as ?username=, ?user=, ?u= or in POST body
        target = (
            (request.values.get("username")
             or request.values.get("user")
             or request.values.get("u")
             or "")
        ).strip().lower()

        # If someone targets 'admin' on a mutating route and current session isn't admin -> forbid
        if target == "admin" and any(tok in path for tok in ("disable", "enable", "toggle", "delete", "remove", "deactivate", "activate", "set", "update", "create")):
            if not _is_admin_session():
                return jsonify({"ok": False, "error": "cannot_modify_admin"}), 403
    except Exception:
        # Never take the site down because of the guard
        pass

# ---- Quick diagnostic (no secrets) ----
@app.get("/__me/diag")
def me_diag_v2():
    try:
        uid = int(session.get("user_id") or 0)
    except Exception:
        uid = 0

    try:
        month_cnt = int(count_usage_month_db(uid)) if (DB_POOL and uid) else 0
    except Exception:
        month_cnt = 0

    try:
        c, t = last_event_for_user(uid) if (DB_POOL and uid) else (None, None)
    except Exception:
        c, t = None, None

    return jsonify({
        "ok": True,
        "logged_in": bool(uid),
        "user_id": uid or None,
        "username": session.get("user") or None,
        "db_pool": bool(DB_POOL),
        "month_usage": month_cnt,
        "last_event": {"candidate": c or "", "ts": t or ""},
    })

# ---------- Director routes (refreshed) ----------

# ---------- Director routes (refreshed) ----------

@app.get("/director")
def director_home():
    # If already a director or admin, go straight to the console
    if session.get("director") or is_admin():
        return redirect(url_for("director_ui"))
    # Otherwise show the director login page
    return render_template_string(DIRECTOR_LOGIN_HTML)

@app.post("/director/login")
def director_login():
    pw = (request.form.get("password") or "").strip()
    if pw == STATS.get("director_pass_override", "director"):
        session["director"] = True
        return redirect(url_for("director_ui"))
    html = DIRECTOR_LOGIN_HTML.replace(
        "<!--DERR-->", "<div class='err'>Incorrect director password</div>"
    )
    return make_response(render_template_string(html), 401)

@app.get("/director/logout")
def director_logout():
    session.pop("director", None)
    return redirect(url_for("app_page"))

# ---------- App polishing + API (org-aware credits) ----------
@app.post("/polish")
def polish():
    # Always reprocess (no caching)
    f = request.files.get("cv")
    if not f:
        abort(400, "No file uploaded")

    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / f.filename
        f.save(str(p))

        text = extract_text_any(p)
        if not text or len(text.strip()) < 30:
            abort(400, "Couldn't read enough text. If it's a scanned PDF, please use a DOCX or an OCRed PDF.")

        # --- Pre-check credits (org-aware). Admin bypasses. ---
        try:
            uid_check = int(session.get("user_id") or 0)
        except Exception:
            uid_check = 0

        try:
            can_bypass = (session.get("user","").strip().lower() == "admin") or bool(session.get("is_admin"))
        except Exception:
            can_bypass = False

        if DB_POOL and uid_check > 0 and not can_bypass:
            # If the user belongs to an org, check the org pool; otherwise check personal balance.
            try:
                org_id = _user_org_id(uid_check)
            except Exception:
                org_id = None

            try:
                if org_id:
                    bal = org_balance(org_id)
                    if bal <= 0:
                        raise PaymentRequired("No credits remaining for your organization. Please top up to continue.")
                    # Optional per-user monthly cap (only applies if a cap is set)
                    cap = get_user_monthly_cap(org_id, uid_check)
                    if cap is not None:
                        spent = org_user_spent_this_month(org_id, uid_check)
                        if spent >= cap:
                            raise PaymentRequired("Your monthly polish limit has been reached. Ask your director to raise your cap.")
                else:
                    row = db_query_one("SELECT COALESCE(SUM(delta),0) FROM credits_ledger WHERE user_id=%s", (uid_check,))
                    bal = int(row[0]) if row else 0
                    if bal <= 0:
                        raise PaymentRequired("No credits remaining for this account. Please top up to continue.")
            except Exception as e:
                # If balance check fails, don't block polishing; just log
                print("credits precheck failed:", e)

        # ---- Polishing logic (unchanged) ----
        data = ai_or_heuristic_structuring(text)
        data["skills"] = extract_top_skills(text)  # keywords-only list as before

        # Optional per-org DOCX template (falls back to default if none)
        template_override = None
        try:
            oid = _current_user_org_id()
            if oid:
                row = db_query_one("SELECT template_path FROM orgs WHERE id=%s", (oid,))
                if row and row[0]:
                    pth = Path(row[0])
                    if pth.exists():
                        template_override = str(pth)
        except Exception as e:
            print("template resolve failed:", e)

        out = build_cv_document(data, template_override=template_override)

        # ---- Update legacy JSON stats (for continuity) ----

        # ---- Update legacy JSON stats (for continuity) ----
        candidate_name = (data.get("personal_info") or {}).get("full_name") or f.filename
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        STATS["downloads"] = int(STATS.get("downloads", 0)) + 1
        STATS["last_candidate"] = candidate_name
        STATS["last_time"] = now
        STATS.setdefault("history", [])
        STATS["history"].append({"candidate": candidate_name, "filename": f.filename, "ts": now})
        _save_stats()

        # --- Log usage + debit one org credit (best-effort; never blocks) ---
        try:
            uid = int(session.get("user_id") or 0)
            if uid:
                # record usage
                log_usage_event(uid, f.filename, candidate_name)

                # debit org pool unless admin bypass
                can_bypass = (session.get("user","").strip().lower() == "admin") or bool(session.get("is_admin"))
                if not can_bypass:
                    oid = _current_user_org_id()
                    if DB_POOL and oid:
                        db_execute(
                            "INSERT INTO org_credits_ledger (org_id, delta, reason, created_by) VALUES (%s, -1, %s, %s)",
                            (oid, 'polish', uid),
                        )
        except Exception as e:
            # Never block the download if this fails
            print("post-polish usage/credit write failed:", e)        

        # ---- Optional: decrement trial credits (legacy session) ----
        try:
            left = int(session.get("trial_credits", 0))
            if left > 0:
                session["trial_credits"] = max(0, left - 1)
        except Exception:
            pass

        # ---- Return the polished file ----
        resp = make_response(send_file(str(out), as_attachment=True, download_name="polished_cv.docx"))
        resp.headers["Cache-Control"] = "no-store"
        return resp
